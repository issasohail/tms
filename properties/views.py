from accounts.access import restrict_queryset_to_properties
import json
import logging
import os
from collections import defaultdict
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation
from io import BytesIO
from urllib.parse import quote

import fitz
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.messages.views import SuccessMessageMixin
from django.core import signing
from django.core.exceptions import ValidationError
from django.db import transaction
from django.db.models import (
    Count,
    DateField,
    Exists,
    IntegerField,
    OuterRef,
    Prefetch,
    Q,
    Subquery,
    Sum,
)
from django.db.models.functions import Coalesce
from django.http import FileResponse, Http404, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse, reverse_lazy
from django.utils import timezone
from django.views.decorators.http import require_POST
from django.views.generic import CreateView, DeleteView, DetailView, UpdateView
from django_filters.views import FilterView
from django_tables2 import SingleTableView
from django_tables2.paginators import LazyPaginator
from django_tables2.views import SingleTableMixin
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from core.models import GlobalSettings
from core.public_urls import build_public_url
from invoices.models import Invoice
from leases.models import Lease, LeaseRenewal, WhatsAppTemplate
from leases.models_parking_inventory import (
    InventoryItemDefinition,
    PropertyInventoryItem,
    UnitInventoryItem,
)
from leases.whatsapp import build_whatsapp_url, render_unit_whatsapp_template
from payments.models import Payment
from tenants.models import Tenant, TenantInterestType
from utils.pdf_export import handle_export

from .filters import UnitFilter
from .forms import PropertyBankAccountForm, PropertyForm, UnitForm
from .models import (
    BuildingType,
    Property,
    PropertyBankAccount,
    PropertyMedia,
    Unit,
    UnitMedia,
)
from .tables import PropertyTable, UnitTable

logger = logging.getLogger(__name__)
UNIT_MEDIA_SHARE_MAX_AGE = 60 * 60 * 48
UNIT_MEDIA_SHARE_SALT = "properties.unit-media-share"


class PropertyListView(SingleTableView):
    model = Property
    table_class = PropertyTable
    template_name = "properties/property_list.html"
    ordering = ["-created_at"]
    context_object_name = "properties"
    table_pagination = {"per_page": 15, "paginator_class": LazyPaginator}

    def get_queryset(self):
        return restrict_queryset_to_properties(Property.objects.all(), self.request.user, "")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        return context

    def get(self, request, *args, **kwargs):
        # Handle export requests first
        if request.GET.get("_export"):
            table = self.get_table()
            export_name = f"properties_{datetime.now().strftime('%Y%m%d')}"
            return handle_export(request, table, export_name)

        # Normal GET request
        return super().get(request, *args, **kwargs)


class PropertyCreateView(LoginRequiredMixin, SuccessMessageMixin, CreateView):
    model = Property
    form_class = PropertyForm
    template_name = "properties/property_form.html"
    success_message = "Property created successfully"
    success_url = reverse_lazy("properties:property_list")

    def form_valid(self, form):
        messages.success(self.request, "Property created successfully.")
        return super().form_valid(form)


class PropertyDetailView(LoginRequiredMixin, DetailView):
    model = Property
    template_name = "properties/property_detail.html"
    context_object_name = "property"
    success_url = reverse_lazy("properties:property_list")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        from smart_meter.rates import resolve_electricity_rate

        context["electricity_rate"] = resolve_electricity_rate(
            property_obj=self.object
        )
        today = timezone.now().date()
        ending_date = today + timedelta(days=40)
        active_lease = Lease.objects.filter(
            unit_id=OuterRef("pk"),
            start_date__lte=today,
            end_date__gte=today,
        ).exclude(status__in=["ended", "terminated"])
        active_lease_history = LeaseRenewal.objects.filter(
            lease__unit_id=OuterRef("pk"),
            start_date__lte=today,
            end_date__gte=today,
        )
        units = list(
            self.object.units.annotate(
                has_active_lease=Exists(active_lease),
                has_active_lease_history=Exists(active_lease_history),
                has_ending_soon_lease=Exists(
                    active_lease.filter(end_date__lte=ending_date)
                ),
                has_ending_soon_lease_history=Exists(
                    active_lease_history.filter(end_date__lte=ending_date)
                ),
                active_lease_end_date=Subquery(
                    active_lease.order_by("end_date", "id").values("end_date")[:1],
                    output_field=DateField(),
                ),
                active_lease_history_end_date=Subquery(
                    active_lease_history.order_by("end_date", "id").values("end_date")[
                        :1
                    ],
                    output_field=DateField(),
                ),
            ).order_by("unit_number")
        )
        active_unit_ids = set(
            Lease.objects.filter(
                unit__property=self.object,
                status="active",
                start_date__lte=today,
                end_date__gte=today,
            )
            .order_by()
            .values_list("unit_id", flat=True)
            .distinct()
        )
        actual_total_units = len(units)
        occupied_units_count = sum(1 for unit in units if unit.id in active_unit_ids)
        vacant_units_count = sum(
            1
            for unit in units
            if unit.status == "vacant" and unit.id not in active_unit_ids
        )
        maintenance_units_count = sum(
            1 for unit in units if unit.status == "maintenance"
        )

        context["units"] = units
        context["actual_total_units"] = actual_total_units
        context["configured_total_units"] = self.object.total_units
        context["occupied_units_count"] = occupied_units_count
        context["vacant_units_count"] = vacant_units_count
        context["maintenance_units_count"] = maintenance_units_count
        context["media_files"] = self.object.media_files.filter(is_active=True).only(
            "id",
            "property_id",
            "file",
            "file_type",
            "description",
            "stamped_file",
            "thumbnail",
            "original_filename",
            "sort_order",
            "uploaded_at",
        )[:6]
        from leases.services.inventory_parking import (
            effective_inventory,
            effective_parking_policy,
            policy_value,
        )

        parking_policy = effective_parking_policy(property_obj=self.object)
        context["effective_inventory"] = effective_inventory(property_obj=self.object)
        context["parking_enabled"] = bool(policy_value(parking_policy, "enabled"))
        context["parking_monthly_rate"] = policy_value(parking_policy, "monthly_rate")
        context["parking_penalty"] = policy_value(
            parking_policy, "unauthorized_parking_penalty"
        )
        context["parking_spaces"] = self.object.parking_spaces.filter(is_active=True)
        bank_accounts = list(self.object.bank_accounts.all())
        for account in bank_accounts:
            account.edit_form = PropertyBankAccountForm(
                instance=account, prefix=f"bank-{account.pk}"
            )
        context["bank_accounts"] = bank_accounts
        context["bank_account_form"] = PropertyBankAccountForm(prefix="new-bank")
        return context


@login_required
@require_POST
def property_bank_account_save(request, property_pk, account_pk=None):
    property_obj = get_object_or_404(Property, pk=property_pk)
    account = None
    if account_pk is not None:
        account = get_object_or_404(
            PropertyBankAccount, pk=account_pk, property=property_obj
        )
    prefix = f"bank-{account.pk}" if account else "new-bank"
    form = PropertyBankAccountForm(request.POST, instance=account, prefix=prefix)
    if form.is_valid():
        saved = form.save(commit=False)
        saved.property = property_obj
        saved.save()
        messages.success(request, f"Bank account {saved.account_label} saved.")
    else:
        messages.error(
            request, "Bank account could not be saved. Check the entered values."
        )
    return redirect("properties:property_detail", pk=property_pk)


@login_required
@require_POST
def property_bank_account_welcome_settings(request, property_pk):
    property_obj = get_object_or_404(Property, pk=property_pk)
    mode = request.POST.get("welcome_bank_account_mode")
    if mode not in {
        Property.WELCOME_BANK_SELECTED,
        Property.WELCOME_BANK_ALL,
    }:
        messages.error(request, "Select a valid welcome-message bank account option.")
        return redirect("properties:property_detail", pk=property_pk)

    if mode == Property.WELCOME_BANK_SELECTED:
        selected = PropertyBankAccount.objects.filter(
            pk=request.POST.get("selected_bank_account"),
            property=property_obj,
            is_active=True,
        ).first()
        if selected is None:
            messages.error(
                request, "Add or activate a bank account before selecting it."
            )
            return redirect("properties:property_detail", pk=property_pk)
        PropertyBankAccount.objects.filter(property=property_obj).update(
            is_default=False
        )
        selected.is_default = True
        selected.save(update_fields=["is_default", "updated_at"])
    property_obj.welcome_bank_account_mode = mode
    property_obj.save(update_fields=["welcome_bank_account_mode", "updated_at"])
    messages.success(request, "Welcome-message bank account selection updated.")
    return redirect("properties:property_detail", pk=property_pk)


class PropertyUpdateView(LoginRequiredMixin, SuccessMessageMixin, UpdateView):
    model = Property
    form_class = PropertyForm
    template_name = "properties/property_form.html"
    success_message = "Property updated successfully"
    success_url = reverse_lazy("properties:property_list")

    def form_valid(self, form):
        messages.success(self.request, "Property updated successfully.")
        return super().form_valid(form)


class PropertyDeleteView(LoginRequiredMixin, DeleteView):
    model = Property
    template_name = "properties/property_confirm_delete.html"
    success_url = reverse_lazy("properties:property_list")

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "Property deleted successfully.")
        return super().delete(request, *args, **kwargs)


class UnitListView(SingleTableMixin, FilterView):
    model = Unit
    table_class = UnitTable
    template_name = "properties/unit_list.html"
    filterset_class = UnitFilter
    table_pagination = {"per_page": 60, "paginator_class": LazyPaginator}

    def _get_building_types(self):
        if not hasattr(self, "_building_types_cache"):
            self._building_types_cache = list(
                BuildingType.objects.filter(is_active=True).order_by(
                    "sort_order", "name"
                )
            )
        return self._building_types_cache

    def _get_inventory_definitions(self):
        if not hasattr(self, "_inventory_definitions_cache"):
            self._inventory_definitions_cache = list(
                InventoryItemDefinition.objects.filter(is_active=True).order_by(
                    "sort_order", "name"
                )
            )
        return self._inventory_definitions_cache

    def get_queryset(self):
        today = timezone.now().date()
        ending_date = today + timedelta(days=40)
        active_lease = Lease.objects.filter(
            unit_id=OuterRef("pk"),
            start_date__lte=today,
            end_date__gte=today,
        ).exclude(status__in=["ended", "terminated"])
        active_lease_history = LeaseRenewal.objects.filter(
            lease__unit_id=OuterRef("pk"),
            start_date__lte=today,
            end_date__gte=today,
        )
        active_lease_end = active_lease.order_by("end_date", "id").values("end_date")[
            :1
        ]
        active_lease_id = active_lease.order_by("end_date", "id").values("id")[:1]
        active_lease_history_end = active_lease_history.order_by(
            "end_date", "id"
        ).values("end_date")[:1]
        active_lease_history_lease_id = active_lease_history.order_by(
            "end_date", "id"
        ).values("lease_id")[:1]
        ending_soon_lease = active_lease.filter(end_date__lte=ending_date)
        ending_soon_lease_history = active_lease_history.filter(
            end_date__lte=ending_date
        )
        queryset = (
            super()
            .get_queryset()
            .select_related("property", "building_type", "interest_type")
            .only(
                "id",
                "property_id",
                "building_type_id",
                "interest_type_id",
                "unit_number",
                "monthly_rent",
                "electric_meter_num",
                "gas_meter_num",
                "society_maintenance",
                "water_charges",
                "internet_charges",
                "is_smart_meter",
                "security_requires",
                "security_deposit_amount",
                "show_publicly",
                "status",
                "property__id",
                "property__property_name",
                "building_type__id",
                "building_type__name",
                "building_type__code",
                "building_type__is_active",
                "building_type__sort_order",
                "interest_type__id",
                "interest_type__name",
                "interest_type__code",
                "interest_type__is_active",
                "interest_type__sort_order",
            )
            .prefetch_related(
                Prefetch(
                    "property__inventory_items",
                    queryset=PropertyInventoryItem.objects.select_related("item"),
                    to_attr="unit_list_inventory",
                ),
                Prefetch(
                    "inventory_items",
                    queryset=UnitInventoryItem.objects.select_related("item"),
                    to_attr="unit_list_inventory",
                ),
            )
            .annotate(
                has_active_lease=Exists(active_lease),
                has_active_lease_history=Exists(active_lease_history),
                has_ending_soon_lease=Exists(ending_soon_lease),
                has_ending_soon_lease_history=Exists(ending_soon_lease_history),
                active_lease_end_date=Subquery(
                    active_lease_end, output_field=DateField()
                ),
                active_lease_id=Subquery(active_lease_id, output_field=IntegerField()),
                active_lease_history_end_date=Subquery(
                    active_lease_history_end, output_field=DateField()
                ),
                active_lease_history_lease_id=Subquery(
                    active_lease_history_lease_id, output_field=IntegerField()
                ),
            )
        )
        queryset = restrict_queryset_to_properties(queryset, self.request.user, "property")

        property_id = self.request.GET.get("property")
        if property_id:
            queryset = queryset.filter(property_id=property_id)
        status = self.request.GET.get("status")
        if status == "vacant":
            queryset = queryset.filter(
                has_active_lease=False,
                has_active_lease_history=False,
            ).exclude(status="maintenance")
        elif status == "occupied":
            queryset = queryset.filter(
                Q(has_active_lease=True) | Q(has_active_lease_history=True)
            )
        elif status == "ending_soon":
            queryset = queryset.filter(
                Q(has_ending_soon_lease=True) | Q(has_ending_soon_lease_history=True)
            )
        elif status == "maintenance":
            queryset = queryset.filter(status="maintenance")
        return queryset

    def get_table_data(self):
        return self.object_list

    def get_table_kwargs(self):
        kwargs = super().get_table_kwargs()
        kwargs["building_types"] = self._get_building_types()
        kwargs["inventory_definitions"] = self._get_inventory_definitions()
        return kwargs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["all_properties"] = (
            Property.objects.only("id", "property_name")
            .annotate(unit_count=Count("units"))
            .order_by("property_name")
        )
        context["current_status"] = self.request.GET.get("status", "")
        context["building_type_options"] = [
            {"id": building_type.pk, "name": building_type.name}
            for building_type in self._get_building_types()
        ]
        return context

    def get(self, request, *args, **kwargs):
        # Handle export requests
        if request.GET.get("_export"):
            export_response = self.handle_export(request)
            if export_response:
                return export_response
        return super().get(request, *args, **kwargs)

    def handle_export(self, request):
        """Handle export functionality"""
        self.object_list = self.get_queryset()
        table = self.get_table()
        export_name = "units_list"
        return handle_export(request, table, export_name)


def _unit_has_current_lease(unit, today=None):
    today = today or timezone.now().date()
    return (
        LeaseRenewal.objects.filter(
            lease__unit=unit,
            start_date__lte=today,
            end_date__gte=today,
        ).exists()
        or Lease.objects.filter(
            unit=unit,
            start_date__lte=today,
            end_date__gte=today,
        )
        .exclude(status__in=["ended", "terminated"])
        .exists()
    )


def _attach_unit_occupancy(unit):
    today = timezone.now().date()
    active_lease_history = (
        LeaseRenewal.objects.select_related("lease", "lease__tenant", "lease__unit")
        .filter(
            lease__unit=unit,
            start_date__lte=today,
            end_date__gte=today,
        )
        .order_by("-renewal_number", "-start_date", "-pk")
        .first()
    )
    current_lease = None
    if active_lease_history:
        current_lease = active_lease_history.lease
    else:
        current_lease = (
            Lease.objects.select_related("tenant", "unit")
            .filter(
                unit=unit,
                start_date__lte=today,
                end_date__gte=today,
            )
            .exclude(status__in=["ended", "terminated"])
            .order_by("-start_date", "-pk")
            .first()
        )

    unit.active_lease_history = active_lease_history
    unit.current_lease = current_lease
    unit.current_tenant = current_lease.tenant if current_lease else None
    unit.has_active_lease_history = active_lease_history is not None
    unit.has_active_lease = current_lease is not None and active_lease_history is None
    unit.is_currently_occupied = current_lease is not None
    return unit


def _unit_lease_history_rows(unit):
    leases = list(
        Lease.objects.filter(unit=unit)
        .select_related("tenant")
        .order_by("-start_date", "-pk")
    )
    if not leases:
        return []

    lease_ids = [lease.id for lease in leases]

    renewals_by_lease = defaultdict(list)
    for renewal in LeaseRenewal.objects.filter(lease_id__in=lease_ids).order_by(
        "renewal_number"
    ):
        renewals_by_lease[renewal.lease_id].append(renewal)

    invoice_totals = {
        row["lease_id"]: row["total"] or Decimal("0.00")
        for row in (
            Invoice.objects.filter(lease_id__in=lease_ids)
            .values("lease_id")
            .annotate(total=Coalesce(Sum("amount"), Decimal("0.00")))
        )
    }

    # Lease ledger balance must use the payment detail's lease_amount (the
    # portion of a payment actually applied to rent/invoices), not the raw
    # payment.amount, since a payment can also cover security deposit etc.
    payment_totals = defaultdict(lambda: Decimal("0.00"))
    for payment in Payment.objects.filter(lease_id__in=lease_ids).select_related(
        "detail"
    ):
        detail = getattr(payment, "detail", None)
        lease_amt = getattr(detail, "lease_amount", None) if detail else None
        payment_totals[payment.lease_id] += (
            lease_amt if lease_amt is not None else (payment.amount or Decimal("0.00"))
        )

    current_lease_id = unit.current_lease.pk if unit.current_lease else None

    rows = []
    for index, lease in enumerate(leases, start=1):
        renewals = renewals_by_lease.get(lease.id, [])
        end_date = renewals[-1].end_date if renewals else lease.end_date
        balance = invoice_totals.get(lease.id, Decimal("0.00")) - payment_totals.get(
            lease.id, Decimal("0.00")
        )
        rows.append(
            {
                "sn": index,
                "lease": lease,
                "tenant": lease.tenant,
                "start_date": lease.start_date,
                "end_date": end_date,
                "renewal_count": len(renewals),
                "balance": balance,
                "is_current": lease.id == current_lease_id,
            }
        )
    return rows


def _unit_detail_context(unit):
    _attach_unit_occupancy(unit)
    current_meter_rows, meter_history_rows = _unit_meter_rows(unit)
    return {
        "unit": unit,
        "units": unit,
        "media_files": unit.media_files.filter(is_active=True).only(
            "id",
            "unit_id",
            "file",
            "file_type",
            "description",
            "stamped_file",
            "thumbnail",
            "original_filename",
            "sort_order",
            "uploaded_at",
        )[:6],
        "current_lease": unit.current_lease,
        "current_tenant": unit.current_tenant,
        "active_lease_history": unit.active_lease_history,
        "current_meter_rows": current_meter_rows,
        "meter_history_rows": meter_history_rows,
        "lease_history_rows": _unit_lease_history_rows(unit),
    }


def _unit_meter_rows(unit):
    from smart_meter.models import (
        LiveReading,
        Meter,
        MeterAssignmentHistory,
        MeterInstallation,
        MeterReading,
    )

    latest_reading = MeterReading.objects.filter(
        meter_id=OuterRef("meter_id")
    ).order_by("-ts", "-id")
    installation_rows = list(
        MeterInstallation.objects.filter(unit=unit)
        .select_related("meter")
        .annotate(
            latest_reading_ts=Subquery(latest_reading.values("ts")[:1]),
            latest_reading_energy=Subquery(
                latest_reading.values("total_energy")[:1]
            ),
        )
        .order_by("-is_active", "-start_date", "-id")
    )
    history_rows = []
    seen_meter_ids = set()

    for installation in installation_rows:
        last_ts = installation.latest_reading_ts
        last_energy = installation.latest_reading_energy
        display_is_active = (
            installation.is_active
            and installation.end_date is None
            and installation.meter.is_active
            and installation.meter.unit_id == unit.id
        )
        display_end_date = installation.end_date
        if not display_is_active and display_end_date is None and last_ts:
            display_end_date = timezone.localtime(last_ts).date()
        display_end_reading = installation.end_reading
        if display_end_reading is None:
            display_end_reading = last_energy
        seen_meter_ids.add(installation.meter_id)
        history_rows.append(
            {
                "meter": installation.meter,
                "start_date": installation.start_date,
                "end_date": display_end_date,
                "start_reading": installation.start_reading,
                "end_reading": display_end_reading,
                "reason": installation.reason or "installation",
                "is_active": display_is_active,
            }
        )

    latest_assignment = MeterAssignmentHistory.objects.filter(
        meter_id=OuterRef("pk"), new_unit=unit
    ).order_by("-change_date", "-id")
    latest_meter_reading = MeterReading.objects.filter(
        meter_id=OuterRef("pk")
    ).order_by("-ts", "-id")
    latest_live_energy = LiveReading.objects.filter(
        meter_id=OuterRef("pk")
    ).values("total_energy")[:1]

    assignment_start_reading = MeterReading.objects.filter(
        meter_id=OuterRef("pk"),
        ts__lte=OuterRef("latest_assignment_date"),
    ).order_by("-ts", "-id")
    cached_unit_meters = list(
        Meter.objects.filter(unit=unit)
        .select_related("unit")
        .annotate(
            latest_assignment_date=Subquery(
                latest_assignment.values("change_date")[:1]
            ),
            assignment_start_reading=Subquery(
                assignment_start_reading.values("total_energy")[:1]
            ),
            latest_reading_ts=Subquery(latest_meter_reading.values("ts")[:1]),
            latest_reading_energy=Subquery(
                latest_meter_reading.values("total_energy")[:1]
            ),
            latest_live_energy=Subquery(latest_live_energy),
        )
        .order_by("-is_active", "-installed_at", "meter_number")
    )

    for meter in cached_unit_meters:
        if meter.pk in seen_meter_ids:
            continue
        assignment_date = meter.latest_assignment_date
        if assignment_date:
            start_date = timezone.localtime(assignment_date).date()
            start_reading = meter.assignment_start_reading
            reason = "assignment"
        else:
            start_date = (
                timezone.localtime(meter.installed_at).date()
                if meter.installed_at
                else None
            )
            start_reading = None
            reason = "cached unit"

        display_is_active = bool(meter.is_active and meter.unit_id == unit.id)
        history_rows.append(
            {
                "meter": meter,
                "start_date": start_date,
                "end_date": None
                if display_is_active
                else (
                    timezone.localtime(meter.latest_reading_ts).date()
                    if meter.latest_reading_ts
                    else None
                ),
                "start_reading": start_reading,
                "end_reading": meter.latest_live_energy
                if meter.latest_live_energy is not None
                else meter.latest_reading_energy,
                "reason": reason,
                "is_active": display_is_active,
            }
        )

    history_rows.sort(
        key=lambda row: (
            row["is_active"],
            row["start_date"] or timezone.datetime.min.date(),
        ),
        reverse=True,
    )
    current_rows = [
        row
        for row in history_rows
        if row["is_active"]
        and row["meter"].unit_id == unit.id
        and row["meter"].is_active
    ]
    return current_rows, history_rows


def _can_manage_media(user, owner):
    if not user.is_authenticated:
        return False
    if user.is_superuser:
        return True
    if isinstance(owner, Unit):
        return (
            user.has_perm("properties.change_unit")
            or user.has_perm("properties.change_unitmedia")
            or user.has_perm("properties.delete_unitmedia")
        )
    return (
        user.has_perm("properties.change_property")
        or user.has_perm("properties.change_propertymedia")
        or user.has_perm("properties.delete_propertymedia")
    )


def _media_permission_denied():
    return JsonResponse({"success": False, "error": "Permission denied"}, status=403)


def _lead_money(value):
    if value in (None, ""):
        return "-"
    try:
        return f"{value:,.0f}"
    except Exception:
        return str(value)


def _default_unit_interest_type(unit):
    linked_interest_type = getattr(
        getattr(unit, "building_type", None), "lead_interest_type", None
    )
    if linked_interest_type:
        return linked_interest_type
    property_name = (unit.property.property_name if unit.property else "").lower()
    code = (
        "single_room_attached_bath_kitchen"
        if "f56" in property_name and "basement" in property_name
        else "two_room_flat"
    )
    return TenantInterestType.objects.filter(code=code, is_active=True).first()


def _vacant_notice_message(request, unit, tenant, photos_link=""):
    details_link = photos_link or build_public_url(
        "properties:unit_detail", args=[unit.pk]
    )
    lines = [
        f"Dear {tenant.get_full_name()},",
        "",
        "A unit is now available for rent:",
        "",
        f"Property: {unit.property.property_name}- {unit.unit_number}",
        f"Rent: {_lead_money(unit.monthly_rent)}",
        f"Maintenance: {_lead_money(unit.society_maintenance)}",
        f"Water: {_lead_money(unit.water_charges)}",
        f"Security Deposit: {unit.security_requires or '-'}",
        f"Agreement Fee: {getattr(unit, 'agreement_fee', None) or getattr(unit, 'agreement_charges', None) or '-'}",
        f"Notes: {unit.comments or '-'}",
        "",
        "Photos/Details:",
        details_link,
    ]
    lines.extend(["", "Please contact us if interested."])
    return "\n".join(lines)


@login_required
def unit_vacant_notice_leads(request, pk):
    today = timezone.now().date()
    unit = get_object_or_404(
        Unit.objects.select_related(
            "property", "building_type__lead_interest_type", "interest_type"
        ),
        pk=pk,
    )
    has_active_lease = (
        Lease.objects.filter(
            unit=unit,
            start_date__lte=today,
            end_date__gte=today,
        )
        .exclude(status__in=["ended", "terminated"])
        .exists()
    )
    has_active_lease_history = LeaseRenewal.objects.filter(
        lease__unit=unit,
        start_date__lte=today,
        end_date__gte=today,
    ).exists()
    unit_interest_type = unit.interest_type or _default_unit_interest_type(unit)
    if not unit_interest_type:
        return JsonResponse(
            {
                "success": True,
                "unit": {
                    "id": unit.pk,
                    "title": f"{unit.property.property_name} - {unit.unit_number} Vacant Notice",
                    "interest_type": None,
                },
                "interest_types": list(
                    TenantInterestType.objects.filter(is_active=True)
                    .order_by("sort_order", "name")
                    .values("id", "name")
                ),
                "tenants": [],
            }
        )

    settings_obj = GlobalSettings.get_solo()
    country_code = getattr(settings_obj, "country_code", "+92")
    has_photos = unit.media_files.filter(is_active=True).exists()
    photos_link = ""
    if has_photos:
        photos_link = build_public_url(
            "properties:unit_media_public_share",
            args=[_sign_unit_media_token(unit.pk)],
        )

    tenants = (
        Tenant.objects.prefetch_related("interested_in")
        .filter(interested_in=unit_interest_type)
        .distinct()
        .order_by("first_name", "last_name")
    )
    interest_types = list(
        TenantInterestType.objects.filter(is_active=True)
        .order_by("sort_order", "name")
        .values("id", "name")
    )
    rows = []
    for tenant in tenants:
        interests = [
            {"id": item.pk, "name": item.name} for item in tenant.interested_in.all()
        ]
        message = _vacant_notice_message(request, unit, tenant, photos_link=photos_link)
        rows.append(
            {
                "id": tenant.pk,
                "full_name": tenant.get_full_name(),
                "phone": tenant.phone or "",
                "photo_url": tenant.photo.url if tenant.photo else None,
                "interested_in": interests,
                "tenant_detail_url": reverse("tenants:tenant_detail", args=[tenant.pk]),
                "tenant_edit_url": reverse("tenants:tenant_update", args=[tenant.pk]),
                "tenant_inline_update_url": reverse(
                    "tenants:tenant_lead_inline_update", args=[tenant.pk]
                ),
                "whatsapp_message": message,
                "whatsapp_url": build_whatsapp_url(
                    tenant.phone, message, country_code=country_code
                ),
            }
        )

    return JsonResponse(
        {
            "success": True,
            "unit": {
                "id": unit.pk,
                "title": f"{unit.property.property_name} - {unit.unit_number} Vacant Notice",
                "interest_type": {
                    "id": unit_interest_type.pk,
                    "name": unit_interest_type.name,
                },
                "detail_url": build_public_url(
                    "properties:unit_detail", args=[unit.pk]
                ),
                "photos_url": photos_link or None,
                "has_active_lease": has_active_lease,
                "has_active_lease_history": has_active_lease_history,
            },
            "interest_types": interest_types,
            "tenants": rows,
        }
    )


@login_required
def unit_vacant_summary_message(request):
    today = timezone.now().date()
    ending_date = today + timedelta(days=40)
    property_id = request.GET.get("property")

    active_lease = Lease.objects.filter(
        unit_id=OuterRef("pk"),
        start_date__lte=today,
        end_date__gte=today,
    ).exclude(status__in=["ended", "terminated"])
    active_lease_history = LeaseRenewal.objects.filter(
        lease__unit_id=OuterRef("pk"),
        start_date__lte=today,
        end_date__gte=today,
    )
    vacant_units = (
        Unit.objects.select_related("property")
        .annotate(
            has_active_lease=Exists(active_lease),
            has_active_lease_history=Exists(active_lease_history),
        )
        .filter(has_active_lease=False, has_active_lease_history=False)
        .exclude(status="maintenance")
        .order_by("property__property_name", "unit_number")
    )
    if property_id:
        vacant_units = vacant_units.filter(property_id=property_id)

    ending_histories = (
        LeaseRenewal.objects.select_related(
            "lease", "lease__tenant", "lease__unit", "lease__unit__property"
        )
        .filter(start_date__lte=today, end_date__gte=today, end_date__lte=ending_date)
        .order_by(
            "end_date",
            "lease__unit__property__property_name",
            "lease__unit__unit_number",
        )
    )
    if property_id:
        ending_histories = ending_histories.filter(lease__unit__property_id=property_id)

    history_lease_ids = set(ending_histories.values_list("lease_id", flat=True))
    ending_leases = (
        Lease.objects.select_related("tenant", "unit", "unit__property")
        .filter(start_date__lte=today, end_date__gte=today, end_date__lte=ending_date)
        .exclude(status__in=["ended", "terminated"])
        .exclude(id__in=history_lease_ids)
        .order_by("end_date", "unit__property__property_name", "unit__unit_number")
    )
    if property_id:
        ending_leases = ending_leases.filter(unit__property_id=property_id)

    settings_obj = GlobalSettings.get_solo()
    country_code = getattr(settings_obj, "country_code", "+92")
    user_phone = getattr(request.user, "whatsapp_number", "") or getattr(
        settings_obj, "whatsapp_number", ""
    )

    lines = [
        "Vacant and Ending Soon Unit Summary",
        f"Date: {today:%Y-%m-%d}",
        "",
        "Vacant units:",
    ]
    vacant_list = list(vacant_units)
    if vacant_list:
        for index, unit in enumerate(vacant_list, start=1):
            lines.append(
                f"{index}. Property: {unit.property.property_name} | Unit: {unit.unit_number} | Status: Vacant | Rent: {_lead_money(unit.monthly_rent)} | Maintenance: {_lead_money(unit.society_maintenance)}"
            )
    else:
        lines.append("1. None")

    lines.extend(["", "Leases ending within 40 days:"])
    ending_rows = []
    for history in ending_histories:
        lease = history.lease
        ending_rows.append(
            (
                history.end_date,
                lease.unit.property.property_name,
                lease.unit.unit_number,
                _lead_money(lease.unit.monthly_rent),
                _lead_money(lease.unit.society_maintenance),
            )
        )
    for lease in ending_leases:
        ending_rows.append(
            (
                lease.end_date,
                lease.unit.property.property_name,
                lease.unit.unit_number,
                _lead_money(lease.unit.monthly_rent),
                _lead_money(lease.unit.society_maintenance),
            )
        )
    ending_rows.sort(key=lambda row: (row[0], row[1], row[2]))
    if ending_rows:
        for index, (
            end_date,
            property_name,
            unit_number,
            rent,
            maintenance,
        ) in enumerate(ending_rows, start=1):
            lines.append(
                f"{index}. Property: {property_name} | Unit: {unit_number} | Status: Lease ending soon ({end_date:%Y-%m-%d}) | Rent: {rent} | Maintenance: {maintenance}"
            )
    else:
        lines.append("1. None")

    message = "\n".join(lines)
    whatsapp_url = (
        build_whatsapp_url(user_phone, message, country_code=country_code)
        if user_phone
        else f"https://wa.me/?text={quote(message)}"
    )
    return JsonResponse(
        {
            "success": True,
            "message": message,
            "whatsapp_url": whatsapp_url,
            "vacant_count": len(vacant_list),
            "ending_soon_count": len(ending_rows),
        }
    )


class UnitDetailView(LoginRequiredMixin, DetailView):
    model = Unit
    template_name = "properties/unit_detail.html"
    context_object_name = "unit"

    def get_queryset(self):
        return (
            super()
            .get_queryset()
            .select_related("property", "building_type", "interest_type")
            .prefetch_related(
                Prefetch(
                    "property__inventory_items",
                    queryset=PropertyInventoryItem.objects.select_related("item"),
                ),
                Prefetch(
                    "inventory_items",
                    queryset=UnitInventoryItem.objects.select_related("item"),
                ),
            )
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context.update(_unit_detail_context(self.object))
        from smart_meter.rates import resolve_electricity_rate

        context["electricity_rate"] = resolve_electricity_rate(unit=self.object)
        from leases.services.inventory_parking import (
            effective_inventory,
            effective_parking_policy,
            policy_value,
        )

        parking_policy = effective_parking_policy(unit=self.object)
        context["effective_inventory"] = effective_inventory(unit=self.object)
        context["parking_enabled"] = bool(policy_value(parking_policy, "enabled"))
        context["parking_monthly_rate"] = policy_value(parking_policy, "monthly_rate")
        context["parking_penalty"] = policy_value(
            parking_policy, "unauthorized_parking_penalty"
        )
        return context


def unit_detail(request, pk):
    unit = get_object_or_404(
        Unit.objects.select_related("property", "building_type", "interest_type"),
        pk=pk,
    )
    from smart_meter.rates import resolve_electricity_rate

    context = _unit_detail_context(unit)
    context["electricity_rate"] = resolve_electricity_rate(unit=unit)
    return render(
        request,
        "properties/unit_detail.html",
        context,
    )


class UnitCreateView(CreateView):
    model = Unit
    form_class = UnitForm
    template_name = "properties/unit_form.html"

    def get_success_url(self):
        messages.success(self.request, "Unit created successfully.")
        return reverse("properties:unit_detail", args=[self.object.pk])


class UnitUpdateView(UpdateView):
    model = Unit
    form_class = UnitForm
    template_name = "properties/unit_form.html"

    def get_success_url(self):
        messages.success(self.request, "Unit updated successfully.")
        return reverse("properties:unit_detail", args=[self.object.pk])


class UnitDeleteView(DeleteView):
    model = Unit
    template_name = "properties/unit_confirm_delete.html"

    def get_success_url(self):
        messages.success(self.request, "Unit deleted successfully.")
        return reverse("properties:unit_list")


@login_required
@require_POST
def unit_inline_update(request):
    try:
        data = json.loads(request.body)
        unit_id = data.get("id")
        field = data.get("field")
        value = data.get("value")

        allowed_fields = {
            "building_type",
            "interest_type",
            "monthly_rent",
            "society_maintenance",
            "water_charges",
            "internet_charges",
            "security_requires",
            "security_deposit_amount",
            "status",
            "show_publicly",
            "is_smart_meter",
        }

        if field not in allowed_fields:
            return JsonResponse(
                {"success": False, "error": "This field cannot be edited inline."},
                status=400,
            )

        unit = get_object_or_404(Unit, pk=unit_id)

        if field == "building_type":
            if value:
                building_type = get_object_or_404(
                    BuildingType, pk=value, is_active=True
                )
                unit.building_type = building_type
                unit.interest_type = getattr(building_type, "lead_interest_type", None)
                new_value = building_type.name
            else:
                unit.building_type = None
                unit.interest_type = None
                new_value = ""
            unit.save(update_fields=["building_type", "interest_type"])
            return JsonResponse(
                {
                    "success": True,
                    "new_value": new_value,
                    "building_type_id": unit.building_type_id or "",
                }
            )

        if field == "interest_type":
            if value:
                interest_type = get_object_or_404(
                    TenantInterestType, pk=value, is_active=True
                )
                unit.interest_type = interest_type
                new_value = interest_type.name
            else:
                unit.interest_type = None
                new_value = ""
            unit.save(update_fields=["interest_type"])
            return JsonResponse(
                {
                    "success": True,
                    "new_value": new_value,
                    "interest_type_id": unit.interest_type_id or "",
                }
            )

        if field in {"show_publicly", "is_smart_meter"}:
            bool_value = str(value).lower() in {"true", "1", "yes", "on"}
            setattr(unit, field, bool_value)
            unit.save(update_fields=[field])
            return JsonResponse(
                {
                    "success": True,
                    "new_value": "Yes" if bool_value else "No",
                    "value": bool_value,
                }
            )

        if field == "status":
            valid_statuses = dict(Unit.UNIT_STATUS)
            if value not in valid_statuses:
                return JsonResponse(
                    {"success": False, "error": "Please choose a valid unit status."},
                    status=400,
                )
            unit.status = value
            unit.save(update_fields=["status"])
            return JsonResponse(
                {
                    "success": True,
                    "new_value": valid_statuses[value],
                    "value": value,
                }
            )

        bulk_editable_fields = {
            "monthly_rent",
            "society_maintenance",
            "water_charges",
            "internet_charges",
            "security_requires",
            "security_deposit_amount",
        }

        if field in bulk_editable_fields:
            apply_to_property = str(data.get("apply_to_property", "")).lower() in {
                "true", "1", "yes", "on"
            }
            if field == "security_requires":
                cleaned_value = str(value or "").strip()
                if len(cleaned_value) > 20:
                    return JsonResponse(
                        {
                            "success": False,
                            "error": "Security text cannot exceed 20 characters.",
                        },
                        status=400,
                    )
            else:
                raw_value = str(value or "0").replace(",", "").strip()
                try:
                    cleaned_value = Decimal(raw_value or "0").quantize(
                        Decimal("0.01")
                    )
                except (InvalidOperation, ValueError):
                    return JsonResponse(
                        {"success": False, "error": "Please enter a valid amount."},
                        status=400,
                    )
                if cleaned_value < 0 or cleaned_value > Decimal("99999999.99"):
                    return JsonResponse(
                        {
                            "success": False,
                            "error": "Amount must be between 0 and 99,999,999.99.",
                        },
                        status=400,
                    )

            if apply_to_property:
                affected_count = Unit.objects.filter(
                    property_id=unit.property_id
                ).update(**{field: cleaned_value})
            else:
                setattr(unit, field, cleaned_value)
                unit.save(update_fields=[field])
                affected_count = 1

            new_value = (
                _lead_money(cleaned_value)
                if field != "security_requires"
                else cleaned_value
            )
            property_unit_count = unit.property.units.count()
            return JsonResponse(
                {
                    "success": True,
                    "new_value": new_value or "",
                    "field": field,
                    "property_id": unit.property_id,
                    "property_name": unit.property.property_name,
                    "property_unit_count": property_unit_count,
                    "affected_count": affected_count,
                    "bulk_applied": apply_to_property,
                    "offer_bulk": not apply_to_property and property_unit_count > 1,
                }
            )

    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)}, status=400)


def _next_media_sort(qs):
    return qs.count() + 1


def _upload_media_files(request, owner, media_model, owner_field):
    files = request.FILES.getlist("files") or request.FILES.getlist("photos")
    description = (request.POST.get("description") or "").strip()[:300]
    if not files:
        messages.error(request, "Please choose at least one file.")
        return 0

    created = 0
    active_qs = owner.media_files.filter(is_active=True)
    serial_start = owner.media_files.count() + 1
    for file_obj in files:
        media = media_model(
            **{owner_field: owner},
            file=file_obj,
            description=description,
            sort_order=_next_media_sort(active_qs) + created,
            uploaded_by=request.user if request.user.is_authenticated else None,
            original_filename=getattr(file_obj, "name", "")[:255],
        )
        media._media_serial = serial_start + created
        try:
            media.full_clean()
            media.save()
            created += 1
        except ValidationError as exc:
            messages.error(
                request, f"{getattr(file_obj, 'name', 'File')}: {exc.messages[0]}"
            )
    if created:
        messages.success(request, f"Uploaded {created} file(s).")
    return created


def _media_export_filename(owner_label, extension):
    safe_label = "".join(
        ch if ch.isalnum() or ch in "-_" else "-" for ch in owner_label
    ).strip("-")
    return f"{safe_label or 'photos'}-{timezone.now():%Y%m%d}.{extension}"


def _image_path(media):
    image_file = media.stamped_file or media.file
    try:
        return image_file.path
    except (NotImplementedError, ValueError):
        return None


def _draw_wrapped_text(pdf, text, x, y, max_chars=95, line_height=5 * mm):
    text = text or ""
    lines = [text[i : i + max_chars] for i in range(0, len(text), max_chars)] or [""]
    for line in lines[:3]:
        pdf.drawString(x, y, line)
        y -= line_height
    return y


def _public_media_url(request, owner_kind, owner_pk, media_id):
    token = _sign_media_token(owner_kind, owner_pk)
    return build_public_url(
        "properties:media_public_file", args=[token, media_id]
    )


def _media_link_rows(request, owner_kind, owner_pk, media_files):
    rows = []
    for index, media in enumerate(media_files, start=1):
        if media.file_type != "image":
            rows.append(
                (
                    index,
                    media.display_filename,
                    _public_media_url(request, owner_kind, owner_pk, media.pk),
                )
            )
    return rows


def _pdf_footer(pdf, title, page_num, total_pages, width):
    pdf.setFont("Helvetica", 8)
    y = 10 * mm
    pdf.drawCentredString(width / 2, y, title)
    pdf.drawRightString(
        width - 12 * mm,
        y,
        f"Page {page_num} of {total_pages}  {timezone.localtime():%Y-%m-%d %H:%M}",
    )


def _draw_photo_page(pdf, media_items, title, page_num, total_pages, pagesize):
    width, height = pagesize
    pdf.setFont("Helvetica-Bold", 13)
    pdf.drawCentredString(width / 2, height - 14 * mm, title)
    cols = 2 if len(media_items) > 1 else 1
    rows = 2 if len(media_items) > 2 else len(media_items)
    if len(media_items) == 2:
        rows = 2
        cols = 1
    margin = 14 * mm
    top = height - 24 * mm
    bottom = 18 * mm
    gap = 6 * mm
    cell_w = (width - (2 * margin) - ((cols - 1) * gap)) / cols
    cell_h = (top - bottom - ((rows - 1) * gap)) / rows if rows else top - bottom

    for idx, media in enumerate(media_items):
        col = idx % cols
        row = idx // cols
        x = margin + col * (cell_w + gap)
        y_top = top - row * (cell_h + gap)
        caption_h = 14 * mm
        img_h = cell_h - caption_h
        path = _image_path(media)
        if path and os.path.exists(path):
            image = ImageReader(path)
            iw, ih = image.getSize()
            scale = min(cell_w / iw, img_h / ih)
            draw_w = iw * scale
            draw_h = ih * scale
            pdf.drawImage(
                image,
                x + (cell_w - draw_w) / 2,
                y_top - draw_h,
                width=draw_w,
                height=draw_h,
                preserveAspectRatio=True,
            )
        pdf.setFont("Helvetica-Bold", 8)
        pdf.drawString(x, y_top - img_h - 4 * mm, media.display_filename[:55])
        pdf.setFont("Helvetica", 8)
        pdf.drawString(
            x, y_top - img_h - 8 * mm, (media.description or "No description")[:70]
        )

    _pdf_footer(pdf, title, page_num, total_pages, width)
    pdf.showPage()


def _draw_pdf_file_pages(pdf, media, title, page_start, total_pages, pagesize):
    width, height = pagesize
    current_page = page_start
    path = _image_path(media)
    if not path or not os.path.exists(path):
        return current_page
    doc = fitz.open(path)
    try:
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            image = ImageReader(BytesIO(pix.tobytes("png")))
            margin = 12 * mm
            top = height - 16 * mm
            bottom = 18 * mm
            max_w = width - 2 * margin
            max_h = top - bottom
            scale = min(max_w / pix.width, max_h / pix.height)
            draw_w = pix.width * scale
            draw_h = pix.height * scale
            pdf.setFont("Helvetica-Bold", 9)
            pdf.drawString(margin, height - 10 * mm, media.display_filename[:85])
            pdf.drawImage(
                image,
                margin + (max_w - draw_w) / 2,
                bottom + (max_h - draw_h) / 2,
                width=draw_w,
                height=draw_h,
            )
            _pdf_footer(pdf, title, current_page, total_pages, width)
            pdf.showPage()
            current_page += 1
    finally:
        doc.close()
    return current_page


def _pdf_page_count_for_media(
    media_files, photos_per_page, export_pdf_files, link_count=0
):
    image_count = sum(1 for media in media_files if media.file_type == "image")
    pages = (image_count + photos_per_page - 1) // photos_per_page
    if export_pdf_files:
        for media in media_files:
            if media.file_type == "file" and media.display_filename.lower().endswith(
                ".pdf"
            ):
                path = _image_path(media)
                if path and os.path.exists(path):
                    doc = fitz.open(path)
                    try:
                        pages += doc.page_count
                    finally:
                        doc.close()
    if link_count:
        pages += (link_count + 11) // 12
    return max(pages, 1)


def _export_media_pdf(
    owner_label,
    parent_label,
    media_files,
    request=None,
    owner_kind="",
    owner_pk=None,
    photos_per_page=1,
    export_pdf_files=False,
):
    buffer = BytesIO()
    photos_per_page = int(photos_per_page or 1)
    photos_per_page = photos_per_page if photos_per_page in {1, 2, 4} else 1
    pagesize = landscape(A4) if photos_per_page == 4 else A4
    pdf = canvas.Canvas(buffer, pagesize=pagesize)
    title = f"{owner_label} Photos"
    pdf.setTitle(title)
    pdf.setAuthor("TMS")
    pdf.setSubject(title)
    link_rows = (
        _media_link_rows(request, owner_kind, owner_pk, media_files)
        if request and owner_kind and owner_pk
        else []
    )
    total_pages = _pdf_page_count_for_media(
        media_files, photos_per_page, export_pdf_files, len(link_rows)
    )
    images = [media for media in media_files if media.file_type == "image"]
    page_num = 1

    if images:
        for start in range(0, len(images), photos_per_page):
            _draw_photo_page(
                pdf,
                images[start : start + photos_per_page],
                title,
                page_num,
                total_pages,
                pagesize,
            )
            page_num += 1

    for media in media_files:
        if (
            export_pdf_files
            and media.file_type == "file"
            and media.display_filename.lower().endswith(".pdf")
        ):
            page_num = _draw_pdf_file_pages(
                pdf, media, title, page_num, total_pages, pagesize
            )

    if link_rows:
        width, height = pagesize
        link_page_num = page_num
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawCentredString(
            width / 2, height - 16 * mm, "48-hour file and video links"
        )
        y = height - 28 * mm
        pdf.setFont("Helvetica", 9)
        for serial, filename, url in link_rows:
            pdf.drawString(14 * mm, y, f"{serial:02d}. {filename[:70]}")
            y -= 5 * mm
            pdf.setFillColorRGB(0, 0, 0.75)
            pdf.drawString(18 * mm, y, url[:130])
            pdf.setFillColorRGB(0, 0, 0)
            y -= 8 * mm
            if y < 24 * mm:
                _pdf_footer(pdf, title, page_num, total_pages, width)
                pdf.showPage()
                page_num += 1
                y = height - 20 * mm
                pdf.setFont("Helvetica-Bold", 12)
                pdf.drawCentredString(
                    width / 2, height - 16 * mm, "48-hour file and video links"
                )
                pdf.setFont("Helvetica", 9)
        if page_num == link_page_num or y < height - 28 * mm:
            _pdf_footer(pdf, title, page_num, total_pages, width)

    if not images and not export_pdf_files and not link_rows:
        width, height = pagesize
        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawString(18 * mm, height - 20 * mm, title)
        pdf.setFont("Helvetica", 10)
        pdf.drawString(18 * mm, height - 30 * mm, "No photos uploaded.")
        _pdf_footer(pdf, title, 1, 1, width)

    pdf.save()
    buffer.seek(0)
    return buffer


def _add_docx_page_field(paragraph, field_name):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    instr_run._r.append(instr_text)

    end_run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_char)


def _set_docx_footer(document, title):
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(3.1), WD_TAB_ALIGNMENT.CENTER
    )
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.2), WD_TAB_ALIGNMENT.RIGHT
    )
    paragraph.add_run("\t")
    paragraph.add_run(title)
    paragraph.add_run("\t")
    paragraph.add_run(f"{timezone.localtime():%Y-%m-%d %H:%M}  Page ")
    _add_docx_page_field(paragraph, "PAGE")
    paragraph.add_run(" of ")
    _add_docx_page_field(paragraph, "NUMPAGES")


def _export_media_docx(
    owner_label, parent_label, media_files, request=None, owner_kind="", owner_pk=None
):
    document = Document()
    title = f"{owner_label} Photos"
    _set_docx_footer(document, title)
    document.add_heading(title, level=1)
    if parent_label:
        document.add_paragraph(parent_label)

    for index, media in enumerate(media_files, start=1):
        document.add_heading(f"{index:02d}. {media.display_filename}", level=2)
        document.add_paragraph(
            timezone.localtime(media.uploaded_at).strftime("%Y-%m-%d %H:%M")
        )
        document.add_paragraph(media.description or "No description")
        if media.file_type == "image":
            path = _image_path(media)
            if path and os.path.exists(path):
                document.add_picture(path, width=Inches(5.8))
        else:
            link = (
                _public_media_url(request, owner_kind, owner_pk, media.pk)
                if request and owner_kind and owner_pk
                else media.display_url
            )
            document.add_paragraph(f"48-hour link: {link}")

    if not media_files:
        document.add_paragraph("No photos uploaded.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


@login_required
def unit_media_page(request, pk):
    unit = get_object_or_404(Unit.objects.select_related("property"), pk=pk)
    if request.method == "POST":
        _upload_media_files(request, unit, UnitMedia, "unit")
        return redirect("properties:unit_media", pk=unit.pk)
    media_files = unit.media_files.filter(is_active=True).order_by(
        "sort_order", "uploaded_at", "pk"
    )
    share_token = _sign_unit_media_token(unit.pk)
    share_url = build_public_url(
        "properties:unit_media_public_share", args=[share_token]
    )
    return render(
        request,
        "properties/media_page.html",
        {
            "owner": unit,
            "owner_label": f"Unit {unit.unit_number}",
            "media_page_title": f"{unit} Photos",
            "parent_label": unit.property.property_name,
            "media_files": media_files,
            "public_token": _sign_media_token("unit", unit.pk),
            "upload_url": reverse("properties:unit_media", args=[unit.pk]),
            "back_url": reverse("properties:unit_detail", args=[unit.pk]),
            "delete_url_name": "properties:unit_media_delete",
            "update_url_name": "properties:unit_media_update",
            "sort_url": reverse("properties:unit_media_sort", args=[unit.pk]),
            "pdf_export_url": reverse(
                "properties:unit_media_export_pdf", args=[unit.pk]
            ),
            "docx_export_url": reverse(
                "properties:unit_media_export_docx", args=[unit.pk]
            ),
            "share_link_url": reverse(
                "properties:unit_media_share_link", args=[unit.pk]
            ),
            "share_url": share_url,
        },
    )


@login_required
def property_media_page(request, pk):
    property_obj = get_object_or_404(Property, pk=pk)
    if request.method == "POST":
        _upload_media_files(request, property_obj, PropertyMedia, "property")
        return redirect("properties:property_media", pk=property_obj.pk)
    media_files = property_obj.media_files.filter(is_active=True).order_by(
        "sort_order", "uploaded_at", "pk"
    )
    share_token = _sign_media_token("property", property_obj.pk)
    share_url = build_public_url("properties:media_public_share", args=[share_token])
    return render(
        request,
        "properties/media_page.html",
        {
            "owner": property_obj,
            "owner_label": property_obj.property_name,
            "media_page_title": f"{property_obj.property_name} Photos",
            "parent_label": "Property / Building",
            "media_files": media_files,
            "public_token": _sign_media_token("property", property_obj.pk),
            "upload_url": reverse("properties:property_media", args=[property_obj.pk]),
            "back_url": reverse("properties:property_detail", args=[property_obj.pk]),
            "delete_url_name": "properties:property_media_delete",
            "update_url_name": "properties:property_media_update",
            "sort_url": reverse(
                "properties:property_media_sort", args=[property_obj.pk]
            ),
            "pdf_export_url": reverse(
                "properties:property_media_export_pdf", args=[property_obj.pk]
            ),
            "docx_export_url": reverse(
                "properties:property_media_export_docx", args=[property_obj.pk]
            ),
            "share_link_url": reverse(
                "properties:property_media_share_link", args=[property_obj.pk]
            ),
            "share_url": share_url,
        },
    )


@login_required
@require_POST
def unit_media_update(request, pk, media_id):
    unit = get_object_or_404(Unit, pk=pk)
    if not _can_manage_media(request.user, unit):
        return _media_permission_denied()
    media = get_object_or_404(UnitMedia, pk=media_id, unit=unit, is_active=True)
    media.description = (request.POST.get("description") or "").strip()[:300]
    media.save(update_fields=["description", "updated_at"])
    media.refresh_image_derivatives()
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return JsonResponse(
            {"success": True, "description": media.description or "No description"}
        )
    messages.success(request, "Photo description updated.")
    return redirect("properties:unit_media", pk=unit.pk)


@login_required
@require_POST
def property_media_update(request, pk, media_id):
    property_obj = get_object_or_404(Property, pk=pk)
    if not _can_manage_media(request.user, property_obj):
        return _media_permission_denied()
    media = get_object_or_404(
        PropertyMedia, pk=media_id, property=property_obj, is_active=True
    )
    media.description = (request.POST.get("description") or "").strip()[:300]
    media.save(update_fields=["description", "updated_at"])
    media.refresh_image_derivatives()
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return JsonResponse(
            {"success": True, "description": media.description or "No description"}
        )
    messages.success(request, "Photo description updated.")
    return redirect("properties:property_media", pk=property_obj.pk)


@login_required
@require_POST
def unit_media_sort(request, pk):
    unit = get_object_or_404(Unit, pk=pk)
    if not _can_manage_media(request.user, unit):
        return _media_permission_denied()
    data = json.loads(request.body.decode("utf-8") or "{}")
    order = data.get("order") or []
    sort_value = data.get("sort_order")
    media_id = data.get("media_id")
    if media_id and sort_value:
        media = get_object_or_404(UnitMedia, pk=media_id, unit=unit, is_active=True)
        media.sort_order = int(sort_value)
        media.save(update_fields=["sort_order", "updated_at"])
    else:
        for index, media_pk in enumerate(order, start=1):
            UnitMedia.objects.filter(pk=media_pk, unit=unit, is_active=True).update(
                sort_order=index
            )
    return JsonResponse({"success": True})


@login_required
@require_POST
def property_media_sort(request, pk):
    property_obj = get_object_or_404(Property, pk=pk)
    if not _can_manage_media(request.user, property_obj):
        return _media_permission_denied()
    data = json.loads(request.body.decode("utf-8") or "{}")
    order = data.get("order") or []
    sort_value = data.get("sort_order")
    media_id = data.get("media_id")
    if media_id and sort_value:
        media = get_object_or_404(
            PropertyMedia, pk=media_id, property=property_obj, is_active=True
        )
        media.sort_order = int(sort_value)
        media.save(update_fields=["sort_order", "updated_at"])
    else:
        for index, media_pk in enumerate(order, start=1):
            PropertyMedia.objects.filter(
                pk=media_pk, property=property_obj, is_active=True
            ).update(sort_order=index)
    return JsonResponse({"success": True})


@login_required
def unit_media_export_pdf(request, pk):
    unit = get_object_or_404(Unit.objects.select_related("property"), pk=pk)
    media_files = list(
        unit.media_files.filter(is_active=True).order_by(
            "sort_order", "uploaded_at", "pk"
        )
    )
    photos_per_page = int(request.GET.get("photos_per_page") or 1)
    export_pdf_files = request.GET.get("export_pdf_files") == "1"
    buffer = _export_media_pdf(
        f"Unit {unit.unit_number}",
        unit.property.property_name,
        media_files,
        request=request,
        owner_kind="unit",
        owner_pk=unit.pk,
        photos_per_page=photos_per_page,
        export_pdf_files=export_pdf_files,
    )
    return FileResponse(
        buffer, as_attachment=True, filename=_media_export_filename(str(unit), "pdf")
    )


@login_required
def unit_media_export_docx(request, pk):
    unit = get_object_or_404(Unit.objects.select_related("property"), pk=pk)
    media_files = list(
        unit.media_files.filter(is_active=True).order_by(
            "sort_order", "uploaded_at", "pk"
        )
    )
    buffer = _export_media_docx(
        f"Unit {unit.unit_number}",
        unit.property.property_name,
        media_files,
        request=request,
        owner_kind="unit",
        owner_pk=unit.pk,
    )
    return FileResponse(
        buffer,
        as_attachment=True,
        filename=_media_export_filename(str(unit), "docx"),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@login_required
def property_media_export_pdf(request, pk):
    property_obj = get_object_or_404(Property, pk=pk)
    media_files = list(
        property_obj.media_files.filter(is_active=True).order_by(
            "sort_order", "uploaded_at", "pk"
        )
    )
    photos_per_page = int(request.GET.get("photos_per_page") or 1)
    export_pdf_files = request.GET.get("export_pdf_files") == "1"
    buffer = _export_media_pdf(
        property_obj.property_name,
        "",
        media_files,
        request=request,
        owner_kind="property",
        owner_pk=property_obj.pk,
        photos_per_page=photos_per_page,
        export_pdf_files=export_pdf_files,
    )
    return FileResponse(
        buffer,
        as_attachment=True,
        filename=_media_export_filename(property_obj.property_name, "pdf"),
    )


@login_required
def property_media_export_docx(request, pk):
    property_obj = get_object_or_404(Property, pk=pk)
    media_files = list(
        property_obj.media_files.filter(is_active=True).order_by(
            "sort_order", "uploaded_at", "pk"
        )
    )
    buffer = _export_media_docx(
        property_obj.property_name,
        "",
        media_files,
        request=request,
        owner_kind="property",
        owner_pk=property_obj.pk,
    )
    return FileResponse(
        buffer,
        as_attachment=True,
        filename=_media_export_filename(property_obj.property_name, "docx"),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@login_required
@require_POST
def unit_media_delete(request, pk, media_id):
    unit = get_object_or_404(Unit, pk=pk)
    if not _can_manage_media(request.user, unit):
        return _media_permission_denied()
    media = get_object_or_404(UnitMedia, pk=media_id, unit=unit, is_active=True)
    if request.POST.get("confirm_delete") != "yes":
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse(
                {"success": False, "error": "Media delete was not confirmed."},
                status=400,
            )
        messages.error(request, "Media delete was not confirmed.")
    else:
        media.is_active = False
        media.save(update_fields=["is_active", "updated_at"])
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"success": True})
        messages.success(request, "Media removed from active list.")
    return redirect("properties:unit_media", pk=unit.pk)


@login_required
@require_POST
def property_media_delete(request, pk, media_id):
    property_obj = get_object_or_404(Property, pk=pk)
    if not _can_manage_media(request.user, property_obj):
        return _media_permission_denied()
    media = get_object_or_404(
        PropertyMedia, pk=media_id, property=property_obj, is_active=True
    )
    if request.POST.get("confirm_delete") != "yes":
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse(
                {"success": False, "error": "Media delete was not confirmed."},
                status=400,
            )
        messages.error(request, "Media delete was not confirmed.")
    else:
        media.is_active = False
        media.save(update_fields=["is_active", "updated_at"])
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"success": True})
        messages.success(request, "Media removed from active list.")
    return redirect("properties:property_media", pk=property_obj.pk)


def _sign_media_token(owner_kind, owner_id):
    return signing.TimestampSigner(salt=UNIT_MEDIA_SHARE_SALT).sign(
        f"{owner_kind}:{owner_id}"
    )


def _sign_unit_media_token(unit_id):
    return _sign_media_token("unit", unit_id)


def _owner_from_share_token(token):
    try:
        signed_value = signing.TimestampSigner(salt=UNIT_MEDIA_SHARE_SALT).unsign(
            token,
            max_age=UNIT_MEDIA_SHARE_MAX_AGE,
        )
    except signing.SignatureExpired:
        raise Http404("This photo link has expired.")
    except signing.BadSignature:
        raise Http404("Invalid photo link.")
    try:
        owner_kind, owner_id = signed_value.split(":", 1)
    except ValueError:
        owner_kind, owner_id = "unit", signed_value
    if owner_kind == "unit":
        return owner_kind, get_object_or_404(
            Unit.objects.select_related("property"), pk=owner_id
        )
    if owner_kind == "property":
        return owner_kind, get_object_or_404(Property, pk=owner_id)
    raise Http404("Invalid photo link.")


def _unit_from_share_token(token):
    owner_kind, owner = _owner_from_share_token(token)
    if owner_kind != "unit":
        raise Http404("Invalid unit photo link.")
    return owner


@login_required
def unit_media_share_link(request, pk):
    unit = get_object_or_404(Unit.objects.select_related("property"), pk=pk)
    token = _sign_unit_media_token(unit.pk)
    share_url = build_public_url(
        "properties:unit_media_public_share", args=[token]
    )
    return render(
        request,
        "properties/unit_media_share_link.html",
        {
            "owner_label": f"{unit.property.property_name} - Unit {unit.unit_number}",
            "token": token,
            "share_url": share_url,
            "expires_hours": 48,
            "back_url": reverse("properties:unit_media", args=[unit.pk]),
        },
    )


@login_required
def property_media_share_link(request, pk):
    property_obj = get_object_or_404(Property, pk=pk)
    token = _sign_media_token("property", property_obj.pk)
    share_url = build_public_url("properties:media_public_share", args=[token])
    return render(
        request,
        "properties/unit_media_share_link.html",
        {
            "owner_label": property_obj.property_name,
            "token": token,
            "share_url": share_url,
            "expires_hours": 48,
            "back_url": reverse("properties:property_media", args=[property_obj.pk]),
        },
    )


def unit_media_public_share(request, token):
    owner_kind, owner = _owner_from_share_token(token)
    media_files = owner.media_files.filter(is_active=True).order_by(
        "sort_order", "uploaded_at", "pk"
    )
    if owner_kind == "unit":
        owner_label = f"{owner.property.property_name} - Unit {owner.unit_number}"
    else:
        owner_label = owner.property_name
    return render(
        request,
        "properties/unit_media_public_share.html",
        {
            "owner_label": owner_label,
            "token": token,
            "media_files": media_files,
        },
    )


def unit_media_public_file(request, token, media_id):
    owner_kind, owner = _owner_from_share_token(token)
    model = UnitMedia if owner_kind == "unit" else PropertyMedia
    lookup = {"unit": owner} if owner_kind == "unit" else {"property": owner}
    media = get_object_or_404(model, pk=media_id, is_active=True, **lookup)
    media_file = (
        media.stamped_file
        if media.file_type == "image" and media.stamped_file
        else media.file
    )
    if not media_file:
        raise Http404("File not found.")
    media_file.open("rb")
    return FileResponse(media_file)


def public_unit_photo_upload(request, token):
    """No-login, lease-bound gallery upload staged for admin approval."""
    import hashlib
    import uuid

    from PIL import Image, UnidentifiedImageError

    from core.upload_utils import compress_uploaded_image
    from leases.models import Lease
    from properties.public_upload_links import read_unit_photo_upload_token
    from whatsapp.models import PendingWhatsAppMedia

    try:
        token_data = read_unit_photo_upload_token(token)
    except signing.SignatureExpired:
        return render(
            request,
            "properties/public_unit_photo_upload.html",
            {"link_error": "This upload link has expired. Ask TMS for a new link."},
            status=410,
        )
    except signing.BadSignature:
        raise Http404("Invalid unit photo upload link.")

    lease = get_object_or_404(
        Lease.objects.select_related("tenant", "unit__property"),
        pk=token_data.get("lease_id"),
        unit_id=token_data.get("unit_id"),
    )
    today = timezone.localdate()
    if (
        lease.status != "active"
        or lease.start_date > today
        or lease.end_date < today
    ):
        return render(
            request,
            "properties/public_unit_photo_upload.html",
            {
                "link_error": (
                    "This link is no longer available because the selected lease "
                    "is not currently active."
                )
            },
            status=410,
        )

    context = {
        "lease": lease,
        "property_obj": lease.unit.property,
        "unit": lease.unit,
        "token": token,
        "expires_hours": 48,
    }
    if request.method != "POST":
        return render(request, "properties/public_unit_photo_upload.html", context)

    uploads = request.FILES.getlist("photos")
    if not uploads:
        context["upload_error"] = "Choose at least one photo from your gallery."
        return render(
            request, "properties/public_unit_photo_upload.html", context, status=400
        )
    if len(uploads) > 24:
        context["upload_error"] = "Upload no more than 24 photos at one time."
        return render(
            request, "properties/public_unit_photo_upload.html", context, status=400
        )

    token_fingerprint = hashlib.sha256(token.encode("utf-8")).hexdigest()[:16]
    upload_marker = f"[Public unit upload {token_fingerprint}]"
    existing_count = PendingWhatsAppMedia.objects.filter(
        lease=lease,
        ai_notes__contains=upload_marker,
    ).count()
    if existing_count + len(uploads) > 100:
        context["upload_error"] = (
            "This link has reached its 100-photo limit. Ask TMS for a new link."
        )
        return render(
            request, "properties/public_unit_photo_upload.html", context, status=400
        )

    prepared = []
    for upload in uploads:
        if upload.size > 20 * 1024 * 1024:
            context["upload_error"] = (
                f"{upload.name}: the original photo exceeds the 20 MiB safety limit."
            )
            return render(
                request,
                "properties/public_unit_photo_upload.html",
                context,
                status=400,
            )
        compressed = compress_uploaded_image(upload, max_side=1800, quality=82)
        try:
            compressed.seek(0)
            with Image.open(compressed) as image:
                if image.width * image.height > 50_000_000:
                    raise ValueError("Image dimensions are too large.")
                image.verify()
            compressed.seek(0)
        except (UnidentifiedImageError, OSError, ValueError):
            context["upload_error"] = f"{upload.name}: this is not a valid photo."
            return render(
                request,
                "properties/public_unit_photo_upload.html",
                context,
                status=400,
            )
        if compressed.size > 5 * 1024 * 1024:
            context["upload_error"] = (
                f"{upload.name}: the prepared photo still exceeds 5 MiB."
            )
            return render(
                request,
                "properties/public_unit_photo_upload.html",
                context,
                status=400,
            )
        prepared.append(compressed)

    batch_key = uuid.uuid4()
    created = []
    with transaction.atomic():
        for index, prepared_photo in enumerate(prepared, start=1):
            original_name = (
                getattr(prepared_photo, "name", "")
                or f"unit-{lease.unit_id}-photo-{index:02d}.jpg"
            )
            pending = PendingWhatsAppMedia(
                phone=getattr(lease.tenant, "phone", "") or "",
                original_filename=original_name[:255],
                media_type="image",
                purpose=PendingWhatsAppMedia.PURPOSE_UNIT,
                target_kind=PendingWhatsAppMedia.TARGET_UNIT_PHOTO,
                batch_key=batch_key,
                tenant=lease.tenant,
                lease=lease,
                property=lease.unit.property,
                unit=lease.unit,
                ai_confidence=100,
                ai_notes=(
                    f"{upload_marker} Submitted through the secure no-login Unit Photo upload "
                    f"link for lease #{lease.pk}. Destination was fixed by TMS."
                ),
                processing=False,
            )
            pending.file.save(original_name, prepared_photo, save=False)
            pending.full_clean()
            pending.save()
            created.append(pending)

    if created:
        try:
            from whatsapp.services.whatsapp_ai import notify_staff_pending_request

            notify_staff_pending_request("upload", created[0])
        except Exception:
            logger.exception(
                "Public Unit Photo batch %s was saved, but staff notification failed.",
                batch_key,
            )
    context.update(
        {
            "upload_success": (
                f"{len(created)} photo(s) uploaded successfully and sent for approval."
            ),
            "uploaded_count": len(created),
        }
    )
    return render(request, "properties/public_unit_photo_upload.html", context)


@login_required
def unit_vacancy_whatsapp(request, pk):
    unit = get_object_or_404(Unit.objects.select_related("property"), pk=pk)
    settings_obj = GlobalSettings.get_solo()
    template, rendered_message = render_unit_whatsapp_template(
        WhatsAppTemplate.TEMPLATE_VACANCY,
        unit,
        request=request,
    )
    phone = (request.GET.get("phone") or "").strip()
    whatsapp_url = build_whatsapp_url(
        phone,
        rendered_message,
        country_code=getattr(settings_obj, "country_code", "+92"),
    )
    return render(
        request,
        "properties/unit_vacancy_whatsapp.html",
        {
            "unit": unit,
            "template": template,
            "phone": phone,
            "message_text": rendered_message,
            "whatsapp_url": whatsapp_url,
        },
    )
