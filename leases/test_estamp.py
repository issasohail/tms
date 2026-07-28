import json
from datetime import datetime, timedelta
from types import SimpleNamespace
from unittest.mock import Mock, patch

from django.core.exceptions import PermissionDenied, ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import RequestFactory, SimpleTestCase
from django.utils import timezone
from pypdf import PdfReader, PdfWriter

from leases.services.estamp import (
    authorize_estamp,
    estamp_status,
    latest_estamp,
    normalize_estamp_pdf,
    compose_stamped_agreement,
)
from leases.services.agreement_package import _redact_default_clause_body


class LegalDeclarationLayoutTests(SimpleTestCase):
    def test_both_declarations_fit_on_one_legal_page_with_thumb_boxes(self):
        from io import BytesIO

        from django.template.loader import render_to_string
        from django.test import RequestFactory

        from leases.services.agreement_package import _pdf

        paragraph = (
            "I confirm that I personally know the Tenant and voluntarily provide "
            "this declaration based on my knowledge of the Tenant's character, "
            "conduct, and financial responsibility."
        )
        party = {
            "name": "Test Party",
            "cnic": "61101-1234567-1",
            "phone": "+92-300-123-4567",
            "cnic_front_url": "",
            "cnic_back_url": "",
        }
        html = render_to_string(
            "leases/proposer_seconder_declaration.html",
            {
                "history": SimpleNamespace(print_on_legal_page=True),
                "signature_config": SimpleNamespace(
                    heading="Proposer and Seconder Declaration",
                    show_thumb_impression=True,
                    footer_text="",
                ),
                "declaration_sections": [
                    {
                        "heading": "Proposer Declaration",
                        "party": party,
                        "paragraphs": [paragraph] * 4,
                    },
                    {
                        "heading": "Seconder Declaration",
                        "party": party,
                        "paragraphs": [paragraph] * 4,
                    },
                ],
            },
        )
        request = RequestFactory().get("/")
        pdf = PdfReader(BytesIO(_pdf(html, request)))

        self.assertEqual(len(pdf.pages), 1)
        self.assertEqual(
            (
                float(pdf.pages[0].mediabox.width),
                float(pdf.pages[0].mediabox.height),
            ),
            (612, 936),
        )
        text = pdf.pages[0].extract_text()
        self.assertIn("Proposer Declaration", text)
        self.assertIn("Seconder Declaration", text)
        self.assertEqual(text.count("Thumb Impression"), 2)


class DefaultLeaseTemplateRedactionTests(SimpleTestCase):
    def test_redacts_both_placeholder_formats_and_keeps_occupant_grid(self):
        body = (
            "Rent [MONTHLY_RENT], owner {{ owner_name }}, "
            "occupants {{authorized_occupants_table}}."
        )

        rendered = _redact_default_clause_body(body)

        self.assertNotIn("MONTHLY_RENT", rendered)
        self.assertNotIn("owner_name", rendered)
        self.assertNotIn("authorized_occupants_table", rendered)
        self.assertGreaterEqual(rendered.count("CONFIDENTIAL"), 6)
        self.assertIn('class="sample-occupants-table"', rendered)


class EStampPolicyTests(SimpleTestCase):
    def _lease_with(self, document):
        queryset = Mock()
        queryset.order_by.return_value.first.return_value = document
        manager = Mock()
        manager.filter.return_value = queryset
        return SimpleNamespace(documents=manager), queryset

    def test_latest_estamp_uses_required_ordering(self):
        document = object()
        lease, queryset = self._lease_with(document)
        self.assertIs(latest_estamp(lease), document)
        lease.documents.filter.assert_called_once_with(
            category="estamp_paper", is_active=True
        )
        queryset.order_by.assert_called_once_with("-uploaded_at", "-pk")

    def test_zero_max_age_disables_restriction(self):
        document = SimpleNamespace(uploaded_at=timezone.now() - timedelta(days=400))
        lease, _ = self._lease_with(document)
        status = estamp_status(
            lease,
            config=SimpleNamespace(estamp_max_age_days=0),
        )
        self.assertFalse(status.is_over_age)

    @patch("leases.services.estamp.estamp_status")
    def test_url_flag_cannot_bypass_age_without_permission(self, status_mock):
        status_mock.return_value = SimpleNamespace(
            document=object(), is_over_age=True, can_override=False
        )
        with self.assertRaises(PermissionDenied):
            authorize_estamp(object(), Mock(), allow_over_age=True)


class EStampUploadTests(SimpleTestCase):
    def _pdf(self, password=None):
        from io import BytesIO

        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)
        writer.add_blank_page(width=612, height=792)
        if password:
            writer.encrypt(password)
        output = BytesIO()
        writer.write(output)
        return SimpleUploadedFile("stamp.pdf", output.getvalue())

    def test_encrypted_pdf_is_saved_unlocked_with_all_pages(self):
        normalized = normalize_estamp_pdf(self._pdf("secret"), "secret")
        reader = PdfReader(normalized)
        self.assertFalse(reader.is_encrypted)
        self.assertEqual(len(reader.pages), 2)

    def test_encrypted_pdf_requires_password(self):
        with self.assertRaisesMessage(ValidationError, "password protected") as caught:
            normalize_estamp_pdf(self._pdf("secret"))
        self.assertEqual(caught.exception.error_list[0].code, "password_required")

    def test_wrong_password_is_reported_without_exposing_value(self):
        with self.assertRaises(ValidationError) as caught:
            normalize_estamp_pdf(self._pdf("secret"), "do-not-log-me")
        self.assertNotIn("do-not-log-me", str(caught.exception))
        self.assertEqual(caught.exception.error_list[0].code, "wrong_password")

    def test_estamp_filename_uses_property_unit_and_upload_date(self):
        from unittest.mock import patch
        from leases.views_lease_files import _estamp_filename

        lease = SimpleNamespace(
            unit=SimpleNamespace(
                unit_number="Flat 04",
                property=SimpleNamespace(property_name="F35 Building"),
            )
        )
        with patch("leases.views_lease_files.timezone.localdate") as localdate:
            localdate.return_value = datetime(2026, 7, 27).date()
            self.assertEqual(
                _estamp_filename(lease),
                "F35_Building-Flat_04_StampPaper_07272026.pdf",
            )

    def test_estamp_storage_keeps_conventional_filename(self):
        from leases.models import lease_document_upload_to

        instance = SimpleNamespace(category="estamp_paper", lease_id=17)
        self.assertEqual(
            lease_document_upload_to(
                instance, "F35_Building-Flat_04_StampPaper_07272026.pdf"
            ),
            "leases/files/17/F35_Building-Flat_04_StampPaper_07272026.pdf",
        )


class LeaseDocumentAjaxTests(SimpleTestCase):
    def setUp(self):
        self.factory = RequestFactory()

    @patch("leases.views_lease_files.get_object_or_404")
    def test_delete_returns_json_and_deactivates_document(self, get_object):
        from leases.views_lease_files import lease_file_deactivate

        document = SimpleNamespace(pk=9, lease_id=3, is_active=True, save=Mock())
        get_object.return_value = document
        request = self.factory.post(
            "/delete/", HTTP_X_REQUESTED_WITH="XMLHttpRequest"
        )
        request.user = SimpleNamespace(is_authenticated=True)

        response = lease_file_deactivate(request, 9)

        self.assertEqual(response.status_code, 200)
        self.assertTrue(json.loads(response.content)["ok"])
        self.assertFalse(document.is_active)
        document.save.assert_called_once_with(update_fields=["is_active"])

    @patch("leases.views_lease_files.get_object_or_404")
    def test_description_updates_inline_without_redirect(self, get_object):
        from leases.views_lease_files import lease_file_description_update

        document = SimpleNamespace(description="", save=Mock())
        get_object.return_value = document
        request = self.factory.post(
            "/description/",
            {"description": "Current unlocked stamp"},
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )
        request.user = SimpleNamespace(is_authenticated=True)

        response = lease_file_description_update(request, 9)
        payload = json.loads(response.content)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(payload["display"], "Current unlocked stamp")
        self.assertEqual(document.description, "Current unlocked stamp")
        document.save.assert_called_once_with(update_fields=["description"])


class EStampCompositionTests(SimpleTestCase):
    def _text_pdf(self, labels, size):
        from io import BytesIO
        from reportlab.pdfgen import canvas

        output = BytesIO()
        pdf = canvas.Canvas(output, pagesize=size)
        for label in labels:
            pdf.drawString(40, size[1] - 40, label)
            pdf.showPage()
        pdf.save()
        return output.getvalue()

    def test_second_stamp_page_maps_to_last_agreement_page(self):
        agreement = self._text_pdf(
            ["AGREEMENT ONE", "AGREEMENT TWO", "AGREEMENT THREE"], (612, 936)
        )
        stamp = self._text_pdf(["STAMP ONE", "STAMP TWO"], (595, 842))
        from io import BytesIO
        result = PdfReader(BytesIO(compose_stamped_agreement(agreement, stamp, "legal")))
        self.assertEqual(len(result.pages), 3)
        self.assertIn("STAMP ONE", result.pages[0].extract_text())
        self.assertNotIn("STAMP", result.pages[1].extract_text())
        self.assertIn("STAMP TWO", result.pages[2].extract_text())
        for page in result.pages:
            self.assertEqual((float(page.mediabox.width), float(page.mediabox.height)), (612, 936))
            for box_name in ("cropbox", "trimbox", "bleedbox", "artbox"):
                box = getattr(page, box_name)
                self.assertEqual(
                    (float(box.left), float(box.bottom), float(box.right), float(box.top)),
                    (0, 0, 612, 936),
                )

    def test_one_stamp_page_is_not_repeated(self):
        agreement = self._text_pdf(["ONE", "TWO"], (612, 792))
        stamp = self._text_pdf(["ONLY STAMP"], (595, 842))
        from io import BytesIO
        result = PdfReader(BytesIO(compose_stamped_agreement(agreement, stamp, "letter")))
        self.assertIn("ONLY STAMP", result.pages[0].extract_text())
        self.assertNotIn("ONLY STAMP", result.pages[1].extract_text())
        self.assertEqual(float(result.pages[0].mediabox.height), 792)

    def test_rotated_non_zero_cropbox_is_normalized_without_mutating_input(self):
        from io import BytesIO
        from pypdf.generic import RectangleObject

        writer = PdfWriter()
        page = writer.add_blank_page(width=300, height=500)
        page.cropbox = RectangleObject((20, 30, 280, 470))
        page.rotate(90)
        source = BytesIO()
        writer.write(source)
        stamp = source.getvalue()
        agreement = self._text_pdf(["AGREEMENT"], (612, 792))
        result = compose_stamped_agreement(agreement, stamp, "letter")
        self.assertEqual(stamp, source.getvalue())
        page = PdfReader(BytesIO(result)).pages[0]
        self.assertEqual((float(page.mediabox.width), float(page.mediabox.height)), (612, 792))

    @patch("leases.services.agreement_package._footer_overlay")
    @patch("leases.services.agreement_package.AgreementSignatureTemplate.current")
    def test_package_footer_uses_separate_legal_and_letter_positions(
        self, config_mock, overlay_mock
    ):
        from pypdf import PageObject
        from leases.services.agreement_package import merge_pdfs

        config_mock.return_value = SimpleNamespace(
            agreement_legal_footer_bottom_points=22,
            agreement_letter_footer_bottom_points=11,
            show_agreement_page_numbers=True,
        )
        overlay_mock.side_effect = lambda width, height, *args: (
            PageObject.create_blank_page(width=width, height=height)
        )

        merge_pdfs(
            [
                self._text_pdf(["LEGAL"], (612, 936)),
                self._text_pdf(["LETTER"], (612, 792)),
            ]
        )

        self.assertEqual(
            [call.args[-1] for call in overlay_mock.call_args_list],
            [22.0, 11.0],
        )
        self.assertEqual(
            [call.args[-2] for call in overlay_mock.call_args_list],
            ["Page 1 of 2", "Page 2 of 2"],
        )

    @patch("leases.services.agreement_package._footer_overlay")
    @patch("leases.services.agreement_package.AgreementSignatureTemplate.current")
    def test_package_page_numbers_can_be_hidden(self, config_mock, overlay_mock):
        from pypdf import PageObject
        from leases.services.agreement_package import merge_pdfs

        config_mock.return_value = SimpleNamespace(
            agreement_letter_footer_bottom_points=4,
            show_agreement_page_numbers=False,
        )
        overlay_mock.side_effect = lambda width, height, *args: (
            PageObject.create_blank_page(width=width, height=height)
        )

        merge_pdfs([self._text_pdf(["LETTER"], (612, 792))])

        self.assertEqual(overlay_mock.call_args.args[-2], "")

    def test_legal_identity_cards_honor_footer_clearance(self):
        from io import BytesIO

        import fitz
        from pypdf import PdfWriter

        from leases.services.agreement_package import _identity_overlay_page

        people = [
            {"role": f"Party {index}", "name": "Test Person", "person": None}
            for index in range(1, 5)
        ]
        overlay = _identity_overlay_page(
            612,
            936,
            people,
            3.10,
            footer_clearance=50,
        )
        writer = PdfWriter()
        writer.add_page(overlay)
        output = BytesIO()
        writer.write(output)

        document = fitz.open(stream=output.getvalue(), filetype="pdf")
        try:
            page = document[0]
            card_bottoms = [
                page.rect.height - drawing["rect"].y1
                for drawing in page.get_drawings()
                if drawing["rect"].width > 100
                and drawing["rect"].height > 100
            ]
        finally:
            document.close()

        self.assertEqual(len(card_bottoms), 4)
        self.assertTrue(all(abs(bottom - 50) < 0.1 for bottom in card_bottoms))

    def test_repeated_page_signature_lines_do_not_repeat_party_details(self):
        from leases.services.agreement_package import (
            _agreement_signature_footer_page,
        )

        lease = SimpleNamespace(
            tenant=SimpleNamespace(
                cnic="71702-0346063-5",
                get_full_name=lambda: "Tenant Full Name",
            ),
            unit=SimpleNamespace(
                property=SimpleNamespace(
                    owner_name="Owner Full Name",
                    owner_cnic="42101-2008010-3",
                )
            ),
        )

        page = _agreement_signature_footer_page(612, 936, lease, 50)
        text = page.extract_text()

        self.assertIn("Owner Signature:", text)
        self.assertIn("Tenant Signature:", text)
        self.assertNotIn("Owner Full Name", text)
        self.assertNotIn("Tenant Full Name", text)
        self.assertNotIn("42101-2008010-3", text)
        self.assertNotIn("71702-0346063-5", text)


class EStampPackageIntegrationTests(SimpleTestCase):
    @patch("leases.services.agreement_package.merge_pdfs", return_value=b"package")
    @patch("leases.services.agreement_package.signature_pdf", return_value=b"signature")
    @patch("leases.services.agreement_package.police_pdf", return_value=b"police")
    @patch("leases.services.agreement_package.inspection_pdf", return_value=b"inspection")
    @patch("leases.services.estamp.compose_stamped_agreement", return_value=b"stamped-core")
    @patch("leases.services.estamp.authorize_estamp")
    @patch(
        "leases.services.agreement_package.AgreementSignatureTemplate.current",
        return_value=SimpleNamespace(estamp_letter_footer_bottom_points=28),
    )
    @patch("leases.services.agreement_package.agreement_pdf", return_value=b"plain-core")
    def test_only_core_agreement_is_stamped(
        self,
        agreement_mock,
        config_mock,
        authorize_mock,
        compose_mock,
        inspection_mock,
        police_mock,
        signature_mock,
        merge_mock,
    ):
        from leases.services.agreement_package import build_package

        file_handle = Mock()
        file_handle.__enter__ = Mock(return_value=SimpleNamespace(read=Mock(return_value=b"stamp")))
        file_handle.__exit__ = Mock(return_value=False)
        authorize_mock.return_value = SimpleNamespace(
            file=SimpleNamespace(open=Mock(return_value=file_handle))
        )
        history = SimpleNamespace(
            print_with_estamp=True,
            estamp_paper_size="letter",
            allow_over_age_estamp=False,
        )
        request = SimpleNamespace(user=Mock())
        lease = Mock()

        payload, filename, _ = build_package(request, lease, history, [])

        self.assertEqual(payload, b"package")
        self.assertTrue(filename.endswith("_Letter.pdf"))
        compose_mock.assert_called_once_with(
            b"plain-core",
            b"stamp",
            "letter",
            stamp_footer_bottom_points=28,
        )
        self.assertEqual(
            merge_mock.call_args.args[0],
            [b"stamped-core", b"inspection", b"police", b"signature"],
        )


class AgreementDocxFormattingTests(SimpleTestCase):
    def test_inline_placeholder_spacing_is_preserved(self):
        from bs4 import BeautifulSoup
        from docx import Document
        from leases.views import _append_inline

        source = BeautifulSoup(
            "<span>Rent is <strong>30,000</strong> per month.</span>",
            "html.parser",
        ).span
        paragraph = Document().add_paragraph()
        for child in source.children:
            _append_inline(paragraph, child)

        self.assertEqual(paragraph.text, "Rent is 30,000 per month.")

    def test_authorized_occupants_html_becomes_word_grid(self):
        from bs4 import BeautifulSoup
        from docx import Document
        from leases.views import _add_docx_html_table

        source = BeautifulSoup(
            """
            <table><tr>
              <td><b>1. Tenant One</b><br>CNIC: 1<br>Relationship: Tenant</td>
              <td>&nbsp;<br>&nbsp;<br>&nbsp;</td>
              <td>&nbsp;<br>&nbsp;<br>&nbsp;</td>
              <td>&nbsp;<br>&nbsp;<br>&nbsp;</td>
            </tr></table>
            """,
            "html.parser",
        ).table
        table = _add_docx_html_table(Document(), source)

        self.assertEqual((len(table.rows), len(table.columns)), (1, 4))
        self.assertEqual(table.style.name, "Table Grid")
        self.assertIn("Relationship: Tenant", table.cell(0, 0).text)

    def test_empty_word_witness_fields_use_matching_tabbed_lines(self):
        from docx import Document
        from leases.views import add_signature_block

        tenant = SimpleNamespace(
            cnic="7170203460635",
            phone="+92-342-159-9177",
            get_full_name=lambda: "Tenant One",
        )
        property_obj = SimpleNamespace(
            owner_name="Owner One",
            owner_cnic="4210120080103",
            owner_phone="+92-312-255-0183",
        )
        lease = SimpleNamespace(
            tenant=tenant,
            unit=SimpleNamespace(property=property_obj),
            witness1_tenant=None,
            witness2_tenant=None,
        )
        doc = Document()
        add_signature_block(doc, lease)

        party_table = doc.tables[0]
        witness_table = doc.tables[1]
        signature_rule = "_" * 39
        for table in (party_table, witness_table):
            self.assertEqual(table.cell(0, 0).text.strip(), signature_rule)
            self.assertEqual(table.cell(0, 1).text.strip(), signature_rule)
            for column in (0, 1):
                self.assertAlmostEqual(
                    table.cell(0, column).width.inches,
                    3.85,
                    places=2,
                )

        owner_details = party_table.cell(1, 0).paragraphs
        tenant_details = party_table.cell(1, 1).paragraphs
        self.assertEqual(len(owner_details), 4)
        self.assertEqual(len(tenant_details), 4)
        self.assertTrue(owner_details[2].text.startswith("Phone:\t+92"))
        self.assertTrue(tenant_details[2].text.startswith("Phone:\t+92"))
        self.assertTrue(owner_details[3].text.startswith("Date:\t"))
        self.assertTrue(tenant_details[3].text.startswith("Date:\t"))
        self.assertTrue(
            all(
                paragraph.paragraph_format.space_after.pt == 0
                for paragraph in owner_details + tenant_details
            )
        )

        blank = "_" * 29
        for column in (0, 1):
            paragraphs = witness_table.cell(1, column).paragraphs
            self.assertEqual(len(paragraphs), 3)
            self.assertTrue(all(p.text.endswith(blank) for p in paragraphs))
            self.assertTrue(all("\t" in p.text for p in paragraphs))

    def test_word_page_uses_compact_side_margins(self):
        from docx import Document
        from leases.views import set_doc_margins

        doc = Document()
        set_doc_margins(doc)

        self.assertAlmostEqual(doc.sections[0].left_margin.inches, 0.35, places=2)
        self.assertAlmostEqual(doc.sections[0].right_margin.inches, 0.35, places=2)
