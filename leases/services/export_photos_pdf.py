# leases/services/export_photos_pdf.py
import io

from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.utils import timezone
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageOps
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


def _tenant_name(lease):
    tenant = getattr(lease, "tenant", None)
    if tenant and hasattr(tenant, "get_full_name"):
        return tenant.get_full_name() or str(tenant)
    return str(tenant or "Tenant")


def _property_name(lease):
    unit = getattr(lease, "unit", None)
    prop = getattr(unit, "property", None)
    return getattr(prop, "property_name", "") or str(prop or "Property")


def _unit_name(lease):
    unit = getattr(lease, "unit", None)
    return getattr(unit, "unit_number", "") or str(unit or "")


def _lease_meta(lease):
    return {
        "tenant": _tenant_name(lease),
        "property": _property_name(lease),
        "unit": _unit_name(lease),
        "start": getattr(lease, "start_date", "") or "",
        "end": getattr(lease, "end_date", "") or "",
    }


def _sorted_existing_photos(pcr):
    photos = []
    for photo in pcr.photos.order_by("sort_order", "created_at", "id"):
        if photo.image and photo.image.name and default_storage.exists(photo.image.name):
            photos.append(photo)
    return photos


def _photo_bytes(photo):
    with default_storage.open(photo.image.name, "rb") as fh:
        raw = fh.read()
    img = Image.open(io.BytesIO(raw))
    img = ImageOps.exif_transpose(img).convert("RGB")
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=90)
    out.seek(0)
    return out


def _photo_contained_bytes(photo, size=(1600, 1200)):
    with default_storage.open(photo.image.name, "rb") as fh:
        raw = fh.read()
    img = Image.open(io.BytesIO(raw))
    img = ImageOps.exif_transpose(img).convert("RGB")
    img.thumbnail(size, Image.LANCZOS)
    canvas_img = Image.new("RGB", size, "white")
    x = (size[0] - img.width) // 2
    y = (size[1] - img.height) // 2
    canvas_img.paste(img, (x, y))
    out = io.BytesIO()
    canvas_img.save(out, format="JPEG", quality=90)
    out.seek(0)
    return out


def _filename(prefix, lease, ext):
    def part(value, limit=None):
        cleaned = "".join(ch if ch.isalnum() else "_" for ch in str(value or "")).strip("_")
        while "__" in cleaned:
            cleaned = cleaned.replace("__", "_")
        cleaned = cleaned or "NA"
        return cleaned[:limit] if limit else cleaned

    meta = _lease_meta(lease)
    return (
        f"{part(meta['property'])}_{part(meta['unit'])}_"
        f"{part(meta['tenant'], 20)}_{part(meta['end'])}_Photos.{ext}"
    )


def export_photos_pdf(pcr):
    lease = pcr.lease
    meta = _lease_meta(lease)
    photos = _sorted_existing_photos(pcr)
    generated = timezone.localtime(timezone.now()).strftime("%Y-%m-%d %H:%M")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4
    c.setTitle(f"Lease #{lease.id} - PCR Photos")

    footer_center = (
        f"{meta['tenant']} | {meta['property']} | Unit {meta['unit']} | "
        f"{meta['start']} to {meta['end']}"
    )

    def footer(page_no):
        c.setFont("Helvetica", 7)
        c.drawString(12 * mm, 9 * mm, generated)
        c.drawCentredString(W / 2, 9 * mm, footer_center[:120])
        c.drawRightString(W - 12 * mm, 9 * mm, f"Page {page_no}")

    def header():
        c.setFont("Helvetica-Bold", 14)
        c.drawString(16 * mm, H - 18 * mm, "Property Condition Report Photos")
        c.setFont("Helvetica", 9)
        c.drawString(16 * mm, H - 26 * mm, f"Tenant: {meta['tenant']}")
        c.drawString(16 * mm, H - 32 * mm, f"Property: {meta['property']}    Unit: {meta['unit']}")
        c.drawString(16 * mm, H - 38 * mm, f"Lease: {meta['start']} to {meta['end']}")

    if not photos:
        header()
        c.setFont("Helvetica", 11)
        c.drawString(16 * mm, H - 55 * mm, "No available PCR photos were found.")
        footer(1)
        c.showPage()
    else:
        for page_no, photo in enumerate(photos, start=1):
            header()
            image_box_x = 16 * mm
            image_box_y = 54 * mm
            image_box_w = 128 * mm
            image_box_h = 175 * mm
            sign_x = image_box_x + image_box_w + 8 * mm
            sign_w = W - sign_x - 16 * mm

            try:
                image_data = _photo_bytes(photo)
                c.rect(image_box_x, image_box_y, image_box_w, image_box_h, stroke=1, fill=0)
                c.drawImage(
                    ImageReader(image_data),
                    image_box_x,
                    image_box_y,
                    image_box_w,
                    image_box_h,
                    preserveAspectRatio=True,
                    anchor="c",
                )
            except Exception:
                c.rect(image_box_x, image_box_y, image_box_w, image_box_h, stroke=1, fill=0)
                c.setFont("Helvetica", 10)
                c.drawCentredString(image_box_x + image_box_w / 2, image_box_y + image_box_h / 2, "Photo file missing")

            caption = f"{photo.room or 'Area'}"
            if photo.comment:
                caption += f" - {photo.comment}"
            if photo.taken_at:
                caption += f" - Taken: {timezone.localtime(photo.taken_at).strftime('%Y-%m-%d %H:%M')}"
            c.setFont("Helvetica", 8)
            c.drawString(image_box_x, image_box_y - 6 * mm, caption[:130])

            c.setFont("Helvetica-Bold", 9)
            c.drawString(sign_x, image_box_y + image_box_h - 8 * mm, "Tenant Initials")
            c.line(sign_x, image_box_y + image_box_h - 20 * mm, sign_x + sign_w, image_box_y + image_box_h - 20 * mm)
            c.setFont("Helvetica", 8)
            c.drawString(sign_x, image_box_y + image_box_h - 28 * mm, "Initial/sign beside photo")

            footer(page_no)
            c.showPage()

    page_no = len(photos) + 1
    header()
    y = H - 58 * mm
    c.setFont("Helvetica-Bold", 12)
    c.drawString(16 * mm, y, "Final Tenant Signature")
    y -= 14 * mm
    c.setFont("Helvetica", 10)
    c.drawString(16 * mm, y, f"Tenant Name: {meta['tenant']}")
    y -= 18 * mm
    c.drawString(16 * mm, y, "Tenant Signature:")
    c.line(55 * mm, y - 1 * mm, 145 * mm, y - 1 * mm)
    y -= 14 * mm
    c.drawString(16 * mm, y, f"Date Generated: {generated}")
    footer(page_no)
    c.showPage()

    c.save()
    buf.seek(0)
    return ContentFile(buf.read(), name=_filename("pcr-photos", lease, "pdf"))


def _add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def export_photos_docx(pcr):
    lease = pcr.lease
    meta = _lease_meta(lease)
    photos = _sorted_existing_photos(pcr)
    generated = timezone.localtime(timezone.now()).strftime("%Y-%m-%d %H:%M")

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    footer_table = section.footer.add_table(rows=1, cols=3, width=Inches(7.3))
    footer_table.cell(0, 0).text = generated
    footer_table.cell(0, 1).text = (
        f"{meta['tenant']} | {meta['property']} | Unit {meta['unit']} | "
        f"{meta['start']} to {meta['end']}"
    )
    page_para = footer_table.cell(0, 2).paragraphs[0]
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_para.add_run("Page ")
    _add_page_number(page_para)

    doc.add_heading("Property Condition Report Photos", level=1)
    for label, value in (
        ("Tenant", meta["tenant"]),
        ("Property", meta["property"]),
        ("Unit", meta["unit"]),
        ("Lease Start", meta["start"]),
        ("Lease End", meta["end"]),
        ("Generated", generated),
    ):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    if not photos:
        doc.add_paragraph("No available PCR photos were found.")

    for idx, photo in enumerate(photos, start=1):
        doc.add_section(WD_SECTION.NEW_PAGE)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(5.0)
        table.columns[1].width = Inches(2.0)
        left = table.cell(0, 0)
        right = table.cell(0, 1)
        para = left.paragraphs[0]
        try:
            image_data = _photo_contained_bytes(photo)
            para.add_run().add_picture(image_data, width=Inches(5.0), height=Inches(3.75))
        except Exception:
            para.add_run("Photo file missing")
        caption = f"{idx}. {photo.room or 'Area'}"
        if photo.comment:
            caption += f" - {photo.comment}"
        if photo.taken_at:
            caption += f" - Taken: {timezone.localtime(photo.taken_at).strftime('%Y-%m-%d %H:%M')}"
        left.add_paragraph(caption)
        right.paragraphs[0].add_run("Tenant Initials").bold = True
        line = right.add_paragraph()
        line.add_run("\n\n____________________")
        hint = right.add_paragraph("Initial/sign beside photo")
        hint.runs[0].font.size = Pt(8)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Final Tenant Signature", level=2)
    doc.add_paragraph(f"Tenant Name: {meta['tenant']}")
    doc.add_paragraph("\nTenant Signature: ______________________________")
    doc.add_paragraph(f"Date Generated: {generated}")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return ContentFile(out.read(), name=_filename("pcr-photos", lease, "docx"))
