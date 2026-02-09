
# adjust if your form/formset live elsewhere
# adjust if your app paths differ
from .models import Lease, Property, Unit, Tenant  # adjust import paths if needed
from .models import Lease, Property, Unit, Tenant
from datetime import datetime, date
from django.shortcuts import get_object_or_404, render
from django.shortcuts import get_object_or_404, redirect
from django.views.decorators.http import require_POST
from .models import Tenant  # adjust import if needed
from django.utils.html import escape
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_GET
from .forms import LeaseForm, LeaseFamilyFormSet
from django.db.models.functions import Lower
from django.apps import apps
from .utils.billing import (
    preview_initial_billing,
    apply_initial_billing,
    preview_billing_on_change,
    update_billing_on_change,
)
from django.views.generic import CreateView, UpdateView
from django.shortcuts import render
from .utils.billing import update_billing_on_change
from django.utils.timezone import now
from django.db import transaction
from django.shortcuts import redirect
from .models import Lease, LeaseFamily
from django.views.generic import CreateView, UpdateView, DetailView
from django.http import JsonResponse, HttpResponseRedirect
from django.views.decorators.http import require_http_methods
from django.http import FileResponse
import logging
from leases.templatetags.lease_tags import replace_placeholders
from leases.models import Lease, LeaseAgreementClause
from django.shortcuts import render, redirect
from math import ceil
import subprocess
from django.core.files.base import ContentFile
import io
from docxtpl import DocxTemplate
from django.template import Template, Context
import re
from .models import LeaseTemplate
from .forms import LeaseTemplateForm
from invoices.models import Invoice
from django.shortcuts import get_object_or_404
from .models import Lease
from properties.models import Property, Unit
from django_tables2 import SingleTableView
from django_tables2.views import SingleTableView
from datetime import timedelta
from decimal import Decimal
from io import BytesIO
from django.utils import timezone
from bs4 import NavigableString, Tag
from django.core.mail import EmailMessage
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse_lazy
from django.utils import timezone
from django.views import View
from django.views.generic import (
    ListView, DetailView, CreateView, UpdateView, DeleteView
)
from utils.pdf_export import handle_export
from dal import autocomplete
from reportlab.pdfgen import canvas
from weasyprint import HTML, CSS
from django.template.loader import render_to_string
from leases.models import Lease
from leases.forms import LeaseForm
from tenants.models import Tenant
from properties.models import Unit, Property
from payments.pdf_utils import generate_payment_pdf
from django.views.generic import DetailView
from rest_framework.decorators import api_view
from rest_framework.response import Response
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib import admin
from .forms import RenewLeaseForm
from django.utils.html import format_html, mark_safe
from django.db import models
from django.urls import reverse, NoReverseMatch
from django.conf import settings
from django.db.models import Sum
from django.db.models import Q
from django import forms
from .utils import generate_lease_agreement
from django.views.generic import TemplateView
from django.views.generic import FormView
from .forms import RenewLeaseForm, CustomRenewForm
from django.contrib.humanize.templatetags.humanize import intcomma
from django.http import HttpResponse
from openpyxl import Workbook
import tempfile
import os
from django.utils.text import slugify
from datetime import datetime
import csv
from django_tables2.export.export import TableExport
from django.views.generic import ListView
# Ensure this is correct for your setup
from django_tables2.export.export import TableExport
from django.core.exceptions import PermissionDenied
from django_tables2.views import SingleTableView
from payments.models import Payment
from documents.models import LeaseDocument
from .tables import LeaseTable
from django_tables2.export.views import ExportMixin
from django.http import JsonResponse
from django.utils.html import format_html
from .tables import LeaseTable
from django.http import JsonResponse
from properties.models import Unit
from django.template.loader import render_to_string
from weasyprint import HTML
# leases/views.py
from django.views.generic import View
from django.http import HttpResponseRedirect, JsonResponse
from django.urls import reverse
from django.contrib import messages
from .utils.agreement_generator import generate_lease_agreement
from .utils.email_service import send_lease_agreement_email
from django.db.models import F, ExpressionWrapper, DecimalField, Sum
from .models import Lease
from payments.models import Payment  # Add this
from invoices.models import Invoice , SecurityDepositTransaction # Add this
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from datetime import datetime
import openpyxl
from datetime import timedelta, date
from django.core.files import File
from django.core.files.storage import default_storage
from leases.utils import do_replace_placeholders
from invoices.services import security_deposit_balance, security_deposit_totals
from invoices.services import security_deposit_totals
from invoices.services import security_deposit_totals
from typing import Dict, Any

from tenants.models import Tenant, normalize_cnic
# leases/views.py
from .forms import LeaseForm, LeaseFamilyFormSet
# LeaseFamily is needed if you link/create in quick-add
from .models import Lease, LeaseFamily
from django.db.models import Value
from django.db.models.functions import Concat, Lower
# leases/views.py
from django.apps import apps

from django.db import transaction
from django.shortcuts import render
from django.views.generic import CreateView, UpdateView
from django.db.models.functions import Lower
from decimal import Decimal
from decimal import Decimal
from invoices.services import security_deposit_totals
from invoices.models import SecurityDepositTransaction
from decimal import Decimal
from invoices.models import SecurityDepositTransaction

from .models import Lease
from .utils.billing import (
    preview_initial_billing,
    apply_initial_billing,
    preview_billing_on_change,
    update_billing_on_change,
)

from .utils.sql_capture import run_and_capture
import logging
import json
import time
from django.core.serializers.json import DjangoJSONEncoder
from django.db import connection, transaction
from django.forms.models import model_to_dict

# --- Security deposit helpers ------------------------------------
from decimal import Decimal

from decimal import Decimal

ZERO = Decimal("0.00")

import re

def strip_html(text):
    if not text:
        return ""
    return re.sub(r'<[^>]+>', '', text)

def lease_applied_amount(payment) -> Decimal:
    """
    Amount that should affect Lease ledger balance.
    - If PaymentAllocation exists: use allocation.lease_amount
    - Else: legacy fallback to payment.amount
    """
    alloc = getattr(payment, "allocation", None)
    if alloc:
        return alloc.lease_amount or ZERO
    return payment.amount or ZERO

def create_initial_security_required(lease):
    """
    On lease create: write one REQUIRED row for the agreed security deposit,
    if > 0.
    """
    amount = lease.security_deposit or Decimal("0.00")
    if amount <= 0:
        return None

    return SecurityDepositTransaction.objects.create(
        lease=lease,
        type="REQUIRED",
        amount=amount,
        notes="Initial required security deposit set from lease.",
    )



def _plan_summary(plan):
    if not plan:
        return {"has_plan": False}
    return {
        "has_plan": True,
        "requires_security_confirmation": plan.get("requires_security_confirmation"),
        "security_item": {
            "before": plan.get("security_item", {}).get("before"),
            "after":  plan.get("security_item", {}).get("after"),
        },
        "invoices_to_create_count": len(plan.get("invoices_to_create") or []),
        "recurring": plan.get("recurring") or [],
    }


def resolve_placeholders(lease, text):
    """
    Replace placeholders in text with actual values from the lease
    """
    replacements = {
        '[OWNER_NAME]': lease.unit.property.owner_name if lease.unit and lease.unit.property else '',
        '[OWNER_CNIC]': lease.unit.property.owner_cnic if lease.unit and lease.unit.property else '',
        '[OWNER_ADDRESS]': lease.unit.property.owner_address if lease.unit and lease.unit.property else '',
        '[TENANT_NAME]': lease.tenant.get_full_name() if lease.tenant else '',
        '[TENANT_CNIC]': lease.tenant.cnic if lease.tenant else '',
        '[TENANT_ADDRESS]': lease.tenant.address if lease.tenant else '',
        '[PROPERTY_NAME]': lease.unit.property.property_name if lease.unit and lease.unit.property else '',
        '[UNIT_NUMBER]': lease.unit.unit_number if lease.unit else '',
        '[START_DATE]': lease.start_date.strftime('%d %B %Y') if lease.start_date else '',
        '[END_DATE]': lease.end_date.strftime('%d %B %Y') if lease.end_date else '',
        '[MONTHLY_RENT]': str(lease.monthly_rent) if lease.monthly_rent is not None else '',
        '[SECURITY_DEPOSIT]': str(lease.security_deposit) if lease.security_deposit is not None else '',
        '[SOCIETY_MAINTENANCE]': str(lease.society_maintenance) if lease.society_maintenance is not None else '',
        # Add more placeholders as needed
    }

    for placeholder, value in replacements.items():
        text = text.replace(placeholder, value)

    return text


def prepare_transaction_columns(transactions):
    columns = []
    if len(transactions) <= 25:
        # Single column with all transactions
        columns.append([{'index': i+1, **t}
                       for i, t in enumerate(transactions)])
    else:
        # Split into multiple columns with 25 items each
        items_per_column = 25
        for i in range(0, len(transactions), items_per_column):
            column = transactions[i:i+items_per_column]
            columns.append([{'index': j+1, **t}
                           for j, t in enumerate(column, start=i)])
    return columns


class LeaseListView(SingleTableView):
    model = Lease
    table_class = LeaseTable
    template_name = 'leases/lease_list.html'
    paginate_by = 40  # Set to show 40 records per page

    def get_queryset(self):
        queryset = super().get_queryset().select_related(
            'tenant', 'unit', 'unit__property'
        ).order_by('unit__unit_number')  # Order by unit number

        # Get filter parameters
        property_id = self.request.GET.get('property')
        unit_id = self.request.GET.get('unit')
        tenant_id = self.request.GET.get('tenant')
        include_inactive = self.request.GET.get('include_inactive') == 'on'
        nonzero_balance = self.request.GET.get('nonzero_balance') == 'on'

        # Apply filters
        if property_id:
            queryset = queryset.filter(unit__property_id=property_id)
        if unit_id:
            queryset = queryset.filter(unit_id=unit_id)
        if tenant_id:
            queryset = queryset.filter(tenant_id=tenant_id)
        if not include_inactive:
            queryset = queryset.filter(status='active')
        if nonzero_balance:
            queryset = [lease for lease in queryset if lease.get_balance > 0]

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Get all properties for dropdown
        context['all_properties'] = Property.objects.all()

        # Get all tenants ordered by first name
        context['all_tenants'] = Tenant.objects.all().order_by('first_name')

        # Get filtered units based on selected property
        property_id = self.request.GET.get('property')
        if property_id:
            context['filtered_units'] = Unit.objects.filter(
                property_id=property_id)
        else:
            context['filtered_units'] = Unit.objects.none()

        # Add current filter values to context
        context['current_property'] = self.request.GET.get('property', '')
        context['current_unit'] = self.request.GET.get('unit', '')
        context['current_tenant'] = self.request.GET.get('tenant', '')
        context['include_inactive'] = self.request.GET.get(
            'include_inactive', '') == 'on'
        context['nonzero_balance'] = self.request.GET.get(
            'nonzero_balance', '')
        leases = self.get_queryset()

        # 🔹 existing total balance
        context['total_balance'] = sum(
            getattr(l, 'get_balance', 0) for l in leases
        )

        # 🔹 NEW: total security deposit still due across all listed leases
        context['total_security_due'] = sum(
            security_deposit_balance(lease) for lease in leases
        )

        return context

    def get(self, request, *args, **kwargs):
        resp = handle_export(request, self.get_table(), 'leases')
        return resp or super().get(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        # Handle export requests
        export_response = self.handle_export(request)
        if export_response:
            return export_response
        return super().get(request, *args, **kwargs)

    def handle_export(self, request):
        """Handle export functionality"""
        self.object_list = self.get_queryset()
        table = self.get_table()
        export_name = "leases_list"
        return handle_export(request, table, export_name)


# leases/views.py  (replace these three views with the versions below)


def _split_name(full_name: str):
    parts = (full_name or '').strip().split()
    first = parts[0] if parts else ''
    last = ' '.join(parts[1:]) if len(parts) > 1 else ''
    return first, last


# at top


class LeaseTenantOrderMixin:
    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        Tenant = apps.get_model("tenants", "Tenant")
        fields = {f.name for f in Tenant._meta.get_fields()}
        if {"first_name", "last_name"} <= fields:
            qs = Tenant.objects.order_by(
                Lower("first_name"), Lower("last_name"), "id")
        elif "name" in fields:
            qs = Tenant.objects.order_by(Lower("name"), "id")
        elif "full_name" in fields:
            qs = Tenant.objects.order_by(Lower("full_name"), "id")
        else:
            qs = Tenant.objects.order_by("id")
        form.fields["tenant"].queryset = qs
        return form


def _order_tenants(form):
    Tenant = apps.get_model("tenants", "Tenant")
    if "tenant" in form.fields:
        form.fields["tenant"].queryset = Tenant.objects.order_by(
            Lower("first_name"), Lower("last_name"), "id"
        )


def _seed_units_queryset(form, request):
    if "unit" not in form.fields:
        return
    Unit = apps.get_model("properties", "Unit")
    prop_id = request.POST.get("property") or request.GET.get("property")
    if getattr(form.instance, "unit_id", None) and not prop_id:
        prop_id = form.instance.unit.property_id
    form.fields["unit"].queryset = (
        Unit.objects.filter(property_id=prop_id).order_by("unit_number") if prop_id
        else Unit.objects.none()
    )


logger = logging.getLogger(__name__)


def _jsonable(obj):
    try:
        return json.loads(json.dumps(obj, cls=DjangoJSONEncoder, default=str))
    except Exception as e:
        return {"_err": str(e), "repr": repr(obj)}

from tenants.models import Tenant
from django.db.models.functions import Lower
from django.db.models import Value
from django.db.models.functions import Concat
class LeaseCreateView(LoginRequiredMixin, LeaseTenantOrderMixin, CreateView):
    model = Lease
    form_class = LeaseForm
    template_name = 'leases/lease_form.html'

    def get_success_url(self):
        return reverse('leases:lease_detail', kwargs={'pk': self.object.pk})

    # ---------- Form setup ----------
    def get_form(self, form_class=None):
        form = super().get_form(form_class)

        # tenant ordering
        _order_tenants(form)
        _seed_units_queryset(form, self.request)

        from tenants.models import Tenant
        form.fields["tenant"].queryset = (
            Tenant.objects
            .annotate(full_name=Concat("first_name", Value(" "), "last_name"))
            .order_by(Lower("full_name"), "id")
        )

        # property -> pre-select + filter units
        prop_id = self.request.POST.get('property') or self.request.GET.get('property')
        if 'property' in form.fields:
            form.fields['property'].initial = prop_id

        form.fields['unit'].queryset = (
            Unit.objects.filter(property_id=prop_id).order_by('unit_number')
            if prop_id else Unit.objects.none()
        )
        return form

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)

        # family formset (create mode: instance is None)
        ctx.setdefault(
            "family_formset",
            LeaseFamilyFormSet(
                self.request.POST or None,
                prefix="family_members",
                instance=getattr(self, "object", None),
            )
        )
        ctx.setdefault(
            "tenants_for_add",
            Tenant.objects
                  .annotate(full_name=Concat("first_name", Value(" "), "last_name"))
                  .order_by(Lower("full_name"), "id")
        )
        return ctx

    # ---------- Main SAVE logic ----------
    def form_valid(self, form):
        """
        CREATE logic:

        - First click (no confirm_billing): preview initial billing in modal.
        - After user confirms (confirm_billing=1):
            * Save lease
            * Create REQUIRED security transaction
            * Create initial billing
            * Save family members / quick-add
        """
        confirmed = self.request.POST.get("confirm_billing") == "1"
        debug_sql = self.request.GET.get("debug_sql") == "1"

        # Family formset
        family_fs = LeaseFamilyFormSet(
            self.request.POST, prefix="family_members", instance=form.instance
        )
        if not family_fs.is_valid():
            messages.error(
                self.request, "Family form has errors. Please fix and try again."
            )
            ctx = self.get_context_data(form=form)
            ctx["family_formset"] = family_fs
            return self.render_to_response(ctx)

        # ---------- 1st click: PREVIEW ----------
        if not confirmed:
            # Temporarily save to compute billing, then roll back
            with transaction.atomic():
                tmp = form.save()
                plan = preview_initial_billing(tmp)
                transaction.set_rollback(True)

            # Keep in create mode (no pk) – don't show "created" yet
            form.instance.pk = None

            ctx = self.get_context_data(form=form)
            ctx["family_formset"] = family_fs
            ctx.update({
                "object": form.instance,
                "billing_plan": plan,
                "auto_open_billing_modal": _plan_has_changes(plan),
                "is_create": True,
            })
            return self.render_to_response(ctx)

        # ---------- 2nd click: CONFIRMED ----------
        # Actually save the lease
        response = super().form_valid(form)   # self.object is now saved

        # Family links
        family_fs.instance = self.object
        family_fs.save()
        self._handle_quick_add(self.object)

        # Create initial REQUIRED security deposit transaction *after* lease exists
        create_initial_security_required(self.object)

        # Read user choices from modal for billing
        include_backfill = self.request.POST.get("include_backfill") == "1"
        update_existing = self.request.POST.get("update_existing") == "1"

        if debug_sql:
            # Wrap billing creation in SQL logger
            sql = _run_sql(lambda: apply_initial_billing(
                self.object,
                include_backfill=include_backfill,
                update_existing=update_existing,
            ))
            messages.success(self.request, "Initial billing created (debug_sql).")
            return render(self.request, self.template_name, {
                "form": form,
                "object": self.object,
                "sql_log": sql,
            })

        # Normal path
        try:
            apply_initial_billing(
                self.object,
                include_backfill=include_backfill,
                update_existing=update_existing,
            )
            messages.success(self.request, "Lease created. Billing & security initialized.")
        except Exception as e:
            print(f"ERROR in apply_initial_billing: {e}")
            messages.error(self.request, f"Lease created, but billing failed: {e}")

        return redirect(self.get_success_url())

    # ---------- Quick-add family members ----------
    def _handle_quick_add(self, lease):
        total = int(self.request.POST.get('qa-TOTAL', 0) or 0)
        created_links = 0
        for i in range(total):
            tenant_id = self.request.POST.get(f'qa-{i}-tenant_id')
            relation = (self.request.POST.get(f'qa-{i}-relation', '') or '').strip()
            wa = self.request.POST.get(f'qa-{i}-whatsapp') == '1'
            if not tenant_id:
                continue
            try:
                tenant = Tenant.objects.get(pk=tenant_id)
            except Tenant.DoesNotExist:
                continue

            link, created = LeaseFamily.objects.get_or_create(
                lease=lease,
                tenant=tenant,
                defaults={'relation': relation, 'whatsapp_opt_in': wa}
            )
            if not created:
                # Optional: update relation/whatsapp if changed
                changed = False
                if relation and link.relation != relation:
                    link.relation = relation
                    changed = True
                if link.whatsapp_opt_in != wa:
                    link.whatsapp_opt_in = wa
                    changed = True
                if changed:
                    link.save(update_fields=['relation', 'whatsapp_opt_in'])
            else:
                created_links += 1

        if created_links:
            messages.success(
                self.request, f'Added {created_links} family member(s).'
            )

    # ---------- Normal form_invalid ----------
    def form_invalid(self, form):
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({'status': 'error', 'errors': form.errors}, status=400)

        print("Form is invalid in LeaseCreateView")
        print(f"Form errors: {form.errors}")
        return super().form_invalid(form)




def get_units(request):
    property_id = request.GET.get('property_id')
    units = Unit.objects.filter(
        property_id=property_id).values('id', 'unit_number')
    return JsonResponse(list(units), safe=False)


logger = logging.getLogger(__name__)

# leases/views.py (only the relevant parts)


def _lease_success_url(obj):
    return reverse("leases:lease_detail", kwargs={"pk": obj.pk})


def _plan_has_changes(plan: dict) -> bool:
    if not plan:
        return False
    if plan.get("requires_security_confirmation"):
        return True
    sec = plan.get("security_item") or {}
    if sec and (sec.get("before") != sec.get("after")):
        return True
    if plan.get("invoices_to_create"):
        return True
    if plan.get("recurring"):
        return True
    return False


class _SQLCol:
    def __init__(self): self.q = []

    def __call__(self, execute, sql, params, many, ctx):
        t0 = time.perf_counter()
        try:
            return execute(sql, params, many, ctx)
        finally:
            self.q.append({"sql": sql, "params": params,
                          "ms": round((time.perf_counter()-t0)*1000, 2)})


def _run_sql(fn):
    col = _SQLCol()
    with connection.execute_wrapper(col):
        fn()
    return col.q


# leases/views.py


@login_required
@require_POST
def lease_family_add(request, pk):
    lease = get_object_or_404(Lease, pk=pk)
    tenant_id = request.POST.get("tenant_id")
    relation = (request.POST.get("relation") or "").strip()

    if not tenant_id:
        messages.error(request, "Please choose a tenant.")
        return redirect("leases:lease_update", pk=lease.pk)

    tenant = get_object_or_404(Tenant, pk=tenant_id)

    link, created = LeaseFamily.objects.get_or_create(
        lease=lease, tenant=tenant, defaults={"relation": relation}
    )
    # If already linked and you changed relation, update it
    if not created and relation and link.relation != relation:
        link.relation = relation
        link.save(update_fields=["relation"])

    messages.success(request, f"{tenant} added to family.")
    # Optional JSON support:
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return JsonResponse({"ok": True, "id": link.id, "tenant": str(tenant), "relation": link.relation})
    return redirect("leases:lease_update", pk=lease.pk)

# leases/views.py
from decimal import Decimal
from django.contrib import messages
from django.db import transaction
from django.shortcuts import redirect, render
from django.urls import reverse
from django.views.generic import UpdateView
from django.db.models.functions import Lower

from .models import Lease, LeaseFamily
from .forms import LeaseForm, LeaseFamilyFormSet
from tenants.models import Tenant
from properties.models import Unit,Property

from leases.utils.billing import (
    preview_billing_on_change,
    update_billing_on_change,
)
from invoices.services import security_deposit_totals   # for detecting changes / flags

# leases/utils/billing.py

from decimal import Decimal

def detect_lease_changes(old_lease, new_lease) -> Dict[str, Any]:
    """
    Compare two Lease instances (old from DB, new from form.instance) and
    tell us what changed for billing/security logic.
    """
    ZERO = Decimal("0.00")

    old_rent   = old_lease.monthly_rent or ZERO
    new_rent   = new_lease.monthly_rent or ZERO
    old_maint  = old_lease.society_maintenance or ZERO
    new_maint  = new_lease.society_maintenance or ZERO
    old_sec    = old_lease.security_deposit or ZERO
    new_sec    = new_lease.security_deposit or ZERO

    return {
        "security_changed":    (old_sec   != new_sec),
        "rent_changed":        (old_rent  != new_rent),
        "maintenance_changed": (old_maint != new_maint),
        "end_date_changed":    (old_lease.end_date != new_lease.end_date),

        # optional: handy numbers for the modal / messages
        "old_required": old_sec,
        "new_required": new_sec,
        "old_rent_total": old_rent + old_maint,
        "new_rent_total": new_rent + new_maint,
        "old_end_date": old_lease.end_date,
        "new_end_date": new_lease.end_date,
    }


def create_security_required_adjustment(old_lease, new_lease):
    """
    Write an ADJUST transaction when lease.security_deposit changes.
    Positive amount = tenant owes more; negative = tenant owes less.
    """
    from invoices.models import SecurityDepositTransaction  # local import

    ZERO = Decimal("0.00")
    old_required = old_lease.security_deposit or ZERO
    new_required = new_lease.security_deposit or ZERO
    delta = new_required - old_required
    if delta == ZERO:
        return None

    # +delta means you need to collect more, -delta means you reduced requirement
    return SecurityDepositTransaction.objects.create(
        lease=new_lease,
        type="ADJUST",
        amount=delta.copy_abs(),
        notes=f"Required security changed from {old_required} to {new_required}",
    )

class LeaseUpdateView(LoginRequiredMixin, LeaseTenantOrderMixin, UpdateView):
    model = Lease
    form_class = LeaseForm
    template_name = 'leases/lease_form.html'

    def get_success_url(self):
        return reverse('leases:lease_detail', kwargs={'pk': self.object.pk})

    def get_context_data(self, **kwargs):
        """
        Always provide `family_formset` so the template can render the
        management_form and rows. On GET it's unbound; on POST it's bound.
        """
        ctx = super().get_context_data(**kwargs)

        # Figure out which lease instance to attach to the formset
        form = kwargs.get("form")
        if form is not None:
            lease_instance = form.instance
        else:
            lease_instance = getattr(self, "object", None)

        ctx.setdefault(
            "family_formset",
            LeaseFamilyFormSet(
                self.request.POST or None,
                prefix="family_members",
                instance=lease_instance,
            )
        )
        ctx.setdefault(
            "tenants_for_add",
            Tenant.objects
                  .annotate(full_name=Concat("first_name", Value(" "), "last_name"))
                  .order_by(Lower("full_name"), "id")
        )        
        return ctx

    # ---------- Quick-add family members (same as in LeaseCreateView) ----------
    def _handle_quick_add(self, lease):
        total = int(self.request.POST.get('qa-TOTAL', 0) or 0)
        created_links = 0
        for i in range(total):
            tenant_id = self.request.POST.get(f'qa-{i}-tenant_id')
            relation = (self.request.POST.get(f'qa-{i}-relation', '') or '').strip()
            wa = self.request.POST.get(f'qa-{i}-whatsapp') == '1'
            if not tenant_id:
                continue
            try:
                tenant = Tenant.objects.get(pk=tenant_id)
            except Tenant.DoesNotExist:
                continue

            link, created = LeaseFamily.objects.get_or_create(
                lease=lease,
                tenant=tenant,
                defaults={'relation': relation, 'whatsapp_opt_in': wa}
            )
            if not created:
                # Optional: update relation/whatsapp if changed
                changed = False
                if relation and link.relation != relation:
                    link.relation = relation
                    changed = True
                if link.whatsapp_opt_in != wa:
                    link.whatsapp_opt_in = wa
                    changed = True
                if changed:
                    link.save(update_fields=['relation', 'whatsapp_opt_in'])
            else:
                created_links += 1

        if created_links:
            messages.success(
                self.request, f'Added {created_links} family member(s).'
            )

    def form_valid(self, form):
        confirmed = self.request.POST.get("confirm_billing") == "1"
        debug_sql = self.request.GET.get("debug_sql") == "1"

        # Snapshot BEFORE (from DB)
        old = Lease.objects.get(pk=form.instance.pk)

        # Detect changes (uses unsaved form.instance values)
        changes = detect_lease_changes(old, form.instance)
        security_changed = changes["security_changed"]
        rent_changed = changes["rent_changed"]
        end_date_changed = changes["end_date_changed"]

        qa_total = int(self.request.POST.get("qa-TOTAL") or 0)
        important_change = security_changed or rent_changed or end_date_changed

        # ---------- STEP 1: PREVIEW PATH ----------
        if not confirmed and important_change and qa_total == 0:
            with transaction.atomic():
                tmp = form.save(commit=False)
                plan = preview_billing_on_change(tmp, old)
                transaction.set_rollback(True)

            ctx = self.get_context_data(form=form)

            # For preview, show an *unbound* family formset (so it doesn't complain)
            ctx["family_formset"] = LeaseFamilyFormSet(
                instance=old, prefix="family_members"
            )

            ctx.update({
                "object": form.instance,
                "billing_plan": plan,
                "auto_open_billing_modal": True,
                "is_create": False,
                "lease_change_flags": changes,
            })
            return self.render_to_response(ctx)

        # ---------- STEP 2: FINAL VALIDATION ----------
        family_fs = LeaseFamilyFormSet(
            self.request.POST, prefix="family_members", instance=form.instance
        )
        if not family_fs.is_valid():
            messages.error(
                self.request, "Family form has errors. Please fix and try again."
            )
            ctx = self.get_context_data(form=form)
            ctx["family_formset"] = family_fs
            return self.render_to_response(ctx)

        # If quick-add rows exist, treat this as confirmed
        if not confirmed and qa_total > 0:
            confirmed = True

        # ---------- STEP 3: SAVE LEASE & FAMILY ----------
        response = super().form_valid(form)  # self.object is now saved

        family_fs.instance = self.object
        family_fs.save()
        self._handle_quick_add(self.object)

        # ---------- STEP 4: SECURITY LEDGER ADJUSTMENT ----------
        if security_changed:
            create_security_required_adjustment(old, self.object)

        # ---------- STEP 5: APPLY BILLING CHANGES ----------
        include_backfill = self.request.POST.get("include_backfill") == "1"
        update_existing = self.request.POST.get("update_existing") == "1"

        if debug_sql:
            # use the _run_sql helper already defined above
            sql = _run_sql(lambda: update_billing_on_change(
                self.object,
                old,
                confirm_security_update=True,
                include_backfill=include_backfill,
                update_existing=update_existing,
            ))
            messages.success(self.request, "Billing updated (debug_sql).")
            return render(self.request, self.template_name, {
                "form": form,
                "object": self.object,
                "sql_log": sql,
            })

        update_billing_on_change(
            self.object,
            old,
            confirm_security_update=True,
            include_backfill=include_backfill,
            update_existing=update_existing,
        )
        messages.success(self.request, "Lease updated. Billing & security synced.")
        return redirect(self.get_success_url())


class LeaseDeleteView(DeleteView):
    model = Lease
    template_name = 'leases/lease_confirm_delete.html'
    success_url = reverse_lazy('leases:lease_list')


class GetUnitsView(View):
    """AJAX view for getting units by property"""

    def get(self, request, *args, **kwargs):
        property_id = request.GET.get('property_id')
        units = Unit.objects.none()

        if property_id:
            units = Unit.objects.filter(
                property_id=property_id
            ).values('id', 'unit_number')

        return JsonResponse(list(units), safe=False)


class RenewLeaseView(LoginRequiredMixin, View):
    def get(self, request, pk):
        lease = get_object_or_404(Lease, pk=pk)

        if lease.status != 'active':
            messages.error(request, "Only active leases can be renewed")
            return redirect('leases:lease_detail', pk=pk)

        renewed_lease = lease.renew_lease(
            notes=f"Manually renewed by {request.user.username}"
        )

        if renewed_lease:
            messages.success(
                request, f"Lease successfully renewed as #{renewed_lease.id}")
            return redirect('leases:lease_detail', pk=renewed_lease.pk)
        else:
            messages.error(request, "Failed to renew lease")
            return redirect('leases:lease_detail', pk=pk)


class PrintLeaseView(View):
    def get(self, request, pk):
        lease = get_object_or_404(Lease, pk=pk)

        # Render template
        html_string = render_to_string('leases/lease_pdf_detail.html', {
            'lease': lease,
            'static_url': '/static/',
            'media_url': '/media/'
        })

        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=True) as output:
            HTML(string=html_string, base_url=request.build_absolute_uri(
                '/')).write_pdf(output)
            output.flush()
            output.seek(0)
            pdf = output.read()

        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="lease_{lease.id}.pdf"'
        return response


class UnitAutocomplete(autocomplete.Select2QuerySetView):
    """Autocomplete view for units"""

    def get_queryset(self):
        qs = Unit.objects.all()

        if self.q:
            qs = qs.filter(unit_number__icontains=self.q)

        if property_id := self.forwarded.get('property'):
            qs = qs.filter(property_id=property_id)

        return qs


def generate_lease_pdf(request, pk):
    """Generate lease PDF using ReportLab"""
    lease = get_object_or_404(Lease, pk=pk)

    buffer = BytesIO()
    p = canvas.Canvas(buffer)

    # Draw PDF content
    p.drawString(100, 800, f"LEASE AGREEMENT #{lease.id}")
    p.drawString(100, 780, f"Property: {lease.property}")
    p.drawString(100, 760, f"Unit: {lease.unit}")
    # Add more content...

    p.showPage()
    p.save()

    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="lease_{lease.id}.pdf"'
    return response

class LeaseDetailView(LoginRequiredMixin, DetailView):
    model = Lease
    template_name = 'leases/lease_detail.html'
    context_object_name = 'lease'

    def get_queryset(self):
        return super().get_queryset().select_related(
            'tenant', 'unit', 'unit__property'
        ).prefetch_related(
            'payments', 'invoices', 'family_members__tenant'
        )

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        lease = self.object

        # existing context
        ctx.update({
            'family_members': lease.family_members.all(),
            'recent_payments': lease.payments.all().order_by('-payment_date')[:5],
            'invoices': lease.invoices.all().order_by('-issue_date'),
            'balance_due': lease.get_balance,
            'occupancy_count': 1 + lease.family_members.count(),
        })

        # ✅ Security deposit totals + ledger
        ZERO = Decimal('0.00')
        sec_totals = security_deposit_totals(lease)

        dep_balance = ZERO
        deposit_tx = []

        for tx in SecurityDepositTransaction.objects.filter(
            lease=lease
        ).order_by('date', 'id'):
            amt = tx.amount or ZERO

            # running balance of what you currently hold
            if tx.type == 'PAYMENT':
                dep_balance += amt
                signed_amt = amt
            elif tx.type in ('REFUND', 'DAMAGE'):
                dep_balance -= amt
                signed_amt = -amt
            else:
                # REQUIRED / ADJUST etc. – doesn’t change “held” balance
                signed_amt = amt

            deposit_tx.append({
                'date': tx.date,
                'type': tx.get_type_display(),
                'description': tx.notes or '',
                'amount': signed_amt,
                'balance': dep_balance,
            })

        ctx.update({
            'deposit_transactions': deposit_tx,
            'security_required': sec_totals['required'],
            'security_paid_in': sec_totals['paid_in'],
            'security_refunded': sec_totals['refunded'],
            'security_damages': sec_totals['damages'],
            'security_balance_to_collect': sec_totals['balance_to_collect'],
            'security_currently_held': sec_totals['currently_held'],
        })

        return ctx

def lease_print(request, pk):
    lease = get_object_or_404(Lease, pk=pk)
    html_string = render_to_string('leases/lease_print.html', {'lease': lease})
    return HttpResponse(html_string)


def lease_email(request, pk):
    lease = get_object_or_404(Lease, pk=pk)
    # Implement your email sending logic here
    return HttpResponse("Lease email sent")


@api_view(['GET'])
def lease_balance(request, pk):
    lease = get_object_or_404(Lease, pk=pk)
    return Response({
        'balance_due': lease.get_balance_due()
    })


# leases/views.py

# If not already imported in this file:
# from tenants.models import Tenant, normalize_cnic


def _last4(s: str) -> str:
    digits = ''.join(ch for ch in (s or '') if ch.isdigit())
    return digits[-4:] if len(digits) >= 4 else (digits or '')


# leases/views.py

_digits = re.compile(r'\D+')


def only_digits(s: str) -> str:
    return _digits.sub('', s or '')


@login_required
@require_GET
def tenant_quick_search(request):
    q = (request.GET.get("q") or "").strip()
    want_all = request.GET.get("all") == "1"
    q_digits = only_digits(q)

    # If you want some default options before typing:
    if not q and want_all:
        qs = Tenant.objects.order_by('first_name', 'last_name', 'id')[:50]
    elif not q:
        return JsonResponse([], safe=False)
    else:
        # Build a flexible filter over fields that actually exist
        filters = (
            Q(first_name__icontains=q) |
            Q(last_name__icontains=q)
        )
        # phones (search by text and digit-stripped)
        for pf in ["phone", "phone2", "phone3"]:
            filters |= Q(**{f"{pf}__icontains": q})
            if q_digits:
                filters |= Q(**{f"{pf}__icontains": q_digits})
        # CNIC (both raw and digits-only shadow field)
        for cf in ["cnic", "cnic_digits"]:
            filters |= Q(**{f"{cf}__icontains": q})
            if q_digits:
                filters |= Q(**{f"{cf}__icontains": q_digits})

        qs = Tenant.objects.filter(filters).order_by(
            'first_name', 'last_name', 'id')[:50]

    def full_name(t: Tenant) -> str:
        return f"{t.first_name} {t.last_name}".strip()

    payload = [{
        "id": t.id,
        # Select2 'text' is set client-side from this
        "name": full_name(t),
        "phone1": t.phone or "",         # map model 'phone' → API 'phone1'
        "phone2": t.phone2 or "",
        "cnic": t.cnic or t.cnic_digits or "",
    } for t in qs]

    return JsonResponse(payload, safe=False)


def get_tenants(request):
    include_inactive = request.GET.get('include_inactive', False)
    tenants = Tenant.objects.all()
    if not include_inactive:
        tenants = tenants.filter(is_active=True)
    tenants = tenants.order_by('first_name', 'last_name')
    data = {
        'tenants': [{
            'id': t.id,
            'first_name': t.first_name,
            'last_name': t.last_name
        } for t in tenants]
    }
    return JsonResponse(data)


from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from .models import Unit  # adjust import if Unit is in another app

@login_required
def get_units_by_property(request):
    """
    Return units for a given property as JSON:
    { "units": [ { "id": 1, "unit_number": "A-101" }, ... ] }
    """
    prop_id = request.GET.get("property_id")
    if not prop_id:
        return JsonResponse({"units": []})

    qs = Unit.objects.filter(property_id=prop_id).order_by("unit_number")
    data = [{"id": u.id, "unit_number": u.unit_number} for u in qs]
    return JsonResponse({"units": data})



def get_lease_info(request):
    tenant_id = request.GET.get('tenant_id')
    unit_id = request.GET.get('unit_id')

    try:
        lease = Lease.objects.get(
            tenant_id=tenant_id, unit_id=unit_id, status='active')
        today = timezone.now().date()
        is_ending_soon = (lease.end_date - today).days <= 40

        data = {
            'lease': {
                'end_date': lease.end_date.strftime('%b %d, %Y'),
                'monthly_rent': str(lease.monthly_rent),
                'security_deposit': str(lease.security_deposit),
                'balance_due': str(lease.get_balance_due()),
                'maintenance_due': str(lease.get_maintenance_due()),
                'is_ending_soon': is_ending_soon
            }
        }
    except Lease.DoesNotExist:
        data = {'lease': None}

    return JsonResponse(data)


def get_lease(self):
    lease_id = self.request.GET.get('lease_id')
    if lease_id:
        return get_object_or_404(Lease, id=lease_id)
    return None


class DashboardView(LoginRequiredMixin, TemplateView):
    template_name = 'leases/dashboard.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Get leases ending in the next 30 days
        ending_soon_date = timezone.now().date() + timedelta(days=30)
        ending_leases = Lease.objects.filter(
            status='active',
            end_date__lte=ending_soon_date,
            end_date__gte=timezone.now().date()
        ).select_related('tenant', 'unit', 'unit__property')

        context['ending_leases'] = ending_leases
        return context


class RenewLeaseView(LoginRequiredMixin, FormView):
    template_name = 'leases/renew_lease.html'
    form_class = RenewLeaseForm

    def get_initial(self):
        initial = super().get_initial()
        lease = get_object_or_404(Lease, pk=self.kwargs['pk'])
        initial['rent_increase_percent'] = lease.rent_increase_percent
        return initial

    def form_valid(self, form):
        lease = get_object_or_404(Lease, pk=self.kwargs['pk'])
        renewed_lease = lease.renew_lease(
            new_end_date=form.cleaned_data['new_end_date'],
            rent_increase_percent=float(
                form.cleaned_data['rent_increase_percent']),
            notes=form.cleaned_data['notes']
        )

        if renewed_lease:
            messages.success(
                self.request, f"Lease successfully renewed as #{renewed_lease.id}")
            return redirect('leases:lease_detail', pk=renewed_lease.pk)

        messages.error(self.request, "Failed to renew lease")
        return redirect('leases:lease_detail', pk=lease.pk)


class CustomRenewView(LoginRequiredMixin, FormView):
    template_name = 'leases/custom_renew.html'
    form_class = CustomRenewForm

    def dispatch(self, request, *args, **kwargs):
        self.lease = get_object_or_404(Lease, pk=kwargs['pk'])
        return super().dispatch(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['lease'] = self.lease
        return kwargs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['lease'] = self.lease
        return context

    def form_valid(self, form):
        # Get the society maintenance from the form
        society_maintenance = form.data.get(
            'society_maintenance', self.lease.society_maintenance)

        # Renew the lease with custom parameters
        renewed_lease = self.lease.renew_lease(
            new_end_date=form.cleaned_data['new_end_date'],
            rent_increase_percent=float(
                form.cleaned_data['rent_increase_percent']),
            society_maintenance=society_maintenance,
            notes=form.cleaned_data['notes']
        )

        if renewed_lease:
            messages.success(
                self.request,
                f"Lease successfully renewed as #{renewed_lease.id} with "
                f"{form.cleaned_data['rent_increase_percent']}% increase"
            )
            return redirect('leases:lease_detail', pk=renewed_lease.pk)

        messages.error(self.request, "Failed to renew lease")
        return redirect('leases:lease_detail', pk=self.lease.pk)


# views.py (imports you likely already have; add any missing ones)


# If you have a separate Transaction model, you don't need it here since we build from invoices/payments.
# from .models import Transaction

# If this is defined elsewhere, keep your original import
# from .utils import prepare_transaction_columns

def prepare_transaction_columns(transactions):
    """
    Keep your existing implementation; this stub is here only for context.
    """
    return [
        {"key": "date", "label": "Date"},
        {"key": "type", "label": "Type"},
        {"key": "description", "label": "Description"},
        {"key": "amount", "label": "Amount"},
        {"key": "balance", "label": "Balance"},
    ]


# views.py


def _truncate15(s: str) -> str:
    if not s:
        return s
    return (s[:15] + "…") if len(s) > 15 else s





class LeaseLedgerView(LoginRequiredMixin, TemplateView):
    model = Lease
    template_name = 'leases/lease_ledger.html'

    def _build_filter_context(self):
        """
        Build the compact toolbar context:
        - property, unit
        - lease candidates (depends on property/unit and 'show_inactive')
        - current selections
        """
        req = self.request
        today = date.today()

        property_id = req.GET.get('property') or ''
        unit_id = req.GET.get('unit') or ''
        # we keep 'lease_id' for deep-link & selection
        url_pk = self.kwargs.get('pk')
        lease_id = req.GET.get('lease_id') or (str(url_pk) if url_pk else '')
        show_inactive = (req.GET.get('show_inactive') == 'on')

        # Properties
        all_properties = Property.objects.all().order_by('property_name')

        # Units (optionally narrowed by property)
        units_qs = Unit.objects.select_related('property') \
            .order_by('property__property_name', 'unit_number')
        if property_id:
            units_qs = units_qs.filter(property_id=property_id)

        # Build lease candidate list based on property/unit & show_inactive
        leases_qs = Lease.objects.select_related(
            'tenant', 'unit', 'unit__property')

        if property_id:
            leases_qs = leases_qs.filter(unit__property_id=property_id)
        if unit_id:
            leases_qs = leases_qs.filter(unit_id=unit_id)

        if not show_inactive:
            leases_qs = leases_qs.filter(
                start_date__lte=today, end_date__gte=today)

        # Order newest first so the most recent shows first when multiple
        leases_qs = leases_qs.order_by('-start_date', '-id')

        # Auto-select rule: if a unit is selected and there is exactly one lease candidate,
        # and either !show_inactive OR (show_inactive but still only one) -> select it.
        # If show_inactive & there are multiple, do NOT auto-select.
        lease_candidates = list(leases_qs)
        if unit_id and len(lease_candidates) == 1 and not lease_id:
            lease_id = str(lease_candidates[0].id)

        # Build display labels (Tenant(≤15) — Prop-Unit • Ends: MMM DD, YYYY)
        lease_choices = []
        for l in lease_candidates:
            tn = _truncate15(l.tenant.get_full_name())
            unit_str = f"{l.unit.property.property_name}-{l.unit.unit_number}" if l.unit and l.unit.property else ""
            end_str = l.end_date.strftime("%b %d, %Y") if l.end_date else "N/A"
            label = f"{tn} — {unit_str} • Ends: {end_str}"
            lease_choices.append({'id': l.id, 'label': label})

        return {
            'all_properties': all_properties,
            'all_units': units_qs,
            'lease_choices': lease_choices,
            'current_property': property_id,
            'current_unit': unit_id,
            'current_lease': lease_id,          # used by select & deep-link
            'show_inactive': show_inactive,     # checkbox state
        }

    def _pick_lease(self, fc):
        """
        Choose the lease to display:
        - explicit lease_id (from URL or GET) wins
        - else None (we rely on user to pick from Lease picker)
        """
        if fc['current_lease']:
            return get_object_or_404(
                Lease.objects.select_related(
                    'tenant', 'unit', 'unit__property'),
                pk=fc['current_lease']
            )
        return None






    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        fc = self._build_filter_context()
        lease = self._pick_lease(fc)

        # Help banners
        property_only = bool(
            fc['current_property'] and not fc['current_unit'] and not fc['current_lease']
        )
        unit_only = bool(fc['current_unit'] and not fc['current_lease'])

        ZERO = Decimal('0.00')

        # ---------- Security deposit: defaults ----------
        sec_totals = {
            "required": ZERO,
            "paid_in": ZERO,
            "refunded": ZERO,
            "damages": ZERO,
            "balance_to_collect": ZERO,
            "currently_held": ZERO,
        }
        deposit_tx = []

        # If we don't have a lease yet, just return empty ledger + default security
        if not lease:
            ctx.update({
                **fc,
                "lease": None,
                "transactions": [],
                "transaction_columns": [],
                "total_paid": ZERO,
                "total_owed": ZERO,
                "current_balance": ZERO,

                # security deposit context (all zero)
                "deposit_transactions": deposit_tx,
                "security_required": sec_totals["required"],
                "security_paid_in": sec_totals["paid_in"],
                "security_refunded": sec_totals["refunded"],
                "security_damages": sec_totals["damages"],
                "security_balance_to_collect": sec_totals["balance_to_collect"],
                "security_currently_held": sec_totals["currently_held"],

                "intcomma": intcomma,
                "generated_on": datetime.now(),
                "property_only": property_only,
                "unit_only": unit_only,
                "page": 1,
                "total_pages": 1,
                "has_prev": False,
                "has_next": False,
                "tx_col1": [],
                "tx_col2": [],
            })
            return ctx

        # ---------- Security deposit summary (with lease) ----------
        sec_totals = security_deposit_totals(lease)

        dep_balance = ZERO
        for tx in SecurityDepositTransaction.objects.filter(lease=lease).order_by('date', 'id'):
            amt = tx.amount or ZERO

            # running balance = what you currently hold
            if tx.type == 'PAYMENT':
                dep_balance += amt
                signed_amt = amt
            elif tx.type in ('REFUND', 'DAMAGE'):
                dep_balance -= amt
                signed_amt = -amt
            else:  # e.g. REQUIRED / ADJUSTMENT if you use those
                signed_amt = amt

            deposit_tx.append({
                "date": tx.date,
                "type": tx.get_type_display(),
                "description": tx.notes or "",
                "amount": signed_amt,
                "balance": dep_balance,
            })

        # ---------- Main rent ledger (your existing logic) ----------
        transactions = []
        balance = ZERO

        for invoice in lease.invoices.all():
            amt = invoice.amount or ZERO
            balance -= amt
            transactions.append({
                'date': invoice.issue_date,
                'type': 'Invoice',
                'description': invoice.description,
                'amount': -amt,
                'balance': balance,
                'url': reverse('invoices:invoice_detail', args=[invoice.id]),
            })

        # IMPORTANT: Lease ledger must use allocation.lease_amount (NOT payment.amount)
        for payment in lease.payments.all().select_related("allocation").order_by("payment_date", "id"):
            alloc = getattr(payment, "allocation", None)

            # lease portion only
            lease_amt = (getattr(alloc, "lease_amount", None) if alloc else None)
            amt = (lease_amt if lease_amt is not None else (payment.amount or ZERO)) or ZERO

            balance += amt
            transactions.append({
                "date": payment.payment_date,
                "type": "Payment",
                "description": payment.reference_number or f"Payment #{payment.id}",
                "amount": amt,
                "balance": balance,
                # keep payment detail url OR switch to allocation detail if you prefer
                "url": (
                    reverse("payments:allocation_detail", args=[alloc.id])
                    if alloc else reverse("payments:payment_detail", args=[payment.id])
                ),

            })



        transactions.sort(key=lambda x: (x['date'], 0 if x['type'] == 'Invoice' else 1))
        balance = ZERO
        for t in transactions:
            balance += t['amount']
            t['balance'] = balance

        def prepare_transaction_columns(transactions):
            return [
                {"key": "date", "label": "Date"},
                {"key": "type", "label": "Type"},
                {"key": "description", "label": "Description"},
                {"key": "amount", "label": "Amount"},
                {"key": "balance", "label": "Balance"},
            ]

        # ---------- Pagination + 2-column split ----------
        page = int(self.request.GET.get('page', 1))
        PAGE_SIZE = 40
        total_count = len(transactions)
        total_pages = max(1, ceil(total_count / PAGE_SIZE))
        page = max(1, min(page, total_pages))

        start = (page - 1) * PAGE_SIZE
        end = start + PAGE_SIZE
        page_tx = transactions[start:end]

        col1 = page_tx[:20]
        col2 = page_tx[20:40]

        ctx.update({
            **fc,
            "lease": lease,
            "transactions": transactions,
            "transaction_columns": prepare_transaction_columns(transactions),
            "total_paid": sum(t["amount"] for t in transactions if t["type"] == "Payment" and t["amount"] > 0),

            "total_owed": sum(-t["amount"] for t in transactions if t["amount"] < 0),
            "current_balance": balance,

            # security deposit context
            "deposit_transactions": deposit_tx,
            "security_required": sec_totals["required"],
            "security_paid_in": sec_totals["paid_in"],
            "security_refunded": sec_totals["refunded"],
            "security_damages": sec_totals["damages"],
            "security_balance_to_collect": sec_totals["balance_to_collect"],
            "security_currently_held": sec_totals["currently_held"],

            "intcomma": intcomma,
            "generated_on": datetime.now(),
            "property_only": False,
            "unit_only": False,
            "transactions": page_tx,
            "tx_col1": col1,
            "tx_col2": col2,
            "page": page,
            "total_pages": total_pages,
            "has_prev": page > 1,
            "has_next": page < total_pages,
        })
        return ctx

def lease_ledger_pdf(request, lease_id):
    """Generate PDF ledger with proper formatting and filename"""
    try:
        lease = get_object_or_404(Lease, pk=lease_id)

        # Prepare transactions data
        transactions = []
        balance = Decimal('0.00')

        # Process invoices
        for invoice in lease.invoices.all().order_by('issue_date'):
            balance -= invoice.amount
            transactions.append({
                'date': invoice.issue_date,
                'type': 'Invoice',
                'description': invoice.description,
                'amount': -invoice.amount,
                'balance': balance,
            })

        # Process payments
        for payment in lease.payments.all().select_related("allocation").order_by("payment_date", "id"):
            alloc = getattr(payment, "allocation", None)
            lease_amt = (getattr(alloc, "lease_amount", None) if alloc else None)
            amt = (lease_amt if lease_amt is not None else (payment.amount or Decimal("0.00"))) or Decimal("0.00")

            balance += amt
            transactions.append({
                "date": payment.payment_date,
                "type": "Payment",
                "description": payment.reference_number or f"Payment #{payment.id}",
                "amount": amt,
                "balance": balance,
            })



        # Sort transactions by date
        transactions.sort(key=lambda x: (
            x['date'], 0 if x['type'] == 'Invoice' else 1))

        # Now compute running balance in sorted order
        balance = Decimal('0.00')
        for t in transactions:
            balance += t['amount']
            t['balance'] = balance

        # Prepare transaction columns
        transaction_columns = prepare_transaction_columns(transactions)

        context = {
            'lease': lease,
            'tenant': lease.tenant,
            'transactions': transactions,
            'transaction_columns': transaction_columns,  # Add this line
            'total_paid': sum(t['amount'] for t in transactions if t['type'] == 'Payment' and t['amount'] > 0),

            'total_owed': sum(-t['amount'] for t in transactions if t['amount'] < 0),
            'current_balance': balance,
            'date': datetime.now().date(),
            'generated_on': datetime.now().strftime("%B %d, %Y %I:%M %p"),
            'intcomma': intcomma,
        }

        # Render HTML
        html_string = render_to_string('leases/lease_ledger_pdf.html', context)

        # Generate PDF with proper CSS
        html = HTML(string=html_string)
        pdf_file = html.write_pdf()

        # Create filename with tenant name, unit, and date
        filename = f"{slugify(lease.tenant.get_full_name())}-ledger-{slugify(lease.unit.unit_number)}-{datetime.now().strftime('%Y-%m-%d')}.pdf"

        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    except Exception as e:
        return HttpResponse(f"Failed to generate PDF: {str(e)}", status=500)


def send_ledger_email(request, lease_id):
    lease = get_object_or_404(Lease, pk=lease_id)

    try:
        # Generate PDF
        pdf_response = lease_ledger_pdf(request, lease_id)
        pdf_content = pdf_response.content

        # Create email
        subject = f'Lease Ledger for {lease.tenant.get_full_name()} - Lease #{lease.id}'
        message = f'Please find attached the ledger for lease #{lease.id}.'
        from_email = settings.DEFAULT_FROM_EMAIL
        # Send to tenant and current user
        to_email = [lease.tenant.email, request.user.email]

        email = EmailMessage(
            subject,
            message,
            from_email,
            to_email,
        )

        # Attach PDF
        email.attach(
            f'lease_{lease.id}_ledger.pdf',
            pdf_content,
            'application/pdf'
        )

        email.send()

        messages.success(request, 'Ledger has been emailed successfully!')
        return redirect('leases:lease_ledger', lease_id=lease_id)

    except Exception as e:
        messages.error(request, f'Failed to send email: {str(e)}')
        return redirect('leases:lease_ledger', lease_id=lease_id)


def export_ledger_excel(request, lease_id):
    try:
        lease = get_object_or_404(Lease, id=lease_id)
        invoices = Invoice.objects.filter(lease=lease).order_by('issue_date')
        payments = Payment.objects.filter(lease=lease).order_by('payment_date')

        # Calculate running balance
        transactions = []
        balance = Decimal('0.00')

        # Process invoices
        for invoice in invoices:
            balance -= invoice.amount
            transactions.append({
                'date': invoice.issue_date,
                'type': 'Invoice',
                'description': invoice.description,
                'amount': -invoice.amount,
                'balance': balance,
            })

        # Process payments
        for payment in payments:
            balance += payment.amount
            transactions.append({
                'date': payment.payment_date,
                'type': 'Payment',
                'description': payment.reference_number or "Payment",
                'amount': payment.amount,
                'balance': balance,
            })

        # Sort by date
        transactions.sort(key=lambda x: (
            x['date'], 0 if x['type'] == 'Invoice' else 1))

        balance = Decimal('0.00')
        for t in transactions:
            balance += t['amount']
            t['balance'] = balance

        # Create workbook and worksheet
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Lease Ledger"

        # Styles
        header_font = Font(bold=True)
        header_fill = PatternFill(
            start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        centered = Alignment(horizontal='center')
        black_font = Font(color='000000')  # Black color for all text

        # Lease Information - Merged across all 5 columns
        # MAIN HEADING
        heading_text = f"Lease Ledger for {lease.tenant.get_full_name()} for {lease.unit.property.property_name} {lease.unit.unit_number}"

        ws.append([heading_text])
        ws.merge_cells('A1:F1')
        ws['A1'].font = Font(bold=True, size=14)
        ws['A1'].alignment = centered

        # Empty row before transactions
        ws.append([])

        # Tenant and Property Info--- 1st row
        ws.append([
            f"Rent: {lease.monthly_rent}",  # 1st colom
            # 2nd colom
            f"Security Deposit: {lease.security_deposit}",
            f"Deposit Returned: {'Yes' if lease.security_deposit_returned else 'No'}",
            f"Period: {lease.start_date.strftime('%b %d, %Y')} to {lease.end_date.strftime('%b %d, %Y')}",
            f"Status: {lease.get_status_display()}",

        ])

        # Dates and Financial Info---2nd row
        new_rent = lease.monthly_rent
        (1 + (lease.rent_increase_percent / 100)
         ) if lease.rent_increase_percent else lease.monthly_rent
        ws.append([

            # 1st colom, 2nd row
            f"Maintenance: {lease.society_maintenance or 0}",
            f"Deposit Paid: {'Yes' if lease.security_deposit_paid else 'No'}",
            f"Return Date: {lease.security_deposit_return_date.strftime('%b %d, %Y') if lease.security_deposit_returned else 'N/A'}",
            f"Rent Increase: {lease.rent_increase_percent or 0}%",
            f"Lease ID: {lease.id}",

        ])

        # Security Deposit Info- 3rd row
        ws.append([
            f"Total Payment: {lease.total_payment}",
            f"Return Amount: {lease.security_deposit_return_amount or 'N/A'}",
            f"New Rent: {new_rent:.2f}",




        ])

        # Notes (merged across all 5 columns)
        if lease.notes or lease.security_deposit_return_notes:
            notes = []
            if lease.notes:
                notes.append(f"Lease Notes: {lease.notes}")
            if lease.security_deposit_return_notes:
                notes.append(
                    f"Deposit Notes: {lease.security_deposit_return_notes}")

            notes_row = len(ws['A']) + 1
            ws.append([" | ".join(notes)])
            ws.merge_cells(f'A{notes_row}:E{notes_row}')

        # Empty row before "Transaction History"
        ws.append([])  # This will be row 7

        # Transaction History heading in row 8
        ws.append(["Transaction History"])  # Now row 8
        ws.merge_cells('A8:F8')
        ws['A8'].font = Font(bold=True, size=14)
        ws['A8'].alignment = centered

        headers = ["S.No", "Date", "Type", "Description", "Amount", "Balance"]
        ws.append(headers)

        # Apply header styles
        for cell in ws[9]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = centered
            cell.border = thin_border

        # Add transactions with running balance
        for index, t in enumerate(transactions, start=1):
            row = [
                index,
                t['date'].strftime('%b %d, %Y'),
                t['type'],
                t['description'],
                t['amount'],
                t['balance']
            ]
            ws.append(row)

        # Apply styles to all data rows (black text)
        for row in ws.iter_rows(min_row=9, max_row=ws.max_row):
            for cell in row:
                cell.font = black_font
                cell.border = thin_border
                if cell.column_letter in ['E', 'F']:  # Amount and Balance columns
                    cell.number_format = '#,##0.00'

        # Set custom column widths
        column_widths = {
            'A': 22,    # S.No
            'B': 24,   # Date
            'C': 20,   # Type
            'D': 32,   # Description
            'E': 15,   # Amount
            'F': 15    # Balance
        }

        for col, width in column_widths.items():
            ws.column_dimensions[col].width = width

        # Create filename with tenant name and today's date
        today = datetime.now().strftime("%m%d%Y")
        filename = f"Lease Ledger - {lease.tenant.get_full_name()} - {today}.xlsx"

        # Create response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        wb.save(response)
        return response

    except Exception as e:
        return HttpResponse(
            f"Error generating Excel file: {str(e)}",
            status=500,
            content_type='text/plain'
        )


class GenerateLeaseAgreementView(View):
    def post(self, request, pk):
        lease = get_object_or_404(Lease, pk=pk)
        try:
            # Call the PDF generation function with request parameter
            response = generate_agreement_pdf(request, pk)

            # If PDF generation was successful, get the updated lease object
            lease.refresh_from_db()

            msg = 'Lease agreement generated successfully!'

            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True,
                    'message': msg,
                    'docx_url': lease.generated_agreement_docx.url if lease.generated_agreement_docx else None,
                    'pdf_url': lease.generated_agreement_pdf.url if lease.generated_agreement_pdf else None,
                })

            messages.success(request, msg)
        except Exception as e:
            import traceback
            traceback.print_exc()
            msg = f'Error generating agreement: {str(e)}'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'message': msg}, status=400)
            messages.error(request, msg)

        return HttpResponseRedirect(reverse('leases:lease_detail', args=[pk]))


class SendAgreementEmailView(View):
    def post(self, request, pk):
        lease = get_object_or_404(Lease, pk=pk)
        recipient_email = request.POST.get('email', lease.tenant.email)

        if not recipient_email:
            messages.error(request, 'No email address provided')
            return HttpResponseRedirect(reverse('leases:lease_detail', args=[pk]))

        try:
            success = send_lease_agreement_email(lease, [recipient_email])
            if success:
                messages.success(
                    request, 'Agreement sent via email successfully!')
            else:
                messages.error(request, 'Failed to send agreement email')
        except Exception as e:
            messages.error(request, f'Error sending email: {str(e)}')

        return HttpResponseRedirect(reverse('leases:lease_detail', args=[pk]))


class UploadSignedCopyView(View):
    def post(self, request, pk):
        lease = get_object_or_404(Lease, pk=pk)
        if 'signed_copy' in request.FILES:
            lease.signed_copy = request.FILES['signed_copy']
            lease.save()
            messages.success(
                request, 'Signed agreement uploaded successfully!')
        return HttpResponseRedirect(reverse('leases:lease_detail', args=[pk]))


class LeaseTemplateListView(ListView):
    model = LeaseTemplate
    template_name = 'leases/template_list.html'
    context_object_name = 'templates'


class LeaseTemplateCreateView(CreateView):
    model = LeaseTemplate
    form_class = LeaseTemplateForm
    template_name = 'leases/template_form.html'
    success_url = reverse_lazy('leases:template_list')


class LeaseTemplateUpdateView(UpdateView):
    model = LeaseTemplate
    form_class = LeaseTemplateForm
    template_name = 'leases/template_form.html'
    success_url = reverse_lazy('leases:template_list')


class LeaseTemplateDeleteView(DeleteView):
    model = LeaseTemplate
    template_name = 'leases/template_confirm_delete.html'
    success_url = reverse_lazy('leases:template_list')


def set_default_template(request, pk):
    template = get_object_or_404(LeaseTemplate, pk=pk)
    template.is_default = True
    template.save()
    messages.success(request, f"'{template.name}' is now the default template")
    return redirect('leases:template_list')


# leases/views.py


def generate_lease_agreement(request, lease_id):
    lease = Lease.objects.get(id=lease_id)

    # Generate HTML content
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{ font-family: 'Times New Roman'; line-height: 1.6; }}
            .clause {{ margin-bottom: 15px; text-align: justify; }}
            .signature {{ margin-top: 100px; }}
            .footer {{ font-size: 0.8em; text-align: center; }}
        </style>
    </head>
    <body>
        <h1 style="text-align: center;">RENT AGREEMENT</h1>
        <p style="text-align: center;">This RENT AGREEMENT is made at Islamabad on this {timezone.now().strftime('%d-%m-%Y')}</p>
        
        <div>
            <p><strong>BETWEEN</strong></p>
            <p>{resolve_placeholders(lease, '[OWNER_NAME]')} holding CNIC NO. {resolve_placeholders(lease, '[OWNER_CNIC]')}</p>
            <p>{resolve_placeholders(lease, '[OWNER_ADDRESS]')} (Hereinafter called "Owner")</p>
            
            <p><strong>AND</strong></p>
            <p>{resolve_placeholders(lease, '[TENANT_NAME]')} holding CNIC NO. {resolve_placeholders(lease, '[TENANT_CNIC]')}</p>
            <p>{resolve_placeholders(lease, '[TENANT_ADDRESS]')} (Hereinafter called "Tenant")</p>
        </div>
        
        <div>
            {"".join(f'<p class="clause">{resolve_placeholders(lease, clause.template_text)}</p>' for clause in lease.clauses.order_by("clause_number"))}
        </div>

        <div class="signature">
            <div style="float: left; width: 40%;">
                <p>_________________________</p>
                <p>Owner: {resolve_placeholders(lease, '[OWNER_NAME]')}</p>
            </div>
            <div style="float: right; width: 40%;">
                <p>_________________________</p>
                <p>Tenant: {resolve_placeholders(lease, '[TENANT_NAME]')}</p>
            </div>
            <div style="clear: both;"></div>
        </div>

        <div class="footer">
            <p>Generated at: {timezone.now().strftime('%d-%m-%Y %H:%M:%S')}</p>
        </div>
    </body>
    </html>
    """

    # Generate PDF
    html = HTML(string=html_content)
    pdf = html.write_pdf()

    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="lease_agreement_{lease_id}.pdf"'
    return response


@require_http_methods(["GET", "POST"])
def edit_clauses(request, pk):
    lease = get_object_or_404(Lease, pk=pk)
    clauses = lease.clauses.all().order_by('clause_number')

    # ✅ Handle AJAX save
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        for key, value in request.POST.items():
            if key.startswith("clause_"):
                clause_id = key.split("_")[1]
                try:
                    clause = lease.clauses.get(id=clause_id)
                    clause.template_text = value
                    clause.save()
                except Exception as e:
                    continue

        return JsonResponse({"status": "success", "message": "Clauses updated successfully"})

    # ✅ Generate PDF if not already
    if not lease.generated_agreement_pdf or not os.path.exists(lease.generated_agreement_pdf.path):
        context = {
            'lease': lease,
            'property': lease.unit.property,
            'tenant': lease.tenant,
            'clauses': [
                {
                    'clause_number': clause.clause_number,
                    'text': do_replace_placeholders(clause.template_text, lease)
                } for clause in clauses
            ]
        }

        html_string = render_to_string('leases/agreement_preview.html', context)
        html = HTML(string=html_string)

        with tempfile.NamedTemporaryFile(delete=True, suffix='.pdf') as output:
            html.write_pdf(output)
            output.flush()

            if lease.generated_agreement_pdf:
                lease.generated_agreement_pdf.delete(save=False)

            tenant_full_name = lease.tenant.get_full_name().replace(' ', '_')
            tenant_full_name = ''.join(
                c for c in tenant_full_name if c.isalnum() or c in ('_', '-'))
            cnic = lease.tenant.cnic.replace(
                '-', '') if lease.tenant.cnic else "nocnic"
            cnic = ''.join(c for c in cnic if c.isalnum())
            filename = f"{tenant_full_name}_{cnic}_{lease.id}_agreement_unsign.pdf"

            lease.generated_agreement_pdf.save(
                filename, File(output), save=True)

    return render(request, 'leases/edit_clause.html', {
        'lease': lease,
        'clauses': clauses,
        'agreement_date': lease.agreement_date,
    })


from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from .models import Lease

@login_required
def generate_agreement_pdf(request, pk):
    lease = get_object_or_404(Lease, pk=pk)

    # 1) build HTML using your real preview template (best)
    html = render_to_string("leases/agreement_preview.html", {"lease": lease})

    # 2) convert HTML -> PDF bytes (use your existing PDF tool)
    pdf_bytes = generate_agreement_pdf(html)   # use whatever you already use

    filename = f"LeaseAgreement_{lease.pk}.pdf"
    resp = HttpResponse(pdf_bytes, content_type="application/pdf")
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp



def send_agreement_email(request, pk):
    lease = get_object_or_404(Lease, pk=pk)
    if lease.generated_agreement_pdf:
        email = EmailMessage(
            f'Lease Agreement for {lease.unit.property.property_name}',
            'Please find attached the lease agreement.',
            'noreply@yourdomain.com',
            [lease.tenant.email, lease.unit.property.owner_email]
        )
        email.attach_file(lease.generated_agreement_pdf.path)
        email.send()

    return redirect('leases:lease_detail', pk=lease.pk)

from io import BytesIO
from django.http import FileResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Lease
from leases.utils import do_replace_placeholders  # you already use this in PDF view


from io import BytesIO
from django.http import FileResponse
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string
from weasyprint import HTML

def download_preview_pdf(request, lease_id):
    lease = get_object_or_404(Lease, pk=lease_id)
    clauses = lease.clauses.all().order_by('clause_number')

    # same logic preview uses
    for clause in clauses:
        clause.rendered_text = do_replace_placeholders(
            clause.template_text, lease
        )

    context = {
        'lease': lease,
        'clauses': clauses,
        'agreement_date': getattr(lease, 'agreement_date', None),
    }

    # 🔥 SINGLE SOURCE OF TRUTH
    html_string = render_to_string(
        'leases/agreement_preview.html',
        context
    )

    pdf_file = BytesIO()
    HTML(string=html_string).write_pdf(pdf_file)
    pdf_file.seek(0)

    filename = (
        f"{lease.tenant.get_full_name().replace(' ', '_')}-"
        f"{lease.tenant.cnic}-{lease.id}-agreement_not_sign.pdf"
    )

    return FileResponse(
        pdf_file,
        as_attachment=True,
        filename=filename,
        content_type="application/pdf"
    )

# leases/views_lease_photos.py


@login_required
def photos_page(request, lease_id):
    lease = get_object_or_404(Lease, pk=lease_id)
    return render(request, "leases/photos_page.html", {"lease": lease})

# leases/views.py
from decimal import Decimal
from math import ceil

from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import get_object_or_404, redirect
from django.urls import reverse, reverse_lazy
from django.views.generic import (
    ListView, CreateView, UpdateView, DeleteView
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from leases.models import Lease
from invoices.models import SecurityDepositTransaction
from invoices.forms import SecurityDepositTransactionForm
from invoices.services import security_deposit_totals

class LeaseSecurityMixin(LoginRequiredMixin):
    """
    Common bits for all security-deposit views, tied to a specific lease.
    URL uses <int:lease_pk>.
    """

    lease_url_kwarg = 'lease_pk'

    def dispatch(self, request, *args, **kwargs):
        self.lease = get_object_or_404(Lease, pk=kwargs[self.lease_url_kwarg])
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['lease'] = self.lease
        ctx['security_totals'] = security_deposit_totals(self.lease)
        return ctx

    def get_success_url(self):
        # After add/edit/delete go back to the security ledger list
        return reverse('leases:lease_security_list', kwargs={'lease_pk': self.lease.pk})
class SecurityDepositListView(LeaseSecurityMixin, ListView):
    model = SecurityDepositTransaction
    template_name = 'leases/security_deposit_list.html'
    context_object_name = 'deposit_transactions'

    def get_queryset(self):
        return (
            SecurityDepositTransaction.objects
            .filter(lease=self.lease)
            .order_by('date', 'id')
        )



    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        rows = []
        balance = Decimal('0.00')
        ZERO = Decimal('0.00')

        for tx in ctx['deposit_transactions']:
            amt = tx.amount or ZERO
            if tx.type == 'PAYMENT':
                signed_amt = amt
                balance += amt
            elif tx.type in ('REFUND', 'DAMAGE'):
                signed_amt = -amt
                balance -= amt
            elif tx.type == 'ADJUST':
                signed_amt = amt
                balance += amt
            else:
                signed_amt = amt

            rows.append({
                'obj': tx,
                'date': tx.date,
                'type': tx.get_type_display(),
                'raw_type': tx.type,
                'description': tx.notes or '',
                'amount': signed_amt,
                'balance': balance,
            })

        ctx['ledger_rows'] = rows

        # NEW: add sec totals (same as PaymentDetailView)
        ctx["sec_totals"] = security_deposit_totals(self.lease) if self.lease else {
            "required": 0,
            "paid_in": 0,
            "refunded": 0,
            "damages": 0,
            "balance_to_collect": 0,
            "currently_held": 0,
        }
        return ctx

class SecurityDepositCreateView(LeaseSecurityMixin, CreateView):
    model = SecurityDepositTransaction
    form_class = SecurityDepositTransactionForm
    template_name = 'leases/security_deposit_form.html'

    def form_valid(self, form):
        obj = form.save(commit=False)
        obj.lease = self.lease
        obj.save()
        messages.success(self.request, "Security deposit transaction added.")
        return redirect(self.get_success_url())
class SecurityDepositUpdateView(LeaseSecurityMixin, UpdateView):
    model = SecurityDepositTransaction
    form_class = SecurityDepositTransactionForm
    template_name = 'leases/security_deposit_form.html'
    pk_url_kwarg = 'pk'   # transaction id

    def get_queryset(self):
        # Only allow editing transactions for this lease
        return SecurityDepositTransaction.objects.filter(lease=self.lease)

    def form_valid(self, form):
        messages.success(self.request, "Security deposit transaction updated.")
        return super().form_valid(form)
class SecurityDepositDeleteView(LeaseSecurityMixin, DeleteView):
    model = SecurityDepositTransaction
    template_name = 'leases/security_deposit_confirm_delete.html'
    pk_url_kwarg = 'pk'

    def get_queryset(self):
        return SecurityDepositTransaction.objects.filter(lease=self.lease)

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "Security deposit transaction deleted.")
        return super().delete(request, *args, **kwargs)
from django.contrib.auth.decorators import login_required, permission_required
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse, reverse_lazy

from .models import Lease, DefaultClause
from .forms import DefaultClauseForm, LeaseClauseFormSet


# ---- DefaultClause UI (global defaults) ----

@login_required
@permission_required("leases.change_defaultclause", raise_exception=True)
def default_clause_list(request):
    clauses = DefaultClause.objects.order_by("clause_number")
    return render(request, "leases/default_clause_list.html", {
        "clauses": clauses,
    })


@login_required
@permission_required("leases.change_defaultclause", raise_exception=True)
def default_clause_create(request):
    if request.method == "POST":
        form = DefaultClauseForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect("leases:default_clause_list")
    else:
        form = DefaultClauseForm()

    return render(request, "leases/default_clause_form.html", {
        "form": form,
        "title": "Add Default Clause",
    })


@login_required
@permission_required("leases.change_defaultclause", raise_exception=True)
def default_clause_edit(request, pk):
    clause = get_object_or_404(DefaultClause, pk=pk)

    if request.method == "POST":
        form = DefaultClauseForm(request.POST, instance=clause)
        if form.is_valid():
            form.save()
            return redirect("leases:default_clause_list")
    else:
        form = DefaultClauseForm(instance=clause)

    return render(request, "leases/default_clause_form.html", {
        "form": form,
        "title": f"Edit Default Clause #{clause.clause_number}",
    })


# ---- Per-Lease clause editor ----

@login_required
@permission_required("leases.change_leaseagreementclause", raise_exception=True)
def lease_clause_edit(request, lease_id):
    lease = get_object_or_404(Lease, pk=lease_id)

    # Ensure clauses are initialized from DefaultClause if this is a new lease
    lease.initialize_clauses()

    if request.method == "POST":
        formset = LeaseClauseFormSet(request.POST, instance=lease)
        if formset.is_valid():
            # mark modified rows as customized
            for form in formset.forms:
                if form.has_changed() and form.instance.pk:
                    form.instance.is_customized = True
            formset.save()
            return redirect("leases:lease_clause_edit", lease_id=lease.id)
    else:
        formset = LeaseClauseFormSet(instance=lease)

    return render(request, "leases/lease_clause_edit.html", {
        "lease": lease,
        "formset": formset,
    })

from io import BytesIO
from django.http import FileResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Lease
from leases.templatetags.lease_tags import replace_placeholders

from django.http import FileResponse
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string
from io import BytesIO

def generate_agreement_docx(request, pk):
    lease = get_object_or_404(Lease, pk=pk)
    clauses = lease.clauses.all().order_by("clause_number")

    # Apply placeholder replacement exactly like PDF/preview
    for clause in clauses:
        clause.rendered_text = do_replace_placeholders(
            clause.template_text, lease
        )

    context = {
        "lease": lease,
        "clauses": clauses,
        "agreement_date": getattr(lease, "agreement_date", None),
    }

    # IMPORTANT: same template as preview & PDF
    html = render_to_string("leases/agreement_preview.html", context)

    # Convert HTML → DOCX (your fixed converter)
    docx_bytes = html_to_docx_bytes(html)

    bio = BytesIO(docx_bytes)
    bio.seek(0)

    filename = (
        f"{lease.tenant.get_full_name().replace(' ', '_')}-"
        f"{lease.tenant.cnic}-{lease.id}-agreement_not_sign.docx"
    )

    return FileResponse(
        bio,
        as_attachment=True,
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

from io import BytesIO
import re
from django.http import FileResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Lease
from leases.utils import do_replace_placeholders  # adjust if needed


_html_tag_re = re.compile(r"<[^>]+>")

def strip_html(s: str) -> str:
    # python-docx does not render HTML; remove tags
    return _html_tag_re.sub("", s or "")



from io import BytesIO
from django.http import FileResponse
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string

from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup, NavigableString, Tag

from .models import Lease
from leases.utils import do_replace_placeholders  # you already use this in PDF view
from io import BytesIO
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

from docx.shared import Inches

def set_doc_margins(doc, left=0.6, right=0.6, top=0.7, bottom=0.7):
    # values in inches
    section = doc.sections[0]
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)

def _set_doc_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)  # ✅ smaller font

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    





from bs4 import NavigableString, Tag

from io import BytesIO
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def _append_inline(paragraph, node, br_as_space=False):
    if isinstance(node, NavigableString):
        txt = str(node)
        if not txt or not txt.strip():
            return
        # normalize whitespace
        paragraph.add_run(" ".join(txt.split()))
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run(" " if br_as_space else "\n")
        return

    if node.name in ("strong", "b"):
        run = paragraph.add_run(node.get_text(" ", strip=True))
        run.bold = True
        return

    for child in node.children:
        _append_inline(paragraph, child, br_as_space=br_as_space)

def _new_p(doc, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p

from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_signature_block(doc, lease):
    # space before signature (like PDF)
    doc.add_paragraph("")
    doc.add_paragraph("")

    # 2 columns, 4 rows:
    # Row1: signature lines
    # Row2: Owner/Tenant
    # Row3: CNIC
    # Row4: blank line (small spacer)
    t = doc.add_table(rows=4, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    for r in t.rows:
        r.cells[0].width = Inches(3.1)
        r.cells[1].width = Inches(3.1)

    def set_cell(cell, lines, bold_prefix=False):
        # clear default paragraph
        cell.text = ""
        for i, line in enumerate(lines):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # allow "Owner: xxx" with bold "Owner:"
            if bold_prefix and ":" in line:
                prefix, rest = line.split(":", 1)
                r1 = p.add_run(prefix.strip() + ": ")
                r1.bold = True
                p.add_run(rest.strip())
            else:
                p.add_run(line)

    # row 0: signature lines
    set_cell(t.cell(0,0), ["_________________________"])
    set_cell(t.cell(0,1), ["_________________________"])

    # row 1: Owner/Tenant labels + names
    set_cell(t.cell(1,0), [f"Owner: {lease.unit.property.owner_name}"], bold_prefix=True)
    set_cell(t.cell(1,1), [f"Tenant: {lease.tenant.get_full_name()}"], bold_prefix=True)

    # row 2: CNICs
    set_cell(t.cell(2,0), [f"CNIC: {lease.unit.property.owner_cnic}"], bold_prefix=True)
    set_cell(t.cell(2,1), [f"CNIC: {lease.tenant.cnic}"], bold_prefix=True)

    # row 3: spacer row (optional)
    set_cell(t.cell(3,0), [""])
    set_cell(t.cell(3,1), [""])

    # Witness table (same structure)
    doc.add_paragraph("")  # small gap

    tw = doc.add_table(rows=3, cols=2)
    tw.alignment = WD_TABLE_ALIGNMENT.CENTER
    tw.autofit = False
    for r in tw.rows:
        r.cells[0].width = Inches(3.1)
        r.cells[1].width = Inches(3.1)

    set_cell(tw.cell(0,0), ["Witness 1: _________________________"], bold_prefix=True)
    set_cell(tw.cell(0,1), ["Witness 2: _________________________"], bold_prefix=True)
    set_cell(tw.cell(1,0), ["CNIC: _________________________"], bold_prefix=True)
    set_cell(tw.cell(1,1), ["CNIC: _________________________"], bold_prefix=True)
    set_cell(tw.cell(2,0), [""])
    set_cell(tw.cell(2,1), [""])

def html_to_docx_bytes(html: str,lease) -> bytes:
    soup = BeautifulSoup(html, "html.parser")
    container = soup.select_one("#agreement-doc") or soup.body or soup

    doc = Document()
    _set_doc_defaults(doc)          # set font size etc
    set_doc_margins(doc, left=0.6, right=0.6, top=0.7, bottom=0.7)
    add_page_number_footer(doc)     # Page X of Y

    # 1) Title
    title = container.select_one("h1,h2,h3")
    if title:
        p = _new_p(doc, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(title.get_text(" ", strip=True))
        r.bold = True

    # 2) Center line under title (the made at Islamabad line)
    # Use direct <p> children only (no nesting duplication)
    subtitle = container.select_one("p.text-center")
    if subtitle:
        p = _new_p(doc, WD_ALIGN_PARAGRAPH.CENTER)
        for child in subtitle.children:
            _append_inline(p, child, br_as_space=True)




    # 3) Parties section (split at AND properly)
    party_wrap = container.select_one(".parties-section")
    if party_wrap:
        # collect HTML inside parties-section
        # (not text) so we keep <strong> etc.
        party_html = str(party_wrap)

        # Remove outer wrapper tags so we only split the content
        party_soup = BeautifulSoup(party_html, "html.parser")
        # Prefer first <p> if present, else use the whole section
        first_p = party_soup.select_one(".parties-section p") or party_soup

        # Convert to HTML string
        html_body = "".join(str(x) for x in first_p.contents)

        # Split at AND (robust: handles <strong>AND</strong>, AND, <br>AND<br>, etc.)
        parts = re.split(r'(?i)\bAND\b', BeautifulSoup(html_body, "html.parser").get_text(" ", strip=False), maxsplit=1)

        # If split failed (no AND found), fallback to old behavior
        if len(parts) < 2:
            p = _new_p(doc)
            for child in first_p.children:
                _append_inline(p, child, br_as_space=True)
        else:
            left_text = parts[0].strip()
            right_text = parts[1].strip()

            # Party 1 paragraph
            p = _new_p(doc)
            p.add_run(" ".join(left_text.split()))

            # AND centered
            and_p = _new_p(doc, WD_ALIGN_PARAGRAPH.CENTER)
            r = and_p.add_run("AND")
            r.bold = True

            # Party 2 paragraph
            p = _new_p(doc)
            p.add_run(" ".join(right_text.split()))

        # Now render any WHEREAS paragraphs that are separate <p> after the first one
        # (if your template has them)
        extra_ps = party_wrap.select("p")[1:]  # remaining <p> after first
        for extra in extra_ps:
            txt = extra.get_text(" ", strip=True)
            if not txt:
                continue
            p = _new_p(doc)
            for child in extra.children:
                _append_inline(p, child, br_as_space=True)



    # 4) Clauses (tight, number + text same line)
    doc.add_paragraph("")   # ✅ blank line like PDF
    for clause_div in container.select(".clauses-section > .clause"):
        p = _new_p(doc, WD_ALIGN_PARAGRAPH.LEFT)

        strong = clause_div.find("strong")
        if strong:
            rn = p.add_run(strong.get_text(strip=True))
            rn.bold = True
            p.add_run(" ")
            strong.extract()

        for child in clause_div.children:
            _append_inline(p, child, br_as_space=False)

    # 5) Signature (table look + spacing)
    add_signature_block(doc, lease)

    # 6) Footer “Generated at”
    for ptag in container.select(".footer p"):
        txt = ptag.get_text(" ", strip=True)
        if not txt:
            continue
        p = _new_p(doc, WD_ALIGN_PARAGRAPH.CENTER)
        for child in ptag.children:
            _append_inline(p, child, br_as_space=True)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

def _add_field(paragraph, field_name: str):
    """Insert Word field e.g. PAGE / NUMPAGES"""
    run = paragraph.add_run()
    r = run._r

    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    r.append(fldBegin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {field_name} "
    r.append(instrText)

    fldSeparate = OxmlElement("w:fldChar")
    fldSeparate.set(qn("w:fldCharType"), "separate")
    r.append(fldSeparate)

    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    r.append(fldEnd)


def add_page_number_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run("Page ")
    _add_field(p, "PAGE")
    p.add_run(" of ")
    _add_field(p, "NUMPAGES")


from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def download_preview_docx(request, lease_id):
    lease = get_object_or_404(Lease, pk=lease_id)
    clauses = lease.clauses.all().order_by('clause_number')

    # same behavior as your PDF: attach rendered_text for clauses
    for clause in clauses:
        clause.rendered_text = do_replace_placeholders(clause.template_text, lease)

    context = {
        "lease": lease,
        "clauses": clauses,
        "agreement_date": getattr(lease, "agreement_date", None),
    }

    # IMPORTANT: render the SAME preview body you show on the left panel
    html = render_to_string("leases/agreement_preview.html", context)

    docx_bytes = html_to_docx_bytes(html,lease)
    bio = BytesIO(docx_bytes)
    bio.seek(0)

    filename = f"{lease.tenant.get_full_name().replace(' ', '_')}-{lease.tenant.cnic}-{lease.id}-agreement_not_sign.docx"
    return FileResponse(
        bio,
        as_attachment=True,
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# leases/views.py
from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect
from django.views.decorators.http import require_POST

@require_POST
def reset_clauses_from_default(request, pk):
    lease = get_object_or_404(Lease, pk=pk)
    default_tpl = LeaseTemplate.objects.filter(is_default=True).first()

    if not default_tpl:
        messages.error(request, "No default template found.")
        return redirect("leases:edit_clauses", pk=pk)

    # OPTION A: overwrite everything
    lease.clauses.all().delete()
    for i, clause_text in enumerate(default_tpl.clauses):
        LeaseAgreementClause.objects.create(
            lease=lease,
            clause_number=i + 1,
            template_text=clause_text,
            is_customized=False,
        )

    messages.success(request, "Clauses reset to default template.")
    return redirect("leases:edit_clauses", pk=pk)
