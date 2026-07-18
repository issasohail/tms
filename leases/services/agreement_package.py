from io import BytesIO
import base64
import mimetypes
import re

from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.html import conditional_escape
from weasyprint import DEFAULT_OPTIONS, Document, HTML
from weasyprint.layout import LayoutContext

from leases.models import AgreementSignatureTemplate
from tenants.models import Tenant
from leases.utils import do_replace_placeholders
from core.utils.identity import format_cnic, format_phone, normalize_cnic


class _QrExclusionShape:
    """A non-painted first-page shape used only by WeasyPrint's line layout."""

    style = {"float": "right"}

    def __init__(self, x, y, width, height):
        self.position_x = x
        self.position_y = y
        self._width = width
        self._height = height

    def margin_width(self):
        return self._width

    def margin_height(self):
        return self._height


class _QrExclusionLayoutContext(LayoutContext):
    def create_block_formatting_context(self):
        super().create_block_formatting_context()
        exclusion = getattr(self, "first_page_qr_exclusion", None)
        # The root formatting context is created immediately before WeasyPrint
        # assigns page number 1. Later pages retain the previous page number.
        if self.current_page is None and exclusion:
            self.excluded_shapes.append(_QrExclusionShape(*exclusion))


class _QrExclusionDocument(Document):
    @classmethod
    def _build_layout_context(cls, html, font_config, counter_style, options):
        context = super()._build_layout_context(
            html, font_config, counter_style, options
        )
        context.__class__ = _QrExclusionLayoutContext
        context.first_page_qr_exclusion = html.first_page_qr_exclusion
        return context


def _pdf(html, request, first_page_qr_exclusion=None):
    source = HTML(string=html, base_url=request.build_absolute_uri("/"))
    if not first_page_qr_exclusion:
        return source.write_pdf()
    source.first_page_qr_exclusion = first_page_qr_exclusion
    document = _QrExclusionDocument._render(
        source, font_config=None, counter_style=None,
        options=DEFAULT_OPTIONS.copy(),
    )
    return document.write_pdf()


def _period(lease, history=None):
    return (
        getattr(history, "start_date", None) or lease.start_date,
        getattr(history, "end_date", None) or lease.end_date,
    )


def _relationship_name(value):
    return str(value) if value else ""


def party_snapshot(lease, history=None):
    def data(person, relationship=""):
        return {
            "name": person.get_full_name() if person else "",
            "cnic": format_cnic(getattr(person, "cnic", "")) if person else "",
            "phone": format_phone(getattr(person, "phone", "")) if person else "",
            "relationship": _relationship_name(relationship),
        }

    witness1 = (getattr(history, "witness1_tenant", None) if history else None) or lease.witness1_tenant
    witness2 = (getattr(history, "witness2_tenant", None) if history else None) or lease.witness2_tenant
    occupants = [{
        "name": row.family_member.get_full_name(),
        "cnic": format_cnic(row.family_member.cnic),
        "relationship": str(row.relationship_type or row.relationship or ""),
    } for row in lease.family_members.select_related("family_member", "relationship_type").filter(lives_with_tenant=True)]

    return {
        "authorized_occupants": occupants,
        "proposer": data(lease.proposer, lease.proposer_relationship),
        "seconder": data(lease.seconder, lease.seconder_relationship),
        "witness1": data(witness1),
        "witness2": data(witness2),
    }


def _declaration_values(lease, history, parties):
    start_date, end_date = _period(lease, history)
    values = {
        "tenant_name": lease.tenant.get_full_name() or "________________",
        "tenant_cnic": format_cnic(lease.tenant.cnic) or "________________",
        "property_unit": f"{lease.unit.property} / {lease.unit}",
        "lease_start_date": start_date.strftime("%B %d, %Y") if start_date else "________________",
        "lease_end_date": end_date.strftime("%B %d, %Y") if end_date else "________________",
    }
    for role in ("proposer", "seconder"):
        row = parties[role]
        values.update({
            f"{role}_name": row["name"] or "________________",
            f"{role}_cnic": row["cnic"] or "________________",
            f"{role}_phone": row["phone"] or "________________",
            f"{role}_relationship": row["relationship"] or "________________",
        })
    return values


def _render_template_text(text, values):
    rendered = text or ""
    for key, value in values.items():
        rendered = rendered.replace("{{ " + key + " }}", str(value))
        rendered = rendered.replace("{{" + key + "}}", str(value))
    return rendered



def _agreement_layout_settings():
    config = AgreementSignatureTemplate.current()
    return {
        "config": config,
        "first_top": float(getattr(config, "legal_first_page_top_reserve", 4.80) or 4.80),
        "qr_width": float(getattr(config, "legal_qr_reserve_width", 4.00) or 0),
        "qr_height": float(getattr(config, "legal_qr_reserve_height", 2.00) or 0),
        "identity_bottom": float(getattr(config, "legal_identity_bottom_reserve", 3.10) or 3.10),
        "clause_spacing": float(getattr(config, "legal_clause_spacing", 5.00) or 0),
    }


def _field_bytes(field):
    try:
        if not field:
            return b""
        with field.open("rb") as source:
            return source.read()
    except (OSError, ValueError, AttributeError):
        return b""


def _draw_fitted_image(canvas_obj, payload, x, y, width, height):
    if not payload:
        return False
    try:
        from reportlab.lib.utils import ImageReader
        image = ImageReader(BytesIO(payload))
        image_width, image_height = image.getSize()
        if not image_width or not image_height:
            return False
        scale = min(width / image_width, height / image_height)
        draw_width = image_width * scale
        draw_height = image_height * scale
        canvas_obj.drawImage(
            image,
            x + (width - draw_width) / 2,
            y + (height - draw_height) / 2,
            width=draw_width,
            height=draw_height,
            preserveAspectRatio=True,
            mask="auto",
        )
        return True
    except Exception:
        return False


def _identity_overlay_page(width, height, people, reserve_inches):
    from reportlab.pdfgen import canvas
    packet = BytesIO()
    pdf = canvas.Canvas(packet, pagesize=(float(width), float(height)))

    reserve_height = max(2.0, min(5.0, float(reserve_inches))) * 72
    left = 0.55 * 72
    right = 0.55 * 72
    footer_clearance = 28
    panel_y = footer_clearance
    panel_height = max(112, reserve_height - footer_clearance - 7)
    gap = 6
    available = float(width) - left - right
    card_width = (available - gap * 3) / 4
    header_height = 24
    inner_gap = 4
    image_height = max(38, (panel_height - header_height - inner_gap * 3) / 2)

    pdf.setLineWidth(0.55)
    for index, person in enumerate((people or [])[:4]):
        x = left + index * (card_width + gap)
        pdf.rect(x, panel_y, card_width, panel_height, stroke=1, fill=0)

        role = str(person.get("role") or "")
        name = str(person.get("name") or "________________")
        pdf.setFont("Helvetica-Bold", 7.5)
        pdf.drawCentredString(x + card_width / 2, panel_y + panel_height - 10, role)
        pdf.setFont("Helvetica", 5.8)
        clipped_name = name if len(name) <= 31 else name[:28] + "..."
        pdf.drawCentredString(x + card_width / 2, panel_y + panel_height - 19, clipped_name)

        source_person = person.get("person")
        fields = (
            ("CNIC Front", getattr(source_person, "cnic_front", None) if source_person else None),
            ("CNIC Back", getattr(source_person, "cnic_back", None) if source_person else None),
        )
        image_top = panel_y + panel_height - header_height - inner_gap
        for image_index, (label, field) in enumerate(fields):
            image_y = image_top - (image_index + 1) * image_height - image_index * inner_gap
            image_x = x + inner_gap
            image_width = card_width - inner_gap * 2
            pdf.setLineWidth(0.35)
            pdf.rect(image_x, image_y, image_width, image_height, stroke=1, fill=0)
            if not _draw_fitted_image(pdf, _field_bytes(field), image_x + 2, image_y + 2, image_width - 4, image_height - 4):
                pdf.setFont("Helvetica", 6)
                pdf.drawCentredString(image_x + image_width / 2, image_y + image_height / 2 - 2, label)

    pdf.save()
    packet.seek(0)
    from pypdf import PdfReader
    return PdfReader(packet).pages[0]


def _pin_identity_cards_to_second_page(pdf_bytes, lease, history, reserve_inches):
    """Keep the four CNIC cards at the bottom of Legal agreement page 2.

    Agreement pages reserve the same bottom region from page 2 onward. The
    cards are then drawn as a PDF overlay, so longer clause text flows to page
    3 instead of pushing the cards away from page 2.
    """
    try:
        from pypdf import PageObject, PdfReader, PdfWriter
    except ImportError as exc:
        raise RuntimeError("pypdf is required for fixed Legal-page CNIC placement.") from exc

    reader = PdfReader(BytesIO(pdf_bytes))
    pages = list(reader.pages)
    if not pages:
        return pdf_bytes
    while len(pages) < 2:
        pages.append(PageObject.create_blank_page(
            width=pages[0].mediabox.width,
            height=pages[0].mediabox.height,
        ))

    page = pages[1]
    people = identity_context(lease, history).get("identity_people", [])
    overlay = _identity_overlay_page(
        page.mediabox.width,
        page.mediabox.height,
        people,
        reserve_inches,
    )
    page.merge_page(overlay, over=True)

    writer = PdfWriter()
    for item in pages:
        writer.add_page(item)
    out = BytesIO()
    writer.write(out)
    return out.getvalue()


def _agreement_signature_footer_page(width, height, lease, y, right_boundary=None):
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase.pdfmetrics import stringWidth

    packet = BytesIO()
    pdf = canvas.Canvas(packet, pagesize=(float(width), float(height)))
    left = 0.55 * 72
    boundary = float(right_boundary or (float(width) - left))
    column_gap = 12
    line_width = (boundary - left - column_gap) / 2
    right = left + line_width + column_gap
    owner_name = str(getattr(lease.unit.property, "owner_name", "") or "________________")
    tenant_name = lease.tenant.get_full_name() or "________________"
    owner_cnic = format_cnic(getattr(lease.unit.property, "owner_cnic", "")) or "________________"
    tenant_cnic = format_cnic(getattr(lease.tenant, "cnic", "")) or "________________"

    for x, role, name, cnic in (
        (left, "Owner", owner_name, owner_cnic),
        (right, "Tenant", tenant_name, tenant_cnic),
    ):
        pdf.setLineWidth(0.55)
        pdf.setFont("Helvetica-Bold", 7.5)
        pdf.drawString(x, y + 16, f"{role} Signature:")
        signature_start = x + 0.78 * 72
        pdf.line(signature_start, y + 15, x + line_width, y + 15)
        details = f"{role}: {name}    CNIC: {cnic}"
        details_font_size = 6.5
        while details_font_size > 4.5 and stringWidth(
            details, "Helvetica", details_font_size
        ) > line_width:
            details_font_size -= 0.25
        pdf.setFont("Helvetica", details_font_size)
        pdf.drawString(x, y + 4, details)

    pdf.save()
    packet.seek(0)
    from pypdf import PdfReader
    return PdfReader(packet).pages[0]


def _add_agreement_signature_footers(
    pdf_bytes, lease, identity_bottom_reserve, qr_reserve_width
):
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(BytesIO(pdf_bytes))
    pages = list(reader.pages)
    if len(pages) <= 1:
        return pdf_bytes
    for index, page in enumerate(pages[:-1]):
        if index == 0:
            y = 31
            qr_width = max(0.0, float(qr_reserve_width or 0)) * 72
            right_boundary = float(page.mediabox.width) - 0.55 * 72 - qr_width - 8
        elif index == 1:
            y = float(identity_bottom_reserve) * 72 + 8
            right_boundary = None
        else:
            y = 31
            right_boundary = None
        overlay = _agreement_signature_footer_page(
            page.mediabox.width,
            page.mediabox.height,
            lease,
            y,
            right_boundary=right_boundary,
        )
        page.merge_page(overlay, over=True)
    writer = PdfWriter()
    for page in pages:
        writer.add_page(page)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def _agreement_page_count(pdf_bytes):
    from pypdf import PdfReader
    return len(PdfReader(BytesIO(pdf_bytes)).pages)


def _add_first_page_qr_reserve_box(pdf_bytes, reserve_width, reserve_height):
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas

    width_inches = max(0.0, float(reserve_width or 0))
    height_inches = max(0.0, float(reserve_height or 0))
    if not width_inches or not height_inches:
        return pdf_bytes

    reader = PdfReader(BytesIO(pdf_bytes))
    pages = list(reader.pages)
    if not pages:
        return pdf_bytes

    first_page = pages[0]
    page_width = float(first_page.mediabox.width)
    page_height = float(first_page.mediabox.height)
    margin = 0.55 * 72
    box_width = min(width_inches * 72, page_width - margin * 2)
    box_height = min(height_inches * 72, page_height - margin * 2)
    x = page_width - margin - box_width
    y = margin

    packet = BytesIO()
    overlay_canvas = canvas.Canvas(packet, pagesize=(page_width, page_height))
    overlay_canvas.setLineWidth(0.55)
    overlay_canvas.rect(x, y, box_width, box_height, stroke=1, fill=0)
    overlay_canvas.save()
    packet.seek(0)
    overlay = PdfReader(packet).pages[0]
    first_page.merge_page(overlay, over=True)

    writer = PdfWriter()
    for page in pages:
        writer.add_page(page)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def agreement_pdf(request, lease, history, clauses):
    for clause in clauses:
        clause.rendered_text = do_replace_placeholders(clause.template_text, lease)
    legal_page = bool(getattr(history, "print_on_legal_page", False))
    layout = _agreement_layout_settings()
    first_bottom_reserve = 0.95
    later_bottom_reserve = layout["identity_bottom"] + 0.72
    css_px_per_inch = 96.0
    legal_page_width = 8.5
    legal_page_height = 14.0
    qr_flow_height = max(
        0.0, layout["qr_height"] - (first_bottom_reserve - 0.55)
    )
    qr_exclusion = None
    if legal_page and layout["qr_width"] > 0 and qr_flow_height > 0:
        qr_text_gutter = 0.12
        qr_exclusion = (
            (
                legal_page_width - 0.55 - layout["qr_width"] - qr_text_gutter
            ) * css_px_per_inch,
            (legal_page_height - 0.55 - layout["qr_height"]) * css_px_per_inch,
            (layout["qr_width"] + qr_text_gutter) * css_px_per_inch,
            qr_flow_height * css_px_per_inch,
        )

    def render(spacing):
        html = render_to_string(
            "leases/agreement_preview.html",
            {
                "lease": lease,
                "history": history,
                "clauses": clauses,
                "agreement_date": getattr(history, "agreement_date", None)
                or getattr(history, "start_date", lease.start_date),
                "legal_page": legal_page,
                "legal_first_page_top_reserve": layout["first_top"],
                "legal_qr_reserve_width": layout["qr_width"],
                "legal_qr_reserve_height": layout["qr_height"],
                "legal_identity_bottom_reserve": layout["identity_bottom"],
                "legal_first_page_bottom_reserve": first_bottom_reserve,
                "legal_later_page_bottom_reserve": later_bottom_reserve,
                "legal_clause_spacing": spacing,
            },
            request=request,
        )
        return _pdf(html, request, first_page_qr_exclusion=qr_exclusion)

    if legal_page:
        pdf_bytes = None
        requested_spacing = max(0.0, min(12.0, layout["clause_spacing"]))
        spacing_candidates = []
        current_spacing = requested_spacing
        while current_spacing >= 0:
            spacing_candidates.append(round(current_spacing, 2))
            current_spacing -= 1
        if spacing_candidates[-1] != 0:
            spacing_candidates.append(0)
        for spacing in spacing_candidates:
            candidate = render(spacing)
            pdf_bytes = candidate
            if _agreement_page_count(candidate) <= 3:
                break
    else:
        pdf_bytes = render(0)
    if legal_page:
        pdf_bytes = _add_first_page_qr_reserve_box(
            pdf_bytes,
            layout["qr_width"],
            layout["qr_height"],
        )
        pdf_bytes = _pin_identity_cards_to_second_page(
            pdf_bytes, lease, history, layout["identity_bottom"]
        )
        pdf_bytes = _add_agreement_signature_footers(
            pdf_bytes,
            lease,
            layout["identity_bottom"],
            layout["qr_width"],
        )
    return pdf_bytes

def _inspection_template_name_for_lease(lease):
    unit = lease.unit
    bedrooms = getattr(unit, "bedrooms", None)
    if bedrooms is not None:
        try:
            return "Room" if int(bedrooms) <= 1 else "Apartment Standard"
        except (TypeError, ValueError):
            pass
    type_text = " ".join(str(value or "") for value in (
        getattr(unit, "unit_type", ""), getattr(unit, "type", ""),
        getattr(unit.property, "type", ""), getattr(unit.property, "property_type", ""),
    )).lower()
    return "Room" if "single room" in type_text or type_text.strip() == "room" else "Apartment Standard"


def _create_default_move_in_inspection(request, lease):
    from leases.models_inspections import InspectionTemplate, InspectionType, LeaseInspection

    inspection_type = InspectionType.objects.filter(name__iexact="Move In", active=True).first()
    if inspection_type is None:
        raise RuntimeError("Active inspection type 'Move In' was not found.")
    template_name = _inspection_template_name_for_lease(lease)
    inspection_template = InspectionTemplate.objects.filter(name__iexact=template_name, active=True).first()
    if inspection_template is None:
        raise RuntimeError(f"Active inspection template '{template_name}' was not found.")
    user = request.user if getattr(request.user, "is_authenticated", False) else None
    inspection = LeaseInspection.objects.create(
        lease=lease, property=lease.unit.property, unit=lease.unit, tenant=lease.tenant,
        inspection_type=inspection_type, inspection_template=inspection_template,
        inspection_date=timezone.localdate(), inspector=None, inspector_name="",
        status=LeaseInspection.STATUS_DRAFT, created_by=user,
    )
    inspection.snapshot_template_items()
    inspection.add_audit("created automatically for agreement package", user, {
        "template": inspection_template.name, "type": inspection_type.name,
    })
    return inspection


def get_or_create_inspection(request, lease):
    inspection = (
        lease.inspections.select_related(
            "inspection_type", "inspection_template", "property", "unit", "tenant",
            "inspector", "approved_by", "created_by",
        )
        .prefetch_related("details__photos", "meter_readings", "keys", "appliances", "damage_charges__invoice")
        .order_by("-inspection_date", "-id")
        .first()
    )
    return inspection or _create_default_move_in_inspection(request, lease)


def inspection_pdf(request, lease):
    from leases.views_inspections import _inspection_pdf_bytes
    return _inspection_pdf_bytes(request, get_or_create_inspection(request, lease))


def police_context(lease):
    property_obj = lease.unit.property
    owner_person = _owner_tenant(property_obj)
    owner_photo = getattr(property_obj, "owner_photo", None) or getattr(owner_person, "photo", None)
    tenant_photo = getattr(lease.tenant, "photo", None)
    return {
        "lease": lease, "tenant": lease.tenant, "property": property_obj,
        "owner_photo_url": _file_data_uri(owner_photo),
        "tenant_photo_url": _file_data_uri(tenant_photo),
        "unit": lease.unit,
        "family_members": lease.family_members.select_related("family_member", "relationship_type"),
        "vehicles": lease.vehicles.filter(is_active=True).select_related("vehicle_type"),
        "generated_at": timezone.localtime(),
    }


def police_pdf(request, lease):
    return _pdf(render_to_string("leases/police_verification_summary_pdf.html", police_context(lease), request=request), request)


def _safe_file_url(field):
    try:
        return field.url if field else ""
    except (ValueError, AttributeError):
        return ""


def _file_data_uri(field):
    """Embed local media in generated PDFs so reverse-proxy/media URLs are not required."""
    try:
        if not field:
            return ""
        with field.open("rb") as source:
            payload = base64.b64encode(source.read()).decode("ascii")
        mime = mimetypes.guess_type(getattr(field, "name", ""))[0] or "image/jpeg"
        return f"data:{mime};base64,{payload}"
    except (OSError, ValueError, AttributeError):
        return ""


def _normalise_cnic(value):
    return normalize_cnic(value)


def _owner_tenant(property_obj):
    """Find an existing Tenant record for the property owner by CNIC, without adding fields."""
    owner_cnic = _normalise_cnic(getattr(property_obj, "owner_cnic", ""))
    if not owner_cnic:
        return None
    # CNIC values may be stored with or without hyphens, so compare normalized values.
    for person in Tenant.objects.exclude(cnic="").only("id", "cnic", "photo", "cnic_front", "cnic_back"):
        if _normalise_cnic(person.cnic) == owner_cnic:
            return person
    return None


def identity_context(lease, history=None):
    witness1 = getattr(history, "witness1_tenant", None) if history else lease.witness1_tenant
    witness2 = getattr(history, "witness2_tenant", None) if history else lease.witness2_tenant
    tenant = lease.tenant
    property_obj = lease.unit.property
    owner_person = _owner_tenant(property_obj)

    def tenant_person(role, person, show_phone=False):
        return {
            "role": role,
            "name": person.get_full_name() if person else "",
            "cnic": format_cnic(getattr(person, "cnic", "")),
            "phone": format_phone(getattr(person, "phone", "")),
            "show_phone": show_phone,
            "front_url": _file_data_uri(getattr(person, "cnic_front", None)) if person else "",
            "back_url": _file_data_uri(getattr(person, "cnic_back", None)) if person else "",
            "person": person,
        }

    owner_row = {
        "role": "Owner",
        "name": getattr(property_obj, "owner_name", "") or "",
        "cnic": format_cnic(getattr(property_obj, "owner_cnic", "")),
        "phone": format_phone(getattr(property_obj, "owner_phone", "")),
        "show_phone": False,
        "front_url": _file_data_uri(getattr(owner_person, "cnic_front", None)) if owner_person else "",
        "back_url": _file_data_uri(getattr(owner_person, "cnic_back", None)) if owner_person else "",
        "person": owner_person,
    }
    return {"lease": lease, "history": history, "identity_people": [
        owner_row,
        tenant_person("Tenant", tenant),
        tenant_person("Witness 1", witness1, True),
        tenant_person("Witness 2", witness2, True),
    ]}


def identity_pdf(request, lease, history=None):
    return _pdf(render_to_string("leases/agreement_identity_documents.html", identity_context(lease, history), request=request), request)


def _bold(value):
    return "<strong>" + conditional_escape(value or "________________") + "</strong>"


def _declaration_sections(lease, history, parties):
    v = _declaration_values(lease, history, parties)
    common4 = f'If any dispute, misunderstanding, payment issue, complaint, or other matter arises between the Tenant and the Management/Landlord, I shall, when reasonably requested, be willing to assist in good faith in communicating with the parties and helping them reach an amicable resolution. I confirm that I am giving this declaration voluntarily and authorize the Management/Landlord to contact me for verification of my identity, relationship with the Tenant, and the information provided in this declaration. I understand that this declaration is a personal reference only and does not make me financially liable for the Tenant’s obligations unless I separately sign a written guarantee.'
    sections=[]
    for role, heading in (("proposer","Proposer Declaration"),("seconder","Seconder Declaration")):
        p=parties[role]
        para3 = ('Based on my personal knowledge of the Tenant’s character, conduct, and financial responsibility, I believe that the Tenant is trustworthy, responsible, suitable for tenancy, and capable of paying the agreed rent, utility charges, and other lawful amounts on time. I recommend and vouch for the Tenant’s suitability for this tenancy.' if role=="proposer" else 'Based on my personal knowledge of the Tenant’s character, conduct, and financial responsibility, I believe that the Tenant is trustworthy, responsible, suitable for tenancy, and capable of paying the agreed rent, utility charges, and other lawful amounts on time. I support and second the proposal for this tenancy.')
        sections.append({"heading":heading,"party":p,"paragraphs":[f'I, {_bold(v[role+"_name"])}, holding CNIC No. {_bold(v[role+"_cnic"])}, and having the relationship of {_bold(v[role+"_relationship"])} with the Tenant, hereby declare that I personally know {_bold(v["tenant_name"])}, holding CNIC No. {_bold(v["tenant_cnic"])}.', f'I understand that the Tenant is entering into a tenancy for {_bold(v["property_unit"])}, for the period from {_bold(v["lease_start_date"])} to {_bold(v["lease_end_date"])}.', para3, common4]})
    return sections


def _render_declaration_html(text, values):
    rendered = text or ""
    for key, value in values.items():
        replacement = _bold(str(value))
        rendered = rendered.replace("{{ " + key + " }}", replacement)
        rendered = rendered.replace("{{" + key + "}}", replacement)
    return rendered


def _configured_declaration_sections(config, lease, history, values, parties):
    sections = []
    for role, heading, template_text in (
        ("proposer", "Proposer Declaration", config.proposer_declaration),
        ("seconder", "Seconder Declaration", config.seconder_declaration),
    ):
        rendered = _render_declaration_html(template_text, values)
        paragraphs = [
            paragraph.strip()
            for paragraph in re.split(r"(?:\r?\n){2,}", rendered)
            if paragraph.strip()
        ]
        if not paragraphs:
            return _declaration_sections(lease, history, parties)
        sections.append({
            "heading": heading,
            "party": parties[role],
            "paragraphs": paragraphs,
        })
    return sections


def signature_context(lease, history=None, snapshot=None):
    parties = snapshot or party_snapshot(lease, history)
    config = AgreementSignatureTemplate.current()
    values = _declaration_values(lease, history, parties)
    return {
        "lease": lease, "history": history, "parties": parties,
        "signature_config": config,
        "proposer_declaration": _render_template_text(config.proposer_declaration, values),
        "seconder_declaration": _render_template_text(config.seconder_declaration, values),
        "generated_at": timezone.now(),
        "declaration_sections": _configured_declaration_sections(
            config, lease, history, values, parties
        ),
    }


def signature_pdf(request, lease, history=None, snapshot=None):
    return _pdf(render_to_string("leases/proposer_seconder_declaration.html", signature_context(lease, history, snapshot), request=request), request)


def _clean_filename_part(value):
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    text = re.sub(r"[^A-Za-z0-9._ -]+", "", text)
    return text.replace(" ", "_").strip("._-") or "blank"


def _package_labels(lease, history):
    start_date, end_date = _period(lease, history)
    tenant_name = lease.tenant.get_full_name() or "Tenant"
    property_name = getattr(lease.unit.property, "property_name", "") or str(lease.unit.property)
    unit_name = getattr(lease.unit, "unit_number", "") or str(lease.unit)
    start_text = start_date.strftime("%Y-%m-%d") if start_date else "blank"
    end_text = end_date.strftime("%Y-%m-%d") if end_date else "blank"
    title = "Lease_{}-{}-{}-{}To{}".format(
        _clean_filename_part(tenant_name),
        _clean_filename_part(property_name),
        _clean_filename_part(unit_name),
        start_text,
        end_text,
    )
    center = "{}_{}_{}-{} to {}".format(
        tenant_name,
        property_name,
        unit_name,
        start_text,
        end_text,
    )
    return title, center


def _footer_overlay(width, height, left_text, center_text, right_text):
    from reportlab.pdfgen import canvas
    packet = BytesIO()
    pdf = canvas.Canvas(packet, pagesize=(float(width), float(height)))
    pdf.setFont("Helvetica", 7.5)
    y = 16
    pdf.drawString(28, y, left_text)
    pdf.drawCentredString(float(width) / 2, y, center_text)
    pdf.drawRightString(float(width) - 28, y, right_text)
    pdf.save()
    packet.seek(0)
    from pypdf import PdfReader
    return PdfReader(packet).pages[0]


def merge_pdfs(parts, lease=None, history=None):
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError as exc:
        raise RuntimeError("pypdf is required to merge the agreement package.") from exc

    pages = []
    for part in parts:
        pages.extend(PdfReader(BytesIO(part)).pages)

    total = len(pages)
    title, center_text = _package_labels(lease, history) if lease else ("Lease Agreement", "")
    timestamp = timezone.localtime().strftime("%Y-%m-%d %H:%M")
    writer = PdfWriter()
    for index, page in enumerate(pages, 1):
        width = page.mediabox.width
        height = page.mediabox.height
        overlay = _footer_overlay(
            width,
            height,
            timestamp,
            center_text,
            f"Page {index} of {total}",
        )
        page.merge_page(overlay, over=True)
        writer.add_page(page)
    writer.add_metadata({"/Title": title, "/Subject": center_text})
    out = BytesIO()
    writer.write(out)
    return out.getvalue()


def _package_basename(lease, history):
    title, _ = _package_labels(lease, history)
    return title

def build_package(request, lease, history, clauses):
    components = []
    try:
        components.append(agreement_pdf(request, lease, history, clauses))
    except Exception as exc:
        raise RuntimeError(f"Agreement PDF generation failed: {exc}") from exc
    try:
        components.append(inspection_pdf(request, lease))
    except Exception as exc:
        raise RuntimeError(f"Agreement generated, but inspection PDF generation failed: {exc}") from exc
    try:
        components.append(police_pdf(request, lease))
    except Exception as exc:
        raise RuntimeError(f"Agreement and inspection generated, but police report generation failed: {exc}") from exc
    try:
        components.append(signature_pdf(request, lease, history))
    except Exception as exc:
        raise RuntimeError(f"Signature page generation failed: {exc}") from exc
    merged = merge_pdfs(components, lease=lease, history=history)
    filename = _package_basename(lease, history) + ".pdf"
    # Generated agreements are reproducible downloads. Do not persist them in
    # Lease Documents; only user-uploaded signed documents should be retained.
    return merged, filename, None


# ----------------------------- DOCX package -----------------------------

def _set_cell_text(cell, text, bold=False, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = 0
    run = p.add_run(str(text or ""))
    run.bold = bold
    from docx.shared import Pt
    run.font.size = Pt(size)


def _add_heading(doc, text, size=14):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_identity_docx(doc, lease, history):
    from docx.shared import Inches, Pt
    doc.add_page_break(); _add_heading(doc, "IDENTITY DOCUMENTS")
    people = identity_context(lease, history)["identity_people"]
    table = doc.add_table(rows=2, cols=2); table.style="Table Grid"
    for idx, person in enumerate(people):
        cell=table.cell(idx//2, idx%2); cell.text=""
        p=cell.paragraphs[0]; r=p.add_run(person["role"]); r.bold=True; r.font.size=Pt(10)
        for label,key in (("Full Name","name"),("CNIC","cnic")):
            q=cell.add_paragraph(); q.paragraph_format.space_after=Pt(1); q.add_run(label+": ").bold=True; q.add_run(person[key] or "________________").bold=True
        if person["show_phone"]:
            q=cell.add_paragraph(); q.paragraph_format.space_after=Pt(1); q.add_run("Phone: ").bold=True; q.add_run(person["phone"] or "________________").bold=True
        for label,url_key in (("CNIC Front","front_url"),("CNIC Back","back_url")):
            q=cell.add_paragraph(); q.add_run(label).bold=True
            url=person[url_key]
            source_person = person.get("person")
            field = getattr(source_person, "cnic_front" if "Front" in label else "cnic_back", None) if source_person else None
            try:
                if field and field.path: cell.add_paragraph().add_run().add_picture(field.path, width=Inches(2.9))
                else: cell.add_paragraph("[                                                ]")
            except (ValueError, OSError, AttributeError): cell.add_paragraph("[                                                ]")


def _add_inspection_docx(doc, inspection):
    from docx.shared import Pt
    doc.add_page_break()
    _add_heading(doc, "Inspection Sheet")
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    values = [
        ("Lease", f"#{inspection.lease_id}"), ("Type", str(inspection.inspection_type)),
        ("Property / Unit", f"{inspection.property} / {inspection.unit}"),
        ("Date", inspection.inspection_date.strftime("%B %d, %Y")),
    ]
    for idx, (label, value) in enumerate(values):
        row, col = divmod(idx, 2)
        _set_cell_text(meta.cell(row, col * 2), label, True)
        _set_cell_text(meta.cell(row, col * 2 + 1), value)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for i, text in enumerate(("SN", "Category", "Item", "Qty", "Condition", "Remarks")):
        _set_cell_text(table.rows[0].cells[i], text, True)
    for idx, item in enumerate(inspection.details.all(), 1):
        cells = table.add_row().cells
        vals = (idx, item.category, item.item_name, item.quantity, item.status_name, item.remarks)
        for i, value in enumerate(vals):
            _set_cell_text(cells[i], value)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run("Inspector: ____________________    Tenant Signature: ____________________    Date: ____________________")


def _add_police_docx(doc, lease):
    doc.add_page_break()
    _add_heading(doc, "Police Verification Report")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Tenant", lease.tenant.get_full_name()), ("Tenant CNIC", format_cnic(lease.tenant.cnic)),
        ("Phone", format_phone(lease.tenant.phone)), ("Property / Unit", f"{lease.unit.property} / {lease.unit}"),
        ("Lease Period", f"{lease.start_date} to {lease.end_date}"),
        ("Address", lease.tenant.address or ""),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], label, True, 9)
        _set_cell_text(cells[1], value, False, 9)
    doc.add_paragraph("Family Members", style=None).runs[0].bold = True
    fam = doc.add_table(rows=1, cols=4); fam.style = "Table Grid"
    for i, text in enumerate(("SN", "Name", "CNIC", "Relationship")):
        _set_cell_text(fam.rows[0].cells[i], text, True)
    for idx, row in enumerate(lease.family_members.select_related("family_member", "relationship_type"), 1):
        cells = fam.add_row().cells
        vals = (idx, row.family_member.get_full_name(), format_cnic(row.family_member.cnic), str(row.relationship_type or row.relationship or ""))
        for i, value in enumerate(vals): _set_cell_text(cells[i], value)
    doc.add_paragraph("Verification Remarks: ______________________________________________________________")
    doc.add_paragraph("Police Station / Officer: __________________________    Date: __________________________")


def _add_signature_docx(doc, context):
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(context["signature_config"].heading or "Proposer and Seconder Declaration")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    def shade(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    for section in context["declaration_sections"]:
        role = "proposer" if section["heading"].startswith("Proposer") else "seconder"
        title_text = section["heading"]
        heading_table = doc.add_table(rows=1, cols=1)
        heading_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        heading_table.autofit = False
        heading_table.columns[0].width = Inches(7.35)
        cell = heading_table.cell(0, 0)
        cell.width = Inches(7.35)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "EDEDED")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title_text)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(11.5)

        from bs4 import BeautifulSoup
        for paragraph_html in section["paragraphs"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(7)
            p.paragraph_format.line_spacing = 1.3
            soup = BeautifulSoup(paragraph_html, "html.parser")
            for node in soup.contents:
                if getattr(node, "name", None) == "strong":
                    r = p.add_run(node.get_text())
                    r.bold = True
                else:
                    r = p.add_run(str(node))
                r.font.name = "Times New Roman"
                r.font.size = Pt(10.5)

        party = context["parties"][role]
        table = doc.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = (2.15, 1.65, 1.55, 2.00)
        labels = ("Full Name", "CNIC", "Phone Number", "Relationship to Tenant")
        vals = (
            party["name"] or "________________",
            party["cnic"] or "________________",
            party["phone"] or "________________",
            party["relationship"] or "________________",
        )
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
        for i, label in enumerate(labels):
            shade(table.cell(0, i), "F2F2F2")
            _set_cell_text(table.cell(0, i), label, True, 8.5)
        for i, value in enumerate(vals):
            _set_cell_text(table.cell(1, i), value, False, 9.5)

        sig = doc.add_paragraph()
        sig.paragraph_format.space_before = Pt(8)
        sig.paragraph_format.space_after = Pt(6)
        if context["signature_config"].show_thumb_impression:
            text = "Signature: __________________    Date: ______________    Thumb Impression: ______________"
        else:
            text = "Signature: ____________________________    Date: ____________________________"
        r = sig.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)


def build_docx_package(request, lease, history, clauses):
    # Reuse the current agreement converter, then append the three package components.
    for clause in clauses:
        clause.rendered_text = do_replace_placeholders(clause.template_text, lease)
    layout = _agreement_layout_settings()
    html = render_to_string("leases/agreement_preview.html", {
        "lease": lease, "history": history, "clauses": clauses,
        "agreement_date": getattr(history, "agreement_date", None) or getattr(history, "start_date", lease.start_date),
        "legal_page": bool(getattr(history, "print_on_legal_page", False)),
        "legal_first_page_top_reserve": layout["first_top"],
        "legal_qr_reserve_width": layout["qr_width"],
        "legal_qr_reserve_height": layout["qr_height"],
        "legal_identity_bottom_reserve": layout["identity_bottom"],
    }, request=request)
    from leases.views import html_to_docx_bytes
    from docx import Document
    try:
        agreement_bytes = html_to_docx_bytes(html, lease, history=history)
        doc = Document(BytesIO(agreement_bytes))
    except Exception as exc:
        raise RuntimeError(f"Agreement Word generation failed: {exc}") from exc
    try:
        _add_inspection_docx(doc, get_or_create_inspection(request, lease))
    except Exception as exc:
        raise RuntimeError(f"Agreement generated, but inspection Word section failed: {exc}") from exc
    try:
        _add_police_docx(doc, lease)
    except Exception as exc:
        raise RuntimeError(f"Agreement and inspection generated, but police report Word section failed: {exc}") from exc
    try:
        _add_signature_docx(doc, signature_context(lease, history))
    except Exception as exc:
        raise RuntimeError(f"Proposer/seconder Word declaration generation failed: {exc}") from exc
    out = BytesIO(); doc.save(out); content = out.getvalue()
    filename = _package_basename(lease, history) + ".docx"
    # Keep generated Word packages download-only for the same reason as PDFs.
    return content, filename, None
